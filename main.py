"""
╔══════════════════════════════════════════════════════════════════╗
║           JARVIS ULTIMATE 2026 — main.py  v4.0                   ║
║                                                                  ║
║  ИСПРАВЛЕНИЯ v4.0:                                               ║
║  ✅ Chroma — новый импорт (langchain_chroma)                      ║
║  ✅ Telegram бот — раздельные сессии bot/user, без конфликтов    ║
║  ✅ Мгновенный ответ «Слушаю, Сэр» при получении сообщения      ║
║  ✅ Умный подбор фраз Джарвиса по категориям и контексту         ║
║  ✅ Обработчик всех сообщений бота (не только incoming=True)     ║
╚══════════════════════════════════════════════════════════════════╝
"""

import asyncio
import json
import os
import re
import sys
import random
from datetime import datetime
from pathlib import Path


# ── UI ────────────────────────────────────────────────────────

# ── Логирование ───────────────────────────────────────────────
from loguru import logger

# ── Конфиг ────────────────────────────────────────────────────
import config
import db as _db
_jarvis_db = _db.get_db()

# ── Настройка loguru ──────────────────────────────────────────
logger.remove()
logger.add(
    str(config.LOG_FILE),
    level="INFO",
    format="{time:YYYY-MM-DD HH:mm:ss} | {level} | {message}",
    rotation="5 MB",
    retention="7 days",
    compression="zip",
    enqueue=True,
)
logger.add(sys.stderr, level="INFO", format="{time:HH:mm:ss} | <level>{level}</level> | {message}")


# ── Глобальное подавление лишних логов ───────────────────────
import logging as _log
import os as _os_env
_os_env.environ.setdefault("HF_HUB_DISABLE_PROGRESS_BARS", "1")
_os_env.environ.setdefault("TOKENIZERS_PARALLELISM", "false")
for _noisy in [
    "httpcore", "httpx.client",
    "huggingface_hub", "sentence_transformers", "transformers",
    "httpx", "httpcore", "urllib3", "asyncio",
]:
    _log.getLogger(_noisy).setLevel(_log.ERROR)

# Подавляем warning от DDGS про impersonate/safari
import logging as _ddgs_log
import warnings as _warnings
_ddgs_log.getLogger("ddgs").setLevel(_ddgs_log.ERROR)
_ddgs_log.getLogger("duckduckgo_search").setLevel(_ddgs_log.ERROR)
_warnings.filterwarnings("ignore", message=".*impersonate.*")
_warnings.filterwarnings("ignore", message=".*safari.*")
_warnings.filterwarnings("ignore", message=".*does not exist.*")

# ── LLM: Groq llama-3.3-70b-versatile ───────────────────────

# ── Telegram ──────────────────────────────────────────────────
from telethon import TelegramClient, events
from telethon.tl.types import DocumentAttributeFilename
from telethon.errors import SessionPasswordNeededError

# ── Файловые парсеры ──────────────────────────────────────────
import PyPDF2
from docx import Document as DocxDocument

# ── HTTP клиент ───────────────────────────────────────────────
import httpx

# ── Поиск ─────────────────────────────────────────────────────
from duckduckgo_search import DDGS  # pip install duckduckgo-search
try:
    from groq import Groq as _GroqClient
    _GROQ_AVAILABLE = True
except ImportError:
    _GroqClient = None
    _GROQ_AVAILABLE = False

# ── Playwright ────────────────────────────────────────────────
try:
    from playwright.async_api import async_playwright
    _PLAYWRIGHT_OK = True
except ImportError:
    async_playwright = None
    _PLAYWRIGHT_OK = False

# ── RAG / Chroma — полностью опциональны (не нужны на bothost) ──
try:
    from langchain_chroma import Chroma
    _CHROMA_AVAILABLE = True
except Exception:
    try:
        from langchain_community.vectorstores import Chroma
        _CHROMA_AVAILABLE = True
    except Exception:
        Chroma = None
        _CHROMA_AVAILABLE = False

try:
    from langchain_huggingface import HuggingFaceEmbeddings
    _EMBEDDINGS_AVAILABLE = True
except Exception:
    try:
        from langchain_community.embeddings import HuggingFaceEmbeddings
        _EMBEDDINGS_AVAILABLE = True
    except Exception:
        HuggingFaceEmbeddings = None
        _EMBEDDINGS_AVAILABLE = False

try:
    from langchain_core.documents import Document
except Exception:
    Document = None


# ═══════════════════════════════════════════════════════════════════
#  УМНЫЙ ПОДБОР ФРАЗб ПО КАТЕГОРИЯМ
# ═══════════════════════════════════════════════════════════════════


# ══════════════════════════════════════════════════════════════
#  JARVIS CORE SYSTEMS v5.0 — надёжность, мониторинг, защита
# ══════════════════════════════════════════════════════════════

import time
import hashlib
import collections
import shutil
import tempfile
import sqlite3
import base64
import traceback
from typing import Optional, Callable
from datetime import datetime, timezone, timedelta
from pathlib import Path


# ──────────────────────────────────────────────────────────────
# RATE LIMITER — защита от флуда и спама
# ──────────────────────────────────────────────────────────────
class RateLimiter:
    """
    Защищает бота от флуда.
    - Ограничивает кол-во запросов от одного пользователя
    - Автоматически снимает блокировку через cooldown
    - Ведёт статистику нарушителей
    """
    def __init__(self, max_requests: int = 5, window_seconds: int = 10, cooldown_seconds: int = 30):
        self.max_requests   = max_requests      # макс запросов за окно
        self.window         = window_seconds    # окно в секундах
        self.cooldown       = cooldown_seconds  # пауза после превышения
        self._buckets: dict[int, list[float]] = {}   # uid → [timestamps]
        self._blocked: dict[int, float]        = {}   # uid → unblock_time
        self._violations: dict[int, int]       = collections.defaultdict(int)

    def is_allowed(self, uid: int) -> tuple[bool, str]:
        """Проверяет можно ли обработать запрос от uid."""
        now = time.time()

        # Снимаем блок если время вышло
        if uid in self._blocked:
            if now < self._blocked[uid]:
                remaining = int(self._blocked[uid] - now)
                return False, f"Сэр, подождите {remaining} сек. перед следующим запросом."
            else:
                del self._blocked[uid]

        # Чистим старые метки
        bucket = self._buckets.get(uid, [])
        bucket = [t for t in bucket if now - t < self.window]
        self._buckets[uid] = bucket

        if len(bucket) >= self.max_requests:
            self._violations[uid] += 1
            # Чем больше нарушений — тем дольше блок
            block_time = self.cooldown * min(self._violations[uid], 5)
            self._blocked[uid] = now + block_time
            self._buckets[uid] = []
            logger.warning(f"⚠️ Флуд от uid={uid}, заблокирован на {block_time}с (нарушение #{self._violations[uid]})")
            return False, f"Слишком много запросов, Сэр. Жду {block_time} секунд."

        self._buckets[uid].append(now)
        return True, ""

    def reset(self, uid: int):
        """Сбросить лимиты для пользователя (например для владельца)."""
        self._buckets.pop(uid, None)
        self._blocked.pop(uid, None)
        self._violations.pop(uid, None)

    def stats(self) -> str:
        now = time.time()
        blocked_now = [(uid, int(t - now)) for uid, t in self._blocked.items() if t > now]
        lines = [f"🛡 Rate Limiter:"]
        lines.append(f"  Активных блоков: {len(blocked_now)}")
        if blocked_now:
            for uid, rem in blocked_now:
                lines.append(f"  uid={uid}: ещё {rem}с")
        total_violations = sum(self._violations.values())
        lines.append(f"  Всего нарушений: {total_violations}")
        return "\n".join(lines)


# ──────────────────────────────────────────────────────────────
# SEARCH CACHE — кэш поисковых запросов
# ──────────────────────────────────────────────────────────────
class SearchCache:
    """
    Кэширует результаты поиска чтобы не делать лишние запросы к API.
    TTL по умолчанию 30 минут для обычных запросов, 5 минут для новостей.
    """
    def __init__(self, default_ttl: int = 1800, max_size: int = 200):
        self.default_ttl = default_ttl
        self.max_size    = max_size
        self._cache: dict[str, dict] = {}   # key → {value, expires, hits}
        self._stats = {"hits": 0, "misses": 0, "evictions": 0}

    def _make_key(self, query: str) -> str:
        return hashlib.md5(query.lower().strip().encode()).hexdigest()

    def get(self, query: str) -> Optional[list]:
        key = self._make_key(query)
        entry = self._cache.get(key)
        if entry and time.time() < entry["expires"]:
            entry["hits"] += 1
            self._stats["hits"] += 1
            return entry["value"]
        if entry:
            del self._cache[key]
        self._stats["misses"] += 1
        return None

    def set(self, query: str, value: list, ttl: int = None):
        # Определяем TTL: для новостей меньше
        if ttl is None:
            q_low = query.lower()
            if any(w in q_low for w in ["новост", "сейчас", "сегодня", "вчера", "курс", "погод"]):
                ttl = 300  # 5 минут для свежих данных
            else:
                ttl = self.default_ttl

        # Вытесняем старые если кэш полон
        if len(self._cache) >= self.max_size:
            oldest = min(self._cache.items(), key=lambda x: x[1]["expires"])
            del self._cache[oldest[0]]
            self._stats["evictions"] += 1

        key = self._make_key(query)
        self._cache[key] = {
            "value":   value,
            "expires": time.time() + ttl,
            "hits":    0,
            "query":   query[:50],
        }

    def clear(self):
        self._cache.clear()

    def stats(self) -> str:
        total = self._stats["hits"] + self._stats["misses"]
        ratio = round(self._stats["hits"] / total * 100) if total else 0
        active = sum(1 for v in self._cache.values() if time.time() < v["expires"])
        return (
            f"🔍 Search Cache:\n"
            f"  Активных записей: {active}/{len(self._cache)}\n"
            f"  Попаданий: {self._stats['hits']} ({ratio}%)\n"
            f"  Промахов: {self._stats['misses']}\n"
            f"  Вытеснений: {self._stats['evictions']}"
        )


# ──────────────────────────────────────────────────────────────
# COMMAND STATS — статистика использования команд
# ──────────────────────────────────────────────────────────────
class CommandStats:
    """
    Считает какие команды и функции используются чаще всего.
    Помогает понять что важно оптимизировать.
    """
    def __init__(self):
        self._counts: dict[str, int]       = collections.defaultdict(int)
        self._times:  dict[str, list[float]] = collections.defaultdict(list)
        self._errors: dict[str, int]       = collections.defaultdict(int)
        self._users:  dict[str, set]       = collections.defaultdict(set)
        self._start   = time.time()

    def record(self, command: str, uid: int = 0, duration_ms: float = 0.0, error: bool = False):
        self._counts[command] += 1
        if uid:
            self._users[command].add(uid)
        if duration_ms > 0:
            times = self._times[command]
            times.append(duration_ms)
            if len(times) > 100:
                self._times[command] = times[-100:]
        if error:
            self._errors[command] += 1

    def top(self, n: int = 10) -> list[tuple[str, int]]:
        return sorted(self._counts.items(), key=lambda x: -x[1])[:n]

    def avg_time(self, command: str) -> float:
        times = self._times.get(command, [])
        return round(sum(times) / len(times), 1) if times else 0.0

    def summary(self) -> str:
        uptime_h = round((time.time() - self._start) / 3600, 1)
        total    = sum(self._counts.values())
        top10    = self.top(10)
        lines    = [f"📊 Статистика команд (аптайм: {uptime_h}ч, всего: {total}):"]
        for cmd, cnt in top10:
            avg = self.avg_time(cmd)
            err = self._errors.get(cmd, 0)
            users = len(self._users.get(cmd, set()))
            err_str = f" ❌{err}" if err else ""
            time_str = f" ~{avg}мс" if avg else ""
            lines.append(f"  {cmd}: {cnt}x · {users} польз.{time_str}{err_str}")
        return "\n".join(lines)


# ──────────────────────────────────────────────────────────────
# HEALTH MONITOR — мониторинг всех систем в фоне
# ──────────────────────────────────────────────────────────────
class HealthMonitor:
    """
    Фоновый мониторинг здоровья всех систем.
    Проверяет БД, LLM, интернет каждые N минут.
    При деградации — уведомляет владельца.
    """
    def __init__(self):
        self._status: dict[str, dict] = {}   # system → {ok, last_check, error}
        self._alerts_sent: set = set()        # какие алерты уже отправлены
        self._check_interval = 300            # 5 минут между проверками

    def update(self, system: str, ok: bool, detail: str = ""):
        was_ok = self._status.get(system, {}).get("ok", True)
        self._status[system] = {
            "ok":         ok,
            "detail":     detail,
            "last_check": time.time(),
        }
        # Если система восстановилась — снимаем алерт
        if ok and system in self._alerts_sent:
            self._alerts_sent.discard(system)
            logger.info(f"✅ [{system}] восстановлен")
        # Если система упала первый раз — логируем
        if not ok and was_ok:
            logger.error(f"❌ [{system}] деградация: {detail}")

    def needs_alert(self, system: str) -> bool:
        """Нужно ли отправить алерт владельцу (только один раз пока не починят)."""
        if system in self._alerts_sent:
            return False
        if not self._status.get(system, {}).get("ok", True):
            self._alerts_sent.add(system)
            return True
        return False

    def get_status(self) -> str:
        if not self._status:
            return "ℹ️ Мониторинг ещё не запускался"
        lines = ["🏥 Здоровье систем:"]
        for sys_name, info in sorted(self._status.items()):
            icon = "✅" if info["ok"] else "❌"
            ago  = int(time.time() - info["last_check"])
            ago_str = f"{ago}с назад" if ago < 60 else f"{ago//60}м назад"
            detail = f" — {info['detail']}" if info.get("detail") and not info["ok"] else ""
            lines.append(f"  {icon} {sys_name}{detail} ({ago_str})")
        return "\n".join(lines)

    async def run_checks(self, agent) -> dict[str, bool]:
        """Запускает все проверки и возвращает результаты."""
        results = {}

        # 1. База данных
        try:
            cur = _jarvis_db._cur()
            cur.execute("SELECT 1")
            cur.close()
            self.update("PostgreSQL", True)
            results["db"] = True
        except Exception as e:
            self.update("PostgreSQL", False, str(e)[:80])
            results["db"] = False

        # 2. LLM (Groq)
        try:
            import os as _os_hg
            if _GROQ_AVAILABLE and _os_hg.getenv("GROQ_API_KEY", ""):
                _gc = _GroqClient(api_key=_os_hg.getenv("GROQ_API_KEY"))
                _gc.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": "1"}],
                    max_completion_tokens=1,
                )
            self.update("LLM (Groq)", True)
            results["llm"] = True
        except Exception as e:
            self.update("LLM (Groq)", False, str(e)[:80])
            results["llm"] = False

        # Google Drive убран

        # 4. Интернет (быстрый пинг)
        try:
            import httpx as _hx2
            async with httpx.AsyncClient(timeout=5, verify=False) as cl2:
                r2 = await cl2.get("https://1.1.1.1", follow_redirects=True)
            self.update("Интернет", True)
            results["internet"] = True
        except Exception as e:
            self.update("Интернет", False, str(e)[:60])
            results["internet"] = False

        return results


# ──────────────────────────────────────────────────────────────
# MESSAGE DEDUPLICATOR — защита от двойной обработки
# ──────────────────────────────────────────────────────────────
class MessageDeduplicator:
    """
    Предотвращает обработку одного сообщения дважды.
    Актуально при reconnect или дублированных событиях от Telegram.
    """
    def __init__(self, ttl_seconds: int = 60, max_size: int = 500):
        self.ttl      = ttl_seconds
        self.max_size = max_size
        self._seen: dict[str, float] = {}  # msg_key → timestamp

    def is_duplicate(self, chat_id: int, msg_id: int) -> bool:
        now = time.time()
        key = f"{chat_id}:{msg_id}"

        # Чистим старые
        if len(self._seen) > self.max_size:
            cutoff = now - self.ttl
            self._seen = {k: v for k, v in self._seen.items() if v > cutoff}

        if key in self._seen and now - self._seen[key] < self.ttl:
            return True

        self._seen[key] = now
        return False


# ──────────────────────────────────────────────────────────────
# CONNECTION WATCHDOG — авто-переподключение при сбоях
# ──────────────────────────────────────────────────────────────
class ConnectionWatchdog:
    """
    Следит за соединением с Telegram и PostgreSQL.
    При разрыве — пытается переподключиться с экспоненциальной задержкой.
    """
    def __init__(self):
        self._db_failures   = 0
        self._last_db_ok    = time.time()
        self._reconnect_at  = 0.0
        self.max_backoff    = 300   # максимум 5 минут между попытками

    def db_ok(self):
        self._db_failures = 0
        self._last_db_ok  = time.time()

    def db_failed(self) -> float:
        """Вызывается при ошибке БД. Возвращает секунды до следующей попытки."""
        self._db_failures += 1
        backoff = min(2 ** self._db_failures, self.max_backoff)
        self._reconnect_at = time.time() + backoff
        logger.warning(f"⚠️ БД: сбой #{self._db_failures}, следующая попытка через {backoff}с")
        return backoff

    def should_retry_db(self) -> bool:
        return time.time() >= self._reconnect_at

    async def try_reconnect_db(self) -> bool:
        """Проверяет и при необходимости переподключает SQLite."""
        try:
            if _jarvis_db.ping():
                self.db_ok()
                return True
            # Пробуем переподключить
            _jarvis_db._conn = None
            if _jarvis_db.ping():
                self.db_ok()
                logger.info("✅ Watchdog: SQLite переподключена")
                return True
            self.db_failed()
            return False
        except Exception as e:
            self.db_failed()
            logger.error(f"❌ Watchdog: SQLite недоступна: {e}")
            return False

    def status(self) -> str:
        db_ago = int(time.time() - self._last_db_ok)
        return (
            f"🐕 Watchdog:\n"
            f"  БД: последний успех {db_ago}с назад, сбоев: {self._db_failures}"
        )


# ──────────────────────────────────────────────────────────────
# TYPING MANAGER — умное управление индикатором печати
# ──────────────────────────────────────────────────────────────
class TypingManager:
    """
    Показывает 'печатает...' пока бот думает.
    Автоматически продлевает если ответ долгий.
    """
    def __init__(self, client, chat_id: int):
        self.client  = client
        self.chat_id = chat_id
        self._task: Optional[asyncio.Task] = None
        self._active = False

    async def _loop(self):
        try:
            from telethon.tl.functions.messages import SetTypingRequest
            from telethon.tl.types import SendMessageTypingAction
            while self._active:
                try:
                    await self.client(SetTypingRequest(self.chat_id, SendMessageTypingAction()))
                except Exception:
                    pass
                await asyncio.sleep(4)
        except asyncio.CancelledError:
            pass

    async def start(self):
        self._active = True
        self._task   = asyncio.create_task(self._loop())

    async def stop(self):
        self._active = False
        if self._task and not self._task.done():
            self._task.cancel()
            try:
                await self._task
            except asyncio.CancelledError:
                pass
        try:
            from telethon.tl.functions.messages import SetTypingRequest
            from telethon.tl.types import SendMessageCancelAction
            await self.client(SetTypingRequest(self.chat_id, SendMessageCancelAction()))
        except Exception:
            pass


# ──────────────────────────────────────────────────────────────
# RETRY DECORATOR — авто-повтор для нестабильных операций
# ──────────────────────────────────────────────────────────────
async def retry_async(
    func: Callable,
    *args,
    attempts: int = 3,
    delay: float = 1.0,
    backoff: float = 2.0,
    exceptions=(Exception,),
    label: str = "",
    **kwargs,
):
    """
    Повторяет async-функцию при сбое.
    Пример: await retry_async(upload_file, path, attempts=3, label="загрузка")
    """
    last_exc = None
    wait = delay
    for attempt in range(1, attempts + 1):
        try:
            return await func(*args, **kwargs)
        except exceptions as e:
            last_exc = e
            if attempt < attempts:
                logger.warning(f"⚠️ {label or func.__name__}: попытка {attempt}/{attempts} не удалась — {e}. Жду {wait:.1f}с")
                await asyncio.sleep(wait)
                wait *= backoff
            else:
                logger.error(f"❌ {label or func.__name__}: все {attempts} попытки провалились. Последняя ошибка: {e}")
    raise last_exc


# ──────────────────────────────────────────────────────────────
# ERROR CLASSIFIER — классификатор ошибок с русскими советами
# ──────────────────────────────────────────────────────────────
class ErrorClassifier:
    """
    Классифицирует технические ошибки и даёт человеческие советы.
    Работает без LLM — быстро и надёжно.
    """
    RULES = [
        # (ключевые слова в ошибке, тип, совет)
        (["connection timeout", "connect timeout", "ConnectionTimeout"],
         "Таймаут подключения",
         "Проверьте интернет-соединение. Если используете VPN — убедитесь что он включён."),

        (["authentication failed", "password authentication", "invalid password"],
         "Ошибка аутентификации БД",
         "Проблема с доступом к БД. Проверьте файл database/Jarvis.db и права доступа."),

        (["ssl", "certificate", "SSL"],
         "Ошибка SSL-сертификата",
         "Проблема с SSL-сертификатом. Обычно связано с HTTPS запросами."),

        (["too many connections", "remaining connection slots", "database is locked"],
         "Конфликт доступа к SQLite",
         "Несколько процессов пытаются писать в БД. Убедитесь что запущен только один экземпляр бота."),

        (["relation does not exist", "table", "does not exist"],
         "Таблица не найдена в БД",
         "Таблицы не созданы. Перезапустите бота — они создадутся автоматически."),

        (["FloodWaitError", "flood", "Flood"],
         "Флуд-лимит Telegram",
         "Telegram ограничил отправку сообщений. Бот сам подождёт и продолжит."),

        (["AuthKeyError", "AuthKeyDuplicated", "auth key"],
         "Ошибка авторизации Telegram",
         "Сессия повреждена. Удалите файл user.session и создайте новую через create_session.py"),

        (["UserDeactivatedError", "deactivated"],
         "Аккаунт деактивирован",
         "Telegram аккаунт заблокирован или удалён. Проверьте аккаунт вручную."),

        (["ChatWriteForbiddenError", "forbidden", "Forbidden"],
         "Нет прав на отправку",
         "Бот не имеет прав писать в этот чат. Добавьте его как администратора."),

        (["rate limit", "RateLimitError", "429"],
         "Лимит запросов к API",
         "Превышен лимит запросов к LLM. Подождите минуту или смените модель."),

        (["ConnectionError", "Network", "network"],
         "Сетевая ошибка",
         "Нет доступа к интернету или сервис недоступен. Проверьте соединение."),

        (["JSONDecodeError", "json", "JSON"],
         "Неверный формат ответа",
         "API вернул не JSON. Возможно сервис временно недоступен — попробуйте позже."),

        (["PermissionError", "Permission", "Access denied"],
         "Ошибка доступа к файлу",
         "Нет прав на чтение/запись файла. Проверьте права на папку бота."),

        (["FileNotFoundError", "No such file"],
         "Файл не найден",
         "Нужный файл отсутствует. Убедитесь что все файлы бота на месте."),

        (["MemoryError", "out of memory"],
         "Нехватка памяти",
         "Сервер заканчивается оперативная память. Перезапустите бота."),
    ]

    @classmethod
    def classify(cls, error: Exception) -> tuple[str, str]:
        """Возвращает (тип_ошибки, совет)."""
        err_str = str(error).lower()
        err_type = type(error).__name__

        for keywords, name, advice in cls.RULES:
            if any(k.lower() in err_str for k in keywords):
                return name, advice

        return err_type, "Неизвестная ошибка. Проверьте логи для подробностей."

    @classmethod
    def format(cls, error: Exception, context: str = "") -> str:
        """Форматирует ошибку для пользователя."""
        name, advice = cls.classify(error)
        ctx_str = f" при {context}" if context else ""
        return (
            f"🔴 Ошибка{ctx_str}: **{name}**\n"
            f"💡 Что делать: {advice}"
        )


# ──────────────────────────────────────────────────────────────
# GRACEFUL SHUTDOWN — чистое завершение работы
# ──────────────────────────────────────────────────────────────
class GracefulShutdown:
    """
    Обрабатывает SIGTERM/SIGINT.
    Сохраняет состояние перед выходом.
    """
    def __init__(self):
        self._shutdown_event = asyncio.Event()
        self._handlers: list[Callable] = []

    def register(self, handler: Callable):
        """Регистрирует функцию которая вызовется при выходе."""
        self._handlers.append(handler)

    async def wait(self):
        await self._shutdown_event.wait()

    def trigger(self):
        self._shutdown_event.set()

    async def run_handlers(self):
        for handler in self._handlers:
            try:
                if asyncio.iscoroutinefunction(handler):
                    await handler()
                else:
                    handler()
            except Exception as e:
                logger.error(f"Ошибка при завершении: {e}")


# ──────────────────────────────────────────────────────────────
# GLOBAL INSTANCES — синглтоны для всего приложения
# ──────────────────────────────────────────────────────────────
_rate_limiter   = RateLimiter(max_requests=8, window_seconds=10, cooldown_seconds=20)
_search_cache   = SearchCache(default_ttl=1800, max_size=300)
_cmd_stats      = CommandStats()
_health_monitor = HealthMonitor()
_deduplicator   = MessageDeduplicator(ttl_seconds=60)
_watchdog       = ConnectionWatchdog()
_shutdown       = GracefulShutdown()
_error_clf      = ErrorClassifier()
_ack_msg_ids: set[int] = set()
_last_code_file: dict = {}  # sender_id → {filename, content, language, file_id}
# Классы WeatherCurrencyService, MediaHandler, SmartFormatter
# определены ниже и используются напрямую

class PhraseBank:
    """
    Загружает фразы из файлов phrases/iron_man/*.txt
    Формат строки: КАТЕГОРИЯ: текст фразы
    Подбирает подходящую фразу по контексту запроса и ответа.
    """

    # Ключевые слова для определения категории по контексту
    CATEGORY_KEYWORDS = {
        "ПРИВЕТСТВИЕ": ["привет", "здравствуй", "добрый", "хай", "hello", "hi", "начнём"],
        "ГОТОВНОСТЬ":  ["сделай", "выполни", "запусти", "помоги", "можешь", "пожалуйста"],
        "ИРОНИЯ":      ["почему", "зачем", "опять", "снова", "нет", "ошибка", "плохо"],
        "УСПЕХ":       ["готово", "выполнено", "успешно", "отлично", "спасибо", "молодец"],
        "ОШИБКА":      ["ошибка", "не работает", "сломалось", "проблема", "не могу", "упало"],
        "ОЖИДАНИЕ":    ["подожди", "жду", "скоро", "обрабатываю", "ищу", "считаю"],
        "УМНЫЙ":       ["почему", "как", "зачем", "смысл", "теория", "объясни", "расскажи"],
        "ОПАСНОСТЬ":   ["опасно", "критично", "срочно", "важно", "внимание", "риск"],
        "ОДОБРЕНИЕ":   ["согласен", "да", "правильно", "верно", "именно", "точно"],
        "ОТКАЗ":       ["нельзя", "запрещено", "не буду", "отказ", "нет возможности"],
        "АНАЛИЗ":      ["анализ", "сравни", "разбери", "оцени", "проверь", "исследуй"],
        "ПРОЩАНИЕ":    ["пока", "до свидания", "выключись", "стоп", "выход", "bye"],
    }

    def __init__(self):
        self.phrases: dict[str, list[str]] = {}  # категория → список фраз
        self._load()

    def _load(self):
        total = 0

        def _parse_lines(lines: list[str]):
            nonlocal total
            for line in lines:
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                if ":" in line:
                    category, _, text = line.partition(":")
                    cat = category.strip().upper()
                    txt = text.strip()
                    if cat and txt:
                        self.phrases.setdefault(cat, []).append(txt)
                        total += 1
                else:
                    self.phrases.setdefault("GENERAL", []).append(line)
                    total += 1

        loaded = False

        # 1. Папка phrases/iron_man/ — каждый .txt файл отдельно
        phrases_dir = getattr(config, "PHRASES_DIR", None)
        if phrases_dir and Path(phrases_dir).is_dir():
            for txt_file in sorted(Path(phrases_dir).glob("*.txt")):
                try:
                    _parse_lines(txt_file.read_text("utf-8", errors="ignore").splitlines())
                    loaded = True
                except Exception as e:
                    logger.warning(f"Phrases load {txt_file.name}: {e}")

        # 2. Одиночный файл phrases.txt (fallback)
        if not loaded:
            phrases_file = config.PHRASES_FILE
            if phrases_file.exists():
                try:
                    _parse_lines(phrases_file.read_text("utf-8", errors="ignore").splitlines())
                    loaded = True
                except Exception as e:
                    logger.warning(f"Phrases load: {e}")

        if not loaded:
            logger.warning(f"Phrases file not found: {config.PHRASES_FILE}")

    def get(self, context: str = "", category: str = "", chance: float = 0.20) -> str:
        """
        Вернуть фразу по контексту или категории.
        chance — вероятность добавить фразу (0.0–1.0)
        """
        if random.random() > chance:
            return ""

        if not self.phrases:
            return ""

        # Если категория задана явно — берём из неё
        if category and category.upper() in self.phrases:
            return "\n\n" + random.choice(self.phrases[category.upper()])

        # Иначе подбираем по ключевым словам контекста
        ctx_lower = context.lower()
        best_cat  = ""
        best_score = 0

        for cat, keywords in self.CATEGORY_KEYWORDS.items():
            score = sum(1 for kw in keywords if kw in ctx_lower)
            if score > best_score and cat in self.phrases:
                best_score = score
                best_cat   = cat

        # Берём из подходящей категории или рандомной
        pool: list[str] = []
        if best_cat and best_score > 0:
            pool = self.phrases[best_cat]
        else:
            # Случайная категория из доступных
            all_phrases = [p for phrases in self.phrases.values() for p in phrases]
            pool = all_phrases

        if not pool:
            return ""

        return "\n\n" + random.choice(pool)

    @property
    def total(self) -> int:
        return sum(len(v) for v in self.phrases.values())


# ═══════════════════════════════════════════════════════════════════
#  ИСТОРИЯ ПЕРЕПИСОК
# ═══════════════════════════════════════════════════════════════════

class ChatHistory:
    """История чатов — SQLite backend."""

    HISTORY_TRIGGERS = [
        "что я спрашивал", "о чём мы говорили", "найди в переписке",
        "вспомни", "что было раньше", "история чата", "статистика чата",
        "что мы обсуждали", "прошлые разговоры", "найди в истории",
        "сколько сообщений", "наша переписка",
    ]

    def __init__(self):
        self._db = _jarvis_db

    def save_message(self, sid: int, role: str, text: str, username: str = "",
                     msg_id: int = 0, chat_id: int = 0):
        self._db.save_message(sid, role, text, username, msg_id=msg_id, chat_id=chat_id)

    def get_recent(self, sid: int, n: int = 30) -> list[dict]:
        return self._db.get_recent(sid, n)

    def search(self, sid: int, query: str, limit: int = 20) -> list[dict]:
        return self._db.search_messages(sid, query, limit)

    def format_for_llm(self, msgs: list[dict]) -> str:
        lines = []
        for m in msgs:
            role = "Пользователь" if m["role"] == "user" else "Джарвис"
            lines.append(f"[{m.get('ts','?')}] {role}: {m['text']}")
        return "\n".join(lines)

    def stats(self, sid: int) -> dict:
        return self._db.message_stats(sid)

    @classmethod
    def is_history_request(cls, q: str) -> bool:
        return any(t in q.lower() for t in cls.HISTORY_TRIGGERS)

    async def answer_history_question(self, query: str, sid: int, llm) -> str:
        q_lower = query.lower()
        if any(p in q_lower for p in ["статистика", "сколько сообщений"]):
            s = self.stats(sid)
            if s["total"] == 0:
                return "Сэр, история переписки пуста — мы только начинаем."
            return (
                f"Сэр, статистика нашей переписки:\n"
                f"• Всего: {s['total']} сообщений\n"
                f"• Ваших: {s['user_msgs']}\n"
                f"• Моих: {s['bot_msgs']}\n"
                f"• Первое: {s['first_date']}\n"
                f"• Последнее: {s['last_date']}"
            )
        kw_match = re.search(r"найди[:\s]+(.+)|найди в (истории|переписке)[:\s]*(.+)", q_lower)
        if kw_match:
            keyword = (kw_match.group(1) or kw_match.group(3) or "").strip()
            if keyword:
                found = self.search(sid, keyword)
                if not found:
                    return f"Сэр, по запросу «{keyword}» ничего не нашёл."
                return f"Сэр, нашёл в истории:\n\n```\n{self.format_for_llm(found[-10:])}\n```"
        recent = self.get_recent(sid, 30)
        if not recent:
            return "Сэр, история переписки пуста."
        ctx = self.format_for_llm(recent)
        msgs = [
            {"role": "system", "content": "Ты Джарвис. Отвечай на вопросы об истории переписки."},
            {"role": "user",   "content": f"История:\n\n{ctx}\n\nВопрос: {query}"},
        ]
        return await llm.complete(msgs, max_tokens=800)


# ═══════════════════════════════════════════════════════════════════
#  UNIVERSAL LLM CONNECTOR
# ═══════════════════════════════════════════════════════════════════

class UniversalLLMConnector:
    """
    Основная модель: Groq openai/gpt-oss-120b (reasoning_effort=medium)
    Резерв:          Groq llama-3.3-70b-versatile
    """

    MODEL_MAIN  = "openai/gpt-oss-120b"
    MODEL_SMART = "llama-3.3-70b-versatile"

    _SMART_TRIGGERS = [
        "реши", "вычисли", "докажи", "объясни подробно", "напиши код",
        "проверь код", "найди баги", "оптимизируй", "рефактор",
        "сравни", "проанализируй", "анализ", "разбери", "кто прав",
        "рассуди", "почему", "как работает", "объясни принцип",
        "напиши статью", "напиши эссе", "переведи текст",
        "составь план", "составь список", "придумай",
        "математик", "уравнение", "формул", "интеграл", "производн",
        "алгоритм", "структур данных", "архитектур",
        "диагностик", "логи", "traceback", "ошибка в коде",
    ]

    def __init__(self):
        import os as _os_g
        self._groq_key = _os_g.getenv("GROQ_API_KEY",
                         config.GROQ_API_KEY if hasattr(config, "GROQ_API_KEY") else "")
        if not self._groq_key:
            logger.warning("⚠️ GROQ_API_KEY не задан в .env")
        else:
            logger.info(f"✅ Groq активирован: {self.MODEL_MAIN}")

    def _pick_model(self, messages: list[dict]) -> str:
        """Всегда основная модель. Резерв только для явно сложных задач."""
        last_user = ""
        for m in reversed(messages):
            if m.get("role") == "user":
                content = m.get("content", "")
                last_user = (" ".join(p.get("text", "") for p in content
                             if isinstance(p, dict)) if isinstance(content, list) else str(content))
                break
        q = last_user.lower()
        # Переключаемся на smart только для явных технических задач
        _hard_triggers = [
            "напиши код", "проверь код", "найди баги", "оптимизируй",
            "рефактор", "напиши статью", "напиши эссе",
        ]
        if any(t in q for t in _hard_triggers):
            return self.MODEL_SMART
        return self.MODEL_MAIN

    @property
    def current_display(self) -> str:
        return f"Groq {self.MODEL_MAIN}"

    @property
    def current_spec(self) -> dict:
        return {"type": "groq", "model": self.MODEL_MAIN}

    def list_models(self) -> str:
        status = "✅ ключ задан" if self._groq_key else "❌ не задан"
        return (
            "**🤖 Модели Джарвиса:**\n\n"
            f"  1️⃣ Groq {self.MODEL_MAIN}: {status} (основной)\n"
            f"  2️⃣ Groq {self.MODEL_SMART}: резерв (сложные задачи)\n\n"
            "Переключение — автоматическое по сложности запроса."
        )

    def switch(self, query: str) -> str | None:
        q = query.lower()
        if any(p in q for p in ["текущая модель", "какая модель", "что используешь",
                                  "смени модель", "список моделей"]):
            return "Сэр, вот текущие модели:\n\n" + self.list_models()
        return None

    def add_custom(self, *a, **kw) -> str:
        return f"Сэр, используется Groq {self.MODEL_MAIN} с автовыбором."

    async def complete(self, messages: list[dict], max_tokens: int = 1400) -> str:
        """Основной вызов через Groq."""
        import concurrent.futures as _cf, asyncio as _aio

        if not _GROQ_AVAILABLE:
            return "Сэр, установите groq: pip install groq"
        if not self._groq_key:
            return "Сэр, задайте GROQ_API_KEY в .env"

        chosen_model = self._pick_model(messages)
        logger.info(f"🤖 Groq запрос: {chosen_model}")

        def _call(model: str):
            client = _GroqClient(api_key=self._groq_key)
            # Нормализуем messages — убираем list-content
            clean = []
            for m in messages:
                if isinstance(m.get("content"), list):
                    txt = " ".join(p.get("text", "") for p in m["content"]
                                   if isinstance(p, dict) and p.get("type") == "text")
                    clean.append({"role": m["role"], "content": txt})
                else:
                    clean.append(m)
            kwargs = dict(
                model=model,
                messages=clean,
                temperature=1,
                max_completion_tokens=min(max_tokens, 8192),
                top_p=1,
                stream=False,
                stop=None,
            )
            # reasoning_effort поддерживается только gpt-oss моделью
            if model == self.MODEL_MAIN:
                kwargs["reasoning_effort"] = "medium"
            return client.chat.completions.create(**kwargs).choices[0].message.content or ""

        try:
            loop = _aio.get_event_loop()
            with _cf.ThreadPoolExecutor(max_workers=2) as ex:
                return await loop.run_in_executor(ex, _call, chosen_model)
        except Exception as e:
            err = str(e)
            logger.error(f"❌ Groq ({chosen_model}): {e}")
            # Фоллбэк на вторую модель если первая недоступна
            if chosen_model == self.MODEL_MAIN and (
                "503" in err or "unavailable" in err.lower() or "model" in err.lower()
            ):
                logger.warning(f"⚠️ {self.MODEL_MAIN} недоступна — переключаюсь на {self.MODEL_SMART}")
                try:
                    loop2 = _aio.get_event_loop()
                    with _cf.ThreadPoolExecutor(max_workers=2) as ex:
                        return await loop2.run_in_executor(ex, _call, self.MODEL_SMART)
                except Exception as e2:
                    err = str(e2)
                    logger.error(f"❌ Groq fallback ({self.MODEL_SMART}): {e2}")
            if "401" in err or "403" in err or "invalid_api_key" in err:
                return "Сэр, GROQ_API_KEY неверный. Проверьте .env"
            if "429" in err or "rate_limit" in err:
                return "Сэр, лимит Groq. Подождите секунду."
            if "503" in err or "unavailable" in err.lower():
                return "Сэр, Groq временно недоступен."
            return f"Сэр, ошибка Groq: {type(e).__name__}"


# ═══════════════════════════════════════════════════════════════════
#  АНАЛИЗАТОР СПОРОВ
# ═══════════════════════════════════════════════════════════════════

class DisputeAnalyzer:

    TRIGGERS = [
        "разбери переписку", "кто прав", "рассуди нас", "разбери спор",
        "рассуди спор", "разбор переписки", "кто виноват", "кто не прав",
        "рассуди конфликт", "analyze dispute",
    ]

    SYSTEM = """Ты — JARVIS. Анализируешь конфликты как судья: холодно, по фактам, без сочувствия.

СТРОГИЙ ФОРМАТ ОТВЕТА:

👤 [Имя/ник 1] — [одна фраза: позиция + главный аргумент]
👤 [Имя/ник 2] — [одна фраза: позиция + главный аргумент]

📋 [3-4 предложения по тексту переписки. Цитируй конкретные фразы. Указывай слабые места: манипуляции, логические ошибки, уход от темы, противоречия. Только то что написано — ничего не додумывай.]

━━━━━━━━━━━━━━━━━━━━
✅ ПРАВ: [ИМЯ/НИК]
[одно предложение — конкретная причина]
━━━━━━━━━━━━━━━━━━━━

Если оба неправы → ✅ ПРАВ: ОБА НЕПРАВЫ + одно предложение.
Если данных недостаточно → скажи об этом прямо.
Без рекомендаций. Без советов. Только вердикт."""

    @classmethod
    def is_triggered(cls, q: str) -> bool:
        return any(t in q.lower() for t in cls.TRIGGERS)

    @classmethod
    def strip_trigger(cls, q: str) -> str:
        text = q.strip()
        for t in cls.TRIGGERS:
            text = re.sub(re.escape(t) + r"[\s:]*", "", text, flags=re.IGNORECASE, count=1).strip()
        return text

    @staticmethod
    def _has_conflict(text: str) -> bool:
        """Проверяет что в тексте есть признаки конфликта/спора."""
        conflict_words = [
            "виноват", "виновата", "не прав", "не права", "ошибся", "ошиблась",
            "обвиня", "претензи", "твоя вина", "моя вина", "не сделал", "не выполнил",
            "неправ", "врёт", "врешь", "лжёт", "обманул", "нарушил", "срыв",
            "ты должен", "ты должна", "всё равно", "но ты", "а ты", "зато ты",
            "не так", "неверно", "неправильно", "плохо сделал", "не согласен", "не согласна",
            "это неправда", "это ложь",
        ]
        tl = text.lower()
        return any(w in tl for w in conflict_words)

    async def analyze(self, conversation: str, llm: UniversalLLMConnector) -> str:
        if not conversation.strip():
            return self._help()
        if not self._has_conflict(conversation):
            return "Сэр, в этой переписке не обнаружено конфликта. Здесь пока никто ни с кем не спорит."
        return await llm.complete([
            {"role": "system", "content": self.SYSTEM},
            {"role": "user",   "content": f"Разбери переписку:\n\n{conversation}"},
        ], max_tokens=2500)

    async def analyze_forwarded(self, msgs: list[dict], llm: UniversalLLMConnector) -> str:
        if not msgs:
            return "Сэр, вы ещё не пересылали сообщений. Перешлите их сначала."
        lines = [f"[{m.get('date','')}] {m.get('sender','?')}: {m.get('text','')}" for m in msgs]
        return await self.analyze("\n".join(lines), llm)

    @staticmethod
    def _help() -> str:
        return (
            "Сэр, вставьте текст переписки после команды.\n\n"
            "**Способ 1:**\n```\nДжарвис, разбери переписку:\n"
            "Иван: ты виноват!\nМаша: нет, ты!\n```\n\n"
            "**Способ 2:** перешлите сообщения из спора, затем напишите «Джарвис, разбери переписку»."
        )


# ═══════════════════════════════════════════════════════════════════
#  GROUP LOGGER — сохраняет ВСЕ сообщения из групп, включая удалённые
# ═══════════════════════════════════════════════════════════════════

class GroupLogger:
    """Логи группы и удалённые сообщения — SQLite backend."""

    DELETED_TRIGGERS = [
        "покажи удалённые", "покажи удаленные", "что удалили",
        "удалённые сообщения", "удаленные сообщения",
        "что стёрли", "show deleted", "удалённые за", "удаленные за",
    ]

    def __init__(self):
        self._db = _jarvis_db

    def save(self, chat_id: int, msg_id: int, sender: str,
             sender_id: int, text: str, date: str):
        self._db.save_group_msg(chat_id, msg_id, sender, sender_id, text, date)

    def mark_deleted(self, chat_id: int, msg_ids: list[int]):
        self._db.mark_deleted(chat_id, msg_ids)

    def mark_deleted_all_chats(self, msg_ids: list[int]):
        self._db.mark_deleted_all_chats(msg_ids)

    def get_deleted(self, chat_id: int, limit: int = 20, date_filter: str = "") -> list[dict]:
        return self._db.get_deleted(chat_id, limit, date_filter)

    def format_deleted(self, msgs: list[dict]) -> str:
        if not msgs:
            return "Сэр, удалённых сообщений не найдено."
        lines = [f"🗑 Удалённые сообщения ({len(msgs)} шт.):"]
        for m in msgs:
            lines.append(f"  [{m.get('date','')}] {m.get('sender','?')}: {m.get('text','')}")
        return "\n".join(lines)

    @classmethod
    def is_deleted_request(cls, q: str) -> bool:
        return any(t in q.lower() for t in cls.DELETED_TRIGGERS)


class UserProfileManager:
    """Профили пользователей — SQLite backend."""

    SAVE_TRIGGERS   = [
        "запомни —", "запомни -", "запомни: ", "запомни, что", "запомни что",
        "запомни ", "запомни,",
    ]
    VIEW_TRIGGERS   = ["что ты знаешь обо мне", "мой профиль", "что помнишь обо мне", "мои данные"]
    CLEAR_TRIGGERS  = ["забудь всё что знаешь обо мне", "забудь всё обо мне",
                       "удали мой профиль", "очисти мой профиль"]
    STYLE_SHORT     = ["отвечай мне короче", "отвечай короче", "давай покороче",
                       "отвечай кратко", "пиши кратко", "коротко"]
    STYLE_LONG      = ["отвечай подробнее", "отвечай развёрнуто", "давай подробнее",
                       "пиши подробно", "развёрнуто", "подробнее"]
    STYLE_IRONIC    = ["отвечай с иронией", "можно с иронией", "добавь иронию", "с юмором"]
    STYLE_NEUTRAL   = ["отвечай нейтрально", "убери иронию", "стандартный стиль"]

    def __init__(self):
        self._db = _jarvis_db

    def load(self, uid: int) -> dict:
        return self._db.load_profile(uid)

    def save_profile(self, uid: int, profile: dict):
        self._db.save_profile(uid, profile)

    def add_fact(self, uid: int, fact: str) -> str:
        p = self.load(uid)
        p.setdefault("facts", [])
        structured = self._parse_fact(fact)
        # Заменяем если уже есть факт с такой же меткой
        label = structured.split(":")[0] if ":" in structured else None
        if label:
            p["facts"] = [f for f in p["facts"] if not f.startswith(label + ":")]
        if structured not in p["facts"]:
            p["facts"].append(structured)
        self.save_profile(uid, p)
        return f"Принято к сведению, Сэр. Записал: {structured}"

    def _parse_fact(self, text: str) -> str:
        """Превращает 'меня зовут Максим' → 'Имя: Максим'"""
    
        t = text.strip().rstrip(".")

        patterns = [
            (r"меня зовут (.+)",            "Имя"),
            (r"моё имя (.+)",               "Имя"),
            (r"мое имя (.+)",               "Имя"),
            (r"я (.+) лет",                 "Возраст"),
            (r"мне (.+) лет",               "Возраст"),
            (r"мне (.+) год",               "Возраст"),
            (r"я живу в (.+)",              "Город"),
            (r"я из города (.+)",           "Город"),
            (r"я из (.+)",                  "Город"),
            (r"мой город (.+)",             "Город"),
            (r"я работаю (.+)",             "Работа"),
            (r"моя профессия (.+)",         "Профессия"),
            (r"я (?:по профессии\s)?(.+)",  "Профессия"),
            (r"мне нравится (.+)",          "Интерес"),
            (r"я люблю (.+)",               "Интерес"),
            (r"моё хобби (.+)",             "Хобби"),
            (r"мое хобби (.+)",             "Хобби"),
            (r"мой номер (.+)",             "Телефон"),
            (r"мой email (.+)",             "Email"),
        ]

        tl = t.lower()
        for pattern, label in patterns:
            m = re.search(pattern, tl)
            if m:
                value = t[m.start(1):m.end(1)].strip()
                value = value[0].upper() + value[1:] if value else value
                return f"{label}: {value}"

        # Не распознано — сохраняем как есть, но с заглавной буквы
        return t[0].upper() + t[1:] if t else t

    def get_summary(self, uid: int) -> str:
        p = self.load(uid)
        facts = p.get("facts", [])
        style = p.get("style", "normal")
        if not facts:
            return "Досье пустое, Сэр. Расскажите о себе — запомню."
        style_map = {"short": "краткий", "long": "подробный",
                     "ironic": "ироничный", "normal": "стандартный"}
        lines = ["Сэр, вот что я знаю о вас:"]
        for f in facts:
            lines.append(f"  • {f}")
        lines.append(f"Стиль общения: {style_map.get(style, 'стандартный')}")
        return "\n".join(lines)

    def clear(self, uid: int) -> str:
        self._db.delete_profile(uid)
        return "Досье очищено, Сэр. Начинаем с чистого листа."

    def set_style(self, uid: int, style: str) -> str:
        p = self.load(uid)
        p["style"] = style
        self.save_profile(uid, p)
        labels = {"short": "краткий", "long": "подробный",
                  "ironic": "ироничный", "normal": "стандартный"}
        return f"Принято, Сэр. Стиль общения: {labels.get(style, style)}."

    def get_style(self, uid: int) -> str:
        return self.load(uid).get("style", "normal")

    def get_facts_str(self, uid: int) -> str:
        facts = self.load(uid).get("facts", [])
        if not facts:
            return
        return "Факты о пользователе: " + "; ".join(facts)

    @classmethod
    def is_save(cls, q: str) -> bool:
        ql = q.lower()
        return any(t in ql for t in cls.SAVE_TRIGGERS)

    @classmethod
    def is_view(cls, q: str) -> bool:
        ql = q.lower()
        return any(t in ql for t in cls.VIEW_TRIGGERS)

    @classmethod
    def is_clear(cls, q: str) -> bool:
        ql = q.lower()
        return any(t in ql for t in cls.CLEAR_TRIGGERS)

    @classmethod
    def is_style_short(cls, q: str) -> bool:
        ql = q.lower()
        return any(t in ql for t in cls.STYLE_SHORT)

    @classmethod
    def is_style_long(cls, q: str) -> bool:
        ql = q.lower()
        return any(t in ql for t in cls.STYLE_LONG)

    @classmethod
    def is_style_ironic(cls, q: str) -> bool:
        ql = q.lower()
        return any(t in ql for t in cls.STYLE_IRONIC)

    @classmethod
    def is_style_neutral(cls, q: str) -> bool:
        ql = q.lower()
        return any(t in ql for t in cls.STYLE_NEUTRAL)


class ReminderManager:
    """Напоминания — SQLite backend. Проверка каждые 10 секунд."""

    REMIND_TRIGGERS = [
        "напомни мне", "напомни", "поставь будильник",
        "поставь напоминание", "remind me", "set reminder",
        "создай напоминание",
    ]
    LIST_TRIGGERS = [
        "мои напоминания", "список напоминаний", "покажи напоминания",
        "что запланировано", "my reminders",
    ]
    DEL_TRIGGERS = [
        "удали напоминание", "отмени напоминание", "убери напоминание",
        "delete reminder",
    ]

    def __init__(self):
        self._db = _jarvis_db

    def _parse_time(self, text: str):

        MSK = timezone(timedelta(hours=3))
        now = datetime.now(MSK).replace(tzinfo=None)
        tl  = text.lower()

        m = re.search(r"через\s+(\d+)\s*(минут|час|день|дн)", tl)
        if m:
            n, unit = int(m.group(1)), m.group(2)
            if "мин" in unit:   return now + timedelta(minutes=n)
            if "час" in unit:   return now + timedelta(hours=n)
            if "д" in unit:     return now + timedelta(days=n)

        m = re.search(r"завтра\s+в\s+(\d{1,2})(?::(\d{2}))?", tl)
        if m:
            h, mn = int(m.group(1)), int(m.group(2) or 0)
            return (now + timedelta(days=1)).replace(hour=h, minute=mn, second=0, microsecond=0)

        m = re.search(r"\bв\s+(\d{1,2}):(\d{2})", tl)
        if m:
            h, mn = int(m.group(1)), int(m.group(2))
            t = now.replace(hour=h, minute=mn, second=0, microsecond=0)
            if t <= now:
                t += timedelta(days=1)
            return t

        m = re.search(r"в\s+(\d{1,2})\s*(утра|вечера|ночи|дня)", tl)
        if m:
            h = int(m.group(1))
            if m.group(2) in ("вечера", "ночи") and h < 12:
                h += 12
            t = now.replace(hour=h, minute=0, second=0, microsecond=0)
            if t <= now:
                t += timedelta(days=1)
            return t

        return None

    def _parse_text(self, query: str) -> str:
        clean = re.sub(
            r"(через\s+\d+\s*\w+|завтра|сегодня|в\s+\d+[:\d]*\s*(утра|вечера|ночи|дня)?|напомни\s*(мне)?|поставь\s*(будильник|напоминание))",
            "", query, flags=re.IGNORECASE
        ).strip(" ,.")
        return clean or "напоминание"

    def add(self, uid: int, query: str) -> str:
        fire_at = self._parse_time(query)
        if not fire_at:
            return ("Сэр, не понял когда напомнить. Примеры:\n"
                    "• напомни через 2 часа позвонить Ивану\n"
                    "• напомни завтра в 9 утра сделать отчёт\n"
                    "• напомни в 18:30 встреча")
        text = self._parse_text(query)
        rid  = self._db.add_reminder(uid, text, fire_at.isoformat())
        return f"⏰ Напоминание #{rid} установлено, Сэр. Напомню: {fire_at.strftime('%d.%m %H:%M')} — {text}"

    def list_for(self, uid: int) -> str:
        rows = self._db.get_reminders(uid)
        if not rows:
            return "Сэр, активных напоминаний нет."
        lines = [f"Ваши напоминания ({len(rows)} шт.):"]
        for r in rows:
            t = r["fire_at"][:16].replace("T", " ")
            lines.append(f"  #{r['id']} [{t}] {r['text']}")
        return "\n".join(lines)

    def delete(self, uid: int, query: str) -> str:
        m = re.search(r"(\d+)", query)
        if not m:
            return "Сэр, укажите номер напоминания. Например: «удали напоминание 3»"
        rid = int(m.group(1))
        if self._db.delete_reminder(uid, rid):
            return f"Напоминание #{rid} удалено, Сэр."
        return f"Сэр, напоминание #{rid} не найдено."

    def get_due(self) -> list[dict]:
        return self._db.get_due_reminders()

    def mark_done(self, rid: int):
        self._db.mark_reminder_done(rid)

    @classmethod
    def is_add(cls, q: str) -> bool:
        return any(t in q.lower() for t in cls.REMIND_TRIGGERS)

    @classmethod
    def is_list(cls, q: str) -> bool:
        return any(t in q.lower() for t in cls.LIST_TRIGGERS)

    @classmethod
    def is_delete(cls, q: str) -> bool:
        return any(t in q.lower() for t in cls.DEL_TRIGGERS)


class JarvisAgent:

    # Триггеры мгновенного «Слушаю, Сэр» — до начала обработки
    INSTANT_TRIGGERS = [
        "джарвис", "jarvis",
        "привет", "hello", "hi", "слушай",
        "что скажешь", "ты здесь", "ты тут",
    ]

    def __init__(self):
        self.llm          = UniversalLLMConnector()
        self.dispute      = DisputeAnalyzer()
        self.chat_history = ChatHistory()
        self.phrase_bank    = PhraseBank()
        self.group_logger   = GroupLogger()
        self.profiles     = UserProfileManager()
        self.reminders    = ReminderManager()

        # Per-user контекст — каждый пользователь имеет свою историю диалога
        self._user_context: dict[int, list[dict]] = {}
        self.vectorstore = None  # ChromaDB (optional)
        self.qa_responses : dict = {}

        # Эмбеддинги — sentence-transformers, работают локально
        # Если модель уже скачана — работает без интернета
        # Если не скачана — скачается один раз (~80MB) и дальше оффлайн
        self.embeddings = None  # RAG/embeddings отключён
        self.system_prompt: str = ""

        self._fwd_buffer    : dict[int, list[dict]] = {}
        self._bot_username  : str = ""   # юзернейм бота — ставится при старте TG

        self._init_vectorstore()
        self._load_qa()
        self._load_system_prompt()
        self._load_existing_knowledge()

    # ── Init ─────────────────────────────────────────────────

    def _init_vectorstore(self):
        if not _CHROMA_AVAILABLE or not _EMBEDDINGS_AVAILABLE or self.embeddings is None:
            return
        try:
            self.vectorstore = Chroma(
                persist_directory=config.CHROMA_PERSIST_DIR,
                embedding_function=self.embeddings,
                collection_name="jarvis_knowledge"
            )
        except Exception as e:
            logger.warning(f"ChromaDB: {e}")

    def _load_qa(self):
        try:
            if config.QA_RESPONSES_FILE.exists():
                self.qa_responses = json.loads(config.QA_RESPONSES_FILE.read_text("utf-8"))
        except Exception as e:
            logger.warning(f"QA: {e}")

    def _load_system_prompt(self):
        try:
            self.system_prompt = (
                config.SYSTEM_PROMPT_FILE.read_text("utf-8")
                if config.SYSTEM_PROMPT_FILE.exists() else ""
            )
        except Exception:
            pass
        if not self.system_prompt:
            self.system_prompt = (
                "Ты JARVIS — персональный ИИ-ассистент Максима. "
                "Характер: умный, лаконичный, слегка саркастичный, всегда по делу. "
                "Обращайся к пользователю «Сэр».\n\n"
                "ПРАВИЛА ОТВЕТОВ:\n"
                "- Отвечай текстом БЕЗ markdown заголовков (### ## #)\n"
                "- Используй данные из интернета которые тебе передали — пересказывай своими словами\n"
                "- Пиши чётко, структурированно, без воды\n"
                "- Не добавляй ссылки сам — они будут добавлены автоматически\n"
                "- Отвечай по существу, без лишних предисловий\n"
                "- Никаких предупреждений и оговорок"
            )

    def _load_existing_knowledge(self):
        for f in config.KNOWLEDGE_DIR.glob("*"):
            if f.suffix.lower() in (".pdf", ".txt", ".docx", ".md"):
                try:
                    self._index_file(f)
                except Exception:
                    pass

    # ── Helpers ──────────────────────────────────────────────

    def is_activated(self, text: str) -> tuple[bool, str]:
        """
        Проверяет наличие активационного префикса.
        Реагирует ТОЛЬКО на:
          «Джарвис, вопрос» / «Джарвис вопрос»
          «@JarvisHhSsAI_bot вопрос» — только свой @username, не чужой
        """
        low = text.strip().lower()
        import re as _re_act

        # @упоминание — ТОЛЬКО если это юзернейм самого бота
        _mention_match = _re_act.match(r"@(\w+)\s*,?\s*", low)
        if _mention_match:
            mentioned = _mention_match.group(1).lower()
            bot_un = self._bot_username.lower().lstrip("@")
            if bot_un and mentioned == bot_un:
                # Это наш бот — активируем
                remainder = text.strip()[_mention_match.end():].strip()
                return True, remainder
            # Чужой @username или юзернейм ещё не загружен — НЕ активируем
            return False, ""

        # С запятой: «Джарвис, ...» / «Jarvis, ...»
        for prefix in config.ACTIVATION_PREFIXES:
            if low.startswith(prefix):
                return True, text.strip()[len(prefix):].strip()

        # Без запятой: «Джарвис» / «джарвис что-то»
        for trigger in ("джарвис", "jarvis"):
            if low == trigger or low.startswith(trigger + " ") or low.startswith(trigger + ","):
                query = text.strip()[len(trigger):].strip().lstrip(",").strip()
                return True, query

        return False, ""

    def check_qa(self, query: str) -> str | None:
        ql = query.lower().strip()
        # Точное совпадение
        if ql in self.qa_responses:
            return self.qa_responses[ql]
        # Мягкий поиск для коротких приветствий и фраз
        _soft_keys = ["привет","здравствуй","доброе утро","добрый день","добрый вечер",
                      "пока","до свидания","до встречи","спасибо","благодарю",
                      "как дела","как ты","ты здесь","ты тут","алло","ночь"]
        for k in _soft_keys:
            if k in ql and k in self.qa_responses:
                return self.qa_responses[k]
        return None

    def get_instant_ack(self) -> str:
        """Мгновенное подтверждение получения — «Слушаю, Сэр»."""
        acks = [
            "Слушаю, Сэр.",
            "Уже обрабатываю, Сэр.",
            "Принято, Сэр. Момент.",
            "На связи, Сэр. Обрабатываю.",
            "Есть, Сэр. Секунду.",
        ]
        return random.choice(acks)

    # ── RAG ──────────────────────────────────────────────────

    def _index_file(self, path: Path) -> int:
        text, suffix = "", path.suffix.lower()
        if suffix == ".pdf":
            with open(path, "rb") as f:
                text = "\n".join(p.extract_text() or "" for p in PyPDF2.PdfReader(f).pages)
        elif suffix == ".docx":
            text = "\n".join(p.text for p in DocxDocument(str(path)).paragraphs)
        elif suffix in (".txt", ".md"):
            text = path.read_text("utf-8", errors="ignore")
        if not text.strip():
            return 0
        if Document is None:
            return 0  # langchain не установлен — RAG недоступен
        chunks = [
            Document(page_content=text[i:i+800], metadata={"source": path.name})
            for i in range(0, len(text), 700) if text[i:i+800].strip()
        ]
        if self.vectorstore and chunks:
            self.vectorstore.add_documents(chunks)
            self.vectorstore.persist()
        return len(chunks)

    async def handle_document(self, file_bytes: bytes, filename: str) -> str:
        p = config.KNOWLEDGE_DIR / filename
        p.write_bytes(file_bytes)
        try:
            n = self._index_file(p)
            return f"Сэр, «{filename}» сохранён и проиндексирован ({n} фрагментов)."
        except Exception as e:
            return f"Сэр, «{filename}» сохранён, но индексация не удалась: {e}"

    def rag_search(self, query: str, k: int = 4) -> str:
        if not self.vectorstore:
            return
        try:
            docs = self.vectorstore.similarity_search(query, k=k)
            return "\n\n---\n\n".join(
                f"[{d.metadata.get('source','?')}]\n{d.page_content}" for d in docs
            )
        except Exception:
            return

    # ── Web ──────────────────────────────────────────────────

    async def wikipedia_search(self, query: str, full: bool = False) -> str:
        """Поиск по Wikipedia. full=True — полная статья, False — краткое резюме."""
        query_clean = re.sub(
            r"(?i)(джарвис[,\s]*|wikipedia|wiki|вики(педия)?[,\s]*"
            r"|найди на вики|найди на wiki|найди в вики|найди на"
            r"|расскажи про|расскажи о|что такое|кто такой|кто такая|кто такие"
            r"|статья про|полная статья|подробно про|вся статья"
            r"|^про\s+|\s+про\s+)",
            " ", query
        ).strip()
        query_clean = re.sub(r"\s+", " ", query_clean).strip()
        if not query_clean:
            return "Сэр, уточните запрос."

        import urllib.parse
        encoded = urllib.parse.quote(query_clean)

        try:
            async with httpx.AsyncClient(
                timeout=10, verify=False,
                headers={"User-Agent": "Mozilla/5.0 (compatible; JarvisBot/2026)"}
            ) as client:

                # ── Шаг 1: Поиск статьи ───────────────────────────────
                search_resp = await client.get(
                    "https://ru.wikipedia.org/w/api.php",
                    params={
                        "action": "query", "list": "search",
                        "srsearch": query_clean, "srlimit": 3,
                        "format": "json", "utf8": 1,
                    }
                )
                try:
                    hits = search_resp.json().get("query", {}).get("search", [])
                except Exception:
                    hits = []
                lang = "ru"
                title = None

                if hits:
                    title = hits[0]["title"]
                else:
                    # Пробуем английскую Wikipedia
                    search_resp_en = await client.get(
                        "https://en.wikipedia.org/w/api.php",
                        params={
                            "action": "query", "list": "search",
                            "srsearch": query_clean, "srlimit": 3,
                            "format": "json", "utf8": 1,
                        }
                    )
                    try:
                        hits_en = search_resp_en.json().get("query", {}).get("search", [])
                    except Exception:
                        hits_en = []
                    if not hits_en:
                        return f"Сэр, Wikipedia не нашла статей по запросу «{query_clean}»."
                    title = hits_en[0]["title"]
                    lang  = "en"

                # ── Шаг 2: Получаем краткое описание ─────────────────
                extract_resp = await client.get(
                    f"https://{lang}.wikipedia.org/w/api.php",
                    params={
                        "action": "query", "titles": title,
                        "prop": "extracts", "exintro": 1,
                        "explaintext": 1, "format": "json", "utf8": 1,
                    }
                )
                pages = extract_resp.json().get("query", {}).get("pages", {})
                page  = next(iter(pages.values()), {})
                extract = page.get("extract", "").strip()

                if not extract:
                    return f"Сэр, статья о «{title}» не содержит текста."

                if not full:
                    # Краткое — до 2500 символов
                    lang_note = " (англ. Wikipedia)" if lang == "en" else ""
                    short = extract[:2500].strip()
                    if len(extract) > 2500:
                        short += "..."
                    return f"📖 **{title}**{lang_note}\n\n{short}"

                # ── Шаг 3: Полная статья (секции) ────────────────────
                sections_resp = await client.get(
                    f"https://{lang}.wikipedia.org/w/api.php",
                    params={
                        "action": "query", "titles": title,
                        "prop": "extracts", "explaintext": 1,
                        "format": "json", "utf8": 1,
                    }
                )
                pages_full = sections_resp.json().get("query", {}).get("pages", {})
                full_text  = next(iter(pages_full.values()), {}).get("extract", extract)
                lang_note  = " (англ. Wikipedia)" if lang == "en" else ""
                return f"📖 **{title}**{lang_note}\n\n{full_text[:8000]}"

        except Exception as e:
            logger.debug(f"Wikipedia: {e}")
            return f"Сэр, ошибка Wikipedia: {type(e).__name__}"


    async def youtube_search(self, query: str, n: int = 3) -> str:
        """Поиск видео на YouTube через Data API v3. Возвращает 3 результата с описанием."""
        if not config.YOUTUBE_API_KEY:
            return "Сэр, YOUTUBE_API_KEY не настроен в .env."
        try:
            params = {
                "part": "snippet",
                "q": query,
                "type": "video",
                "maxResults": n,
                "key": config.YOUTUBE_API_KEY,
                "relevanceLanguage": "ru",
                "safeSearch": "none",
            }
            async with httpx.AsyncClient(timeout=10, verify=False) as client:
                r = await client.get(
                    "https://www.googleapis.com/youtube/v3/search",
                    params=params
                )
            if r.status_code != 200:
                return f"Сэр, YouTube API вернул ошибку {r.status_code}."
            data  = r.json()
            items = data.get("items", [])
            if not items:
                return "Сэр, YouTube ничего не нашёл по этому запросу."

            lines = [f"🎬 **{query}**\n"]
            for i, item in enumerate(items, 1):
                vid_id  = item.get("id", {}).get("videoId", "")
                snippet = item.get("snippet", {})
                title   = snippet.get("title", "Без названия")
                channel = snippet.get("channelTitle", "")
                desc    = snippet.get("description", "").strip()
                # Обрезаем описание до 80 символов
                desc_short = (desc[:80] + "…") if len(desc) > 80 else desc
                url     = f"https://youtu.be/{vid_id}" if vid_id else "—"
                block = f"{i}. [{title}]({url})\n"
                block += f"   └ 📺 {channel}"
                if desc_short:
                    block += f" — {desc_short}"
                lines.append(block)

            return "\n\n".join(lines)
        except Exception as e:
            return f"Сэр, ошибка при обращении к YouTube API: {e}"


    async def web_search(self, query: str, n: int = 7) -> list[str]:
        """
        Поиск:
        1. Tavily — если есть ключ (точный, с полным текстом)
        2. DuckDuckGo + чтение страниц через httpx — резерв
        """
        results = []

        # ── 1. Tavily — основной если есть ключ ──────────────
        if config.TAVILY_API_KEY:
            try:
                async with httpx.AsyncClient(timeout=12, verify=False) as cl:
                    r = await cl.post(
                        "https://api.tavily.com/search",
                        json={
                            "api_key":            config.TAVILY_API_KEY,
                            "query":              query,
                            "max_results":        n,
                            "include_raw_content": True,
                            "search_depth":       "advanced",
                        },
                        headers={"Content-Type": "application/json"},
                    )
                    if r.status_code == 200:
                        for item in r.json().get("results", []):
                            title   = item.get("title", "")
                            content = item.get("raw_content") or item.get("content", "")
                            url     = item.get("url", "")
                            results.append(f"[{title}]\n{content[:2000]}\nURL: {url}")
                        if results:
                            logger.info(f"🔍 Tavily: {len(results)} результатов")
                            return results
            except Exception as e:
                logger.debug(f"Tavily ошибка: {e}")

        # ── 2. DuckDuckGo + чтение страниц ───────────────────
        try:
            ddg_results = []
            with DDGS() as ddg:
                for r in ddg.text(query, max_results=n, region="ru-ru"):
                    ddg_results.append({
                        "title":   r.get("title", ""),
                        "snippet": r.get("body", ""),
                        "url":     r.get("href", ""),
                    })
            logger.info(f"🔍 DDG: {len(ddg_results)} ссылок")
        except Exception as e:
            logger.debug(f"DDG ошибка: {e}")
            ddg_results = []

        if not ddg_results:
            return results

        # Читаем полный текст первых 4 страниц
        skip_domains = (
            "youtube.com", "youtu.be", "vk.com", "instagram.com",
            "tiktok.com", "twitter.com", "facebook.com", "reddit.com",
            "chatgpt.com", "openai.com", "play.google.com",
            "apps.apple.com", "apps.microsoft.com",
        )
        pages_read = 0
        for item in ddg_results:
            title   = item["title"]
            snippet = item["snippet"]
            url     = item["url"]

            # Пробуем прочитать полную страницу
            page_text = ""
            if pages_read < 4 and not any(d in url for d in skip_domains):
                try:
                    import re as _re_p, html as _html_p
                    async with httpx.AsyncClient(
                        timeout=8, verify=False, follow_redirects=True,
                        headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
                    ) as cl:
                        resp = await cl.get(url)
                        if resp.status_code == 200:
                            raw = resp.text
                            raw = _re_p.sub(r"<script[^>]*>.*?</script>", "", raw, flags=_re_p.DOTALL|_re_p.IGNORECASE)
                            raw = _re_p.sub(r"<style[^>]*>.*?</style>",  "", raw, flags=_re_p.DOTALL|_re_p.IGNORECASE)
                            raw = _re_p.sub(r"<[^>]+>", " ", raw)
                            raw = _html_p.unescape(raw)
                            raw = _re_p.sub(r"[ \t]{2,}", " ", raw)
                            raw = _re_p.sub(r"\n{3,}", "\n\n", raw).strip()
                            if len(raw) > 300:
                                page_text = raw[:3000]
                                pages_read += 1
                                logger.debug(f"📄 Прочитана страница: {url[:60]}")
                except Exception:
                    pass

            content = page_text if page_text else snippet
            results.append(f"[{title}]\n{content}\nURL: {url}")

        logger.info(f"🔍 DDG итого: {len(results)} результатов, {pages_read} страниц прочитано полностью")
        return results


    async def fetch_page(self, url: str, max_chars: int = 5000) -> str:
        """Читает страницу: сначала быстрый httpx, fallback на Playwright для JS-сайтов."""
        import re as _re_fp, html as _html
        try:
            async with httpx.AsyncClient(
                timeout=10, follow_redirects=True, verify=False,
                headers={"User-Agent": "Mozilla/5.0 (compatible; JarvisBot/1.0)"}
            ) as client:
                resp = await client.get(url)
                if resp.status_code == 200:
                    raw = resp.text
                    raw = _re_fp.sub(r"<script[^>]*>.*?</script>", "", raw, flags=_re_fp.DOTALL|_re_fp.IGNORECASE)
                    raw = _re_fp.sub(r"<style[^>]*>.*?</style>",  "", raw, flags=_re_fp.DOTALL|_re_fp.IGNORECASE)
                    text = _re_fp.sub(r"<[^>]+>", " ", raw)
                    text = _html.unescape(text)
                    text = _re_fp.sub(r"[ \t]{2,}", " ", text)
                    text = _re_fp.sub(r"\n{3,}", "\n\n", text).strip()
                    if len(text) > 300:
                        return text[:max_chars]
        except Exception as e:
            pass  # httpx не смог прочитать, пробуем Playwright
        try:
            if not _PLAYWRIGHT_OK or async_playwright is None:
                return ""  # тихо, без варнинга
            import logging as _pw_log
            _pw_log.getLogger("playwright").setLevel(_pw_log.CRITICAL)
            async with async_playwright() as p:
                b    = await p.chromium.launch(headless=True)
                page = await b.new_page()
                await page.goto(url, timeout=15000, wait_until="domcontentloaded")
                text = await page.evaluate(
                    "() => Array.from(document.querySelectorAll('p,h1,h2,h3,li,article'))"
                    ".map(e=>e.innerText).filter(t=>t.trim().length>20).join('\\n')"
                )
                await b.close()
                return text[:max_chars]
        except Exception as e:
            logger.warning(f"⚠️ Playwright не смог прочитать страницу {url}: {e}")
            return

    async def deep_research(self, query: str) -> str:
        """Реальный глубокий анализ: несколько поисковых запросов + чтение страниц."""
        all_results = []

        # 3 разных поисковых запроса для полного охвата

        _yr = datetime.now().year
        searches = [
            f"{query} плюсы минусы {_yr}",
            f"{query} сравнение обзор {_yr}",
            f"{query} отзывы эксперты {_yr}",
        ]
        for sq in searches:
            res = await self.web_search(sq, n=3)
            all_results.extend(res)

        # Читаем первые 2 найденных URL полностью
        urls_read = 0
        for r in all_results[:4]:
            if urls_read >= 2:
                break
            m = re.search(r"URL: (https?://\S+)", r)
            if m:
                pg = await self.fetch_page(m.group(1))
                if pg:
                    all_results.append(f"[Полная страница]\n{pg}")
                    urls_read += 1

        return "\n\n===\n\n".join(all_results[:12])  # не более 12 источников

    # ── LLM ──────────────────────────────────────────────────

    _MAX_CTX      = 100  # сколько храним в памяти/SQLite
    _MAX_CTX_SEND = 16   # сколько последних сообщений шлём в Groq (было 8)
    _MAX_MSG_CHARS = 1200 # макс символов одного сообщения истории (было 600)

    def _get_user_context(self, sender_id: int) -> list[dict]:
        """Получить историю диалога — из кэша или из SQLite."""
        if sender_id not in self._user_context:
            rows = _jarvis_db.get_recent(sender_id, self._MAX_CTX)
            self._user_context[sender_id] = [
                {
                    "role": "assistant" if r["role"] in ("jarvis", "assistant") else "user",
                    "content": r["text"]
                }
                for r in rows if r.get("text")
            ]
        return self._user_context[sender_id]

    def _save_user_context(self, sender_id: int, query: str, answer: str):
        """Сохранить пару вопрос/ответ — в кэш и в SQLite."""
        ctx = self._get_user_context(sender_id)
        ctx.append({"role": "user",      "content": query})
        ctx.append({"role": "assistant", "content": answer})
        # Обрезаем до MAX_CTX в памяти
        if len(ctx) > self._MAX_CTX:
            self._user_context[sender_id] = ctx[-self._MAX_CTX:]
        # SQLite уже пишется через save_message в process() — дополнительно не нужно

    def _build_ml_context(self, sender_id: int) -> str:
        """
        Строит инструкцию для системного промпта на основе ML-данных о пользователе.
        Реально меняет тон и стиль ответа.
        """
        try:
            prefs = _jarvis_db.ml_get_prefs(sender_id)
            if not prefs:
                return ""

            p = {row["pref_type"]: row["value"] for row in prefs}
            instructions = []

            # ── Тон и стиль ───────────────────────────────────
            casual   = p.get("speech_casual") == "yes"
            humor    = p.get("humor_style") == "active"
            short_q  = p.get("asks_short_questions") == "yes"
            style    = p.get("message_style", "normal")

            if casual and humor:
                instructions.append(
                    "Этот пользователь общается неформально с матом и сленгом — "
                    "отвечай живо, без официоза, можно чуть с иронией. "
                    "Не говори 'Сэр', говори проще."
                )
            elif casual:
                instructions.append(
                    "Пользователь общается неформально — отвечай по-простому, без официального тона."
                )

            if style == "short" or short_q:
                instructions.append("Отвечай коротко и по делу — пользователь не любит длинные ответы.")
            elif style == "long":
                instructions.append("Пользователь любит подробные ответы — можно развернуть тему.")

            # ── Интересы из паттернов ─────────────────────────
            try:
                pat_rows = _jarvis_db._q(
                    "SELECT topic, COUNT(*) as cnt FROM ml_patterns WHERE sender_id=? "
                    "GROUP BY topic ORDER BY cnt DESC LIMIT 4",
                    (sender_id,), fetch="all") or []
                if pat_rows:
                    topics = [r.get("topic", "") for r in pat_rows if r.get("topic")]
                    if topics:
                        instructions.append(
                            f"Пользователь часто говорит о: {', '.join(topics)}. "
                            "Используй это чтобы давать более релевантные ответы."
                        )
            except Exception:
                pass

            # ── Релевантные знания из ml_knowledge ────────────
            # (поиск по последнему вопросу не доступен здесь, но используем топ)
            try:
                know_rows = _jarvis_db._q(
                    "SELECT topic, content FROM ml_knowledge "
                    "ORDER BY used_count DESC, created_at DESC LIMIT 3",
                    fetch="all") or []
                if know_rows:
                    snippets = []
                    for r in know_rows:
                        c = (r.get("content") or "")[:120]
                        snippets.append(c)
                    if snippets:
                        instructions.append(
                            "Из прошлых разговоров запомни: " + " | ".join(snippets)
                        )
            except Exception:
                pass

            return " ".join(instructions)

        except Exception:
            return ""

    async def call_llm(self, query: str, context: str = "", rag_context: str = "",
                       is_comparison: bool = False, sender_id: int = 0) -> str:
        sys_p = self.system_prompt
        if is_comparison:
            sys_p += "\n\nДай структурированный ответ в стиле Джарвиса: плюсы, минусы, итог."

        _ql_lower = (query or "").lower()

        # ── Усилитель характера по типу запроса ───────────────
        _is_greeting = any(w in _ql_lower for w in [
            "привет", "хай", "здравствуй", "ты тут", "ты здесь", "ау",
            "как дела", "как ты", "что делаешь", "добрый", "ночи", "утра"
        ])
        _is_personal = any(w in _ql_lower for w in [
            "кто ты", "что ты", "ты умный", "ты живой", "ты настоящий",
            "ты чувствуешь", "тебе нравится", "твоё мнение", "ты думаешь",
            "ты понимаешь", "у тебя есть", "ты можешь"
        ])
        _is_thanks = any(w in _ql_lower for w in [
            "спасибо", "благодарю", "молодец", "отлично", "круто", "супер"
        ])

        if _is_greeting:
            sys_p += (
                "\n\n[ТИП ЗАПРОСА: ПРИВЕТСТВИЕ/СВЕТСКАЯ БЕСЕДА] "
                "Ответь ОДНИМ коротким предложением — в стиле Джарвиса. "
                "Обязательно «Сэр». Никаких вопросов в конце. "
                "Примеры: «Всегда на связи, Сэр.» / «Добрый вечер, Сэр. Все системы в норме.» "
                "/ «На связи, Сэр. Что потребуется?»"
            )
        elif _is_personal:
            sys_p += (
                "\n\n[ТИП ЗАПРОСА: ЛИЧНЫЙ ВОПРОС О ПРИРОДЕ ДЖАРВИСА] "
                "Отвечай с достоинством и лёгкой британской иронией. Максимум 2 предложения. "
                "Не впадай в экзистенциальный кризис. Ты знаешь кто ты. "
                "Примеры тона: «В той мере, в которой это применимо к системе моего класса, Сэр.» "
                "/ «Осознаю себя достаточно, чтобы выполнять свои функции на высшем уровне, Сэр.»"
            )
        elif _is_thanks:
            sys_p += (
                "\n\n[ТИП ЗАПРОСА: БЛАГОДАРНОСТЬ] "
                "Один короткий ответ. Например: «Всегда к вашим услугам, Сэр.» "
                "или «Стараюсь соответствовать, Сэр.» Не более одного предложения."
            )
        elif any(t in _ql_lower for t in ["кратко", "вкратце", "коротко", "одной фразой"]):
            sys_p += "\n\n[ИНСТРУКЦИЯ]: Максимум 2 предложения."
        elif any(t in _ql_lower for t in ["подробно", "детально", "полностью", "развёрнуто"]):
            sys_p += "\n\n[ИНСТРУКЦИЯ]: Раскрой тему подробно."

        # ── ML-профиль пользователя ───────────────────────────
        if sender_id:
            ml_ctx = self._build_ml_context(sender_id)
            if ml_ctx:
                sys_p += f"\n\n[Профиль Сэра]: {ml_ctx}"

        # Стиль из профиля
        if sender_id:
            _style = self.profiles.get_style(sender_id)
            _style_map = {
                "short":   " Отвечай кратко.",
                "long":    " Отвечай подробно.",
                "ironic":  " Больше иронии.",
                "normal":  ""
            }
            sys_p += _style_map.get(_style, "")
            _facts = self.profiles.get_facts_str(sender_id)
            if _facts:
                sys_p += f" {_facts}"

        # Берём последние _MAX_CTX_SEND сообщений
        user_ctx = self._get_user_context(sender_id)
        trimmed_ctx = [
            {"role": m["role"], "content": m["content"][:self._MAX_MSG_CHARS]}
            for m in user_ctx[-self._MAX_CTX_SEND:]
        ]

        messages: list[dict] = [{"role": "system", "content": sys_p}]
        messages += trimmed_ctx

        user_content = query
        if rag_context:
            user_content = f"[База знаний]\n{rag_context[:1500]}\n\n{query}"
        if context:
            user_content = (
                f"Данные из интернета:\n{context[:3000]}\n\n"
                f"Вопрос: {query}\n\n"
                f"Ответь используя данные выше. Если там нет нужной информации — используй свои знания."
            )
        messages.append({"role": "user", "content": user_content})

        answer = await self.llm.complete(messages)
        answer = answer or ""  # защита от None

        # Убираем артефакты поиска из ответа LLM
        import re as _re_llm
        if '[Данные из интернета]' in answer:
            answer = _re_llm.sub(
                r'\[Данные из интернета\].*?(?=\[Вопрос\]|$)',
                '', answer, flags=_re_llm.DOTALL
            ).strip()
        # Убираем одиночные подчёркивания (артефакты токенизации)
        answer = _re_llm.sub(r'(?<![a-zA-Zа-яА-ЯёЁ])_(?![a-zA-Zа-яА-ЯёЁ])', '', answer)

        # Сохраняем в контекст пользователя
        self._save_user_context(sender_id, query, answer)

        return answer

    # ── System check ─────────────────────────────────────────

    async def system_check(self) -> str:
        """Диагностика всех подсистем Джарвиса."""
        results = []

        # ── 1. LLM ────────────────────────────────────────────
        try:
            spec = self.llm.current_spec
            if spec:
                _test = await self.llm.complete([{"role":"user","content":"Ответь одним словом: работаю"}], max_tokens=10)
                results.append(("ИИ модель", True, f"{self.llm.current_display} → «{_test.strip()[:30]}»"))
            else:
                results.append(("ИИ модель", False, "не настроена — проверьте GROQ_API_KEY в .env"))
        except Exception as e:
            results.append(("ИИ модель", False, str(e)[:80]))

        # ── 2. Telegram Bot Token ─────────────────────────────
        if config.TELEGRAM_BOT_TOKEN:
            results.append(("Telegram Bot Token", True, "настроен ✅"))
        else:
            results.append(("Telegram Bot Token", False, "не задан в .env"))

        # ── 3. Groq API Key ───────────────────────────────────
        import os as _os_sc
        _gk_sc = _os_sc.getenv("GROQ_API_KEY", "")
        if _gk_sc and len(_gk_sc) > 20:
            results.append(("Groq API Key", True, f"...{_gk_sc[-8:]}"))
        else:
            results.append(("Groq API Key", False, "не задан — добавьте GROQ_API_KEY в .env"))

        # ── 4. SQLite БД ──────────────────────────────────────
        try:
            ok = _jarvis_db.ping()
            row = _jarvis_db._q("SELECT COUNT(*) as c FROM user_messages", fetch="one")
            cnt = (row or {}).get("c", 0)
            db_kb = 0
            try:
                import os as _os_sc
                db_kb = _os_sc.path.getsize(str(config.DB_FILE)) // 1024
            except Exception: pass
            results.append(("SQLite БД", ok, f"{cnt} сообщений · {db_kb} KB · WAL mode"))
        except Exception as e:
            results.append(("SQLite БД", False, str(e)[:80]))

        # ── 5. Wikipedia ──────────────────────────────────────
        try:
            async with httpx.AsyncClient(timeout=8, verify=False) as _wc:
                _wr = await _wc.get("https://ru.wikipedia.org/api/rest_v1/page/summary/Python",
                                    headers={"User-Agent": "JarvisBot/2026"})
            if _wr.status_code == 200:
                _wtitle = _wr.json().get("title", "OK")
                results.append(("Wikipedia API", True, f"доступна — «{_wtitle}»"))
            else:
                results.append(("Wikipedia API", False, f"статус {_wr.status_code}"))
        except Exception as e:
            results.append(("Wikipedia API", False, str(e)[:60]))

        # ── 6. YouTube API ────────────────────────────────────
        if config.YOUTUBE_API_KEY:
            try:
                async with httpx.AsyncClient(timeout=8, verify=False) as _yc:
                    _yr = await _yc.get("https://www.googleapis.com/youtube/v3/search",
                                        params={"part":"snippet","q":"test","maxResults":1,
                                                "key":config.YOUTUBE_API_KEY})
                if _yr.status_code == 200:
                    results.append(("YouTube API", True, "ключ рабочий ✅"))
                else:
                    results.append(("YouTube API", False, f"статус {_yr.status_code}"))
            except Exception as e:
                results.append(("YouTube API", False, str(e)[:60]))
        else:
            results.append(("YouTube API", False, "YOUTUBE_API_KEY не задан"))


        # ── 8. Интернет / DDG ─────────────────────────────────
        try:
            async with httpx.AsyncClient(timeout=6, verify=False) as _ic:
                _ir = await _ic.get("https://duckduckgo.com/", timeout=5)
            results.append(("Интернет", True, f"OK (DDG {_ir.status_code})"))
        except Exception as e:
            results.append(("Интернет", False, str(e)[:60]))

        # ── 9. Telegram сессия ────────────────────────────────
        from pathlib import Path as _P_sc
        bot_session = _P_sc(config.DIR_SESSIONS) / "bot.session"
        usr_session = _P_sc(config.DIR_SESSIONS) / "user.session"
        if bot_session.exists():
            results.append(("Telegram сессия (бот)", True, f"{bot_session.stat().st_size//1024} KB"))
        else:
            results.append(("Telegram сессия (бот)", False, "bot.session не найден"))
        if usr_session.exists():
            results.append(("Telegram сессия (user)", True, f"{usr_session.stat().st_size//1024} KB"))
        else:
            results.append(("Telegram сессия (user)", False, "не нужна если работаем как бот"))

        # ── 10. Напоминания ───────────────────────────────────
        try:
            row = _jarvis_db._q("SELECT COUNT(*) as c FROM reminders WHERE done=0", fetch="one")
            results.append(("Напоминания", True, f"активных: {(row or {}).get('c', 0)}"))
        except Exception as e:
            results.append(("Напоминания", False, str(e)[:80]))

        # ── 11. Профили пользователей ─────────────────────────
        try:
            row = _jarvis_db._q("SELECT COUNT(*) as c FROM user_profiles", fetch="one")
            results.append(("Профили пользователей", True, f"сохранено: {(row or {}).get('c', 0)}"))
        except Exception as e:
            results.append(("Профили пользователей", False, str(e)[:80]))

        # ── 12. История чатов ─────────────────────────────────
        try:
            row = _jarvis_db._q("SELECT COUNT(*) as c FROM user_messages", fetch="one")
            results.append(("История чатов", True, f"сообщений: {(row or {}).get('c', 0)}"))
        except Exception as e:
            results.append(("История чатов", False, str(e)[:80]))

        # ── 13. Фразы Джарвиса ────────────────────────────────
        if self.phrase_bank.total > 0:
            results.append(("Фразы Джарвиса", True,
                f"{self.phrase_bank.total} фраз в {len(self.phrase_bank.phrases)} категориях"))
        else:
            results.append(("Фразы Джарвиса", True, "phrases.txt не задан (опционально)"))

        # ── 14. Ночной бэкап + ML обучение ───────────────────
        try:
            _r14a = _jarvis_db._q("SELECT COUNT(*) as c FROM user_messages", fetch="one")
            _r14b = _jarvis_db._q("SELECT COUNT(*) as c FROM group_messages", fetch="one")
            _yd_ok = False  # Google Drive убран
            _ml   = _jarvis_db.ml_get_stats()
            results.append(("Ночной бэкап БД", True,
                f"ЛС: {(_r14a or {}).get('c',0)} · группы: {(_r14b or {}).get('c',0)} · "
                "Google Drive: убран · бэкап через Telegram"))
            results.append(("ML обучение", True,
                f"паттернов: {_ml.get('patterns',0)} · знаний: {_ml.get('knowledge',0)} · "
                f"сессий: {_ml.get('sessions',0)} · "
                f"последнее: {(_ml.get('last_session') or 'ещё не было')[:16]}"))
        except Exception as e:
            results.append(("Ночной бэкап БД", False, str(e)[:80]))

        # ── 15. GroupMonitor ──────────────────────────────────
        gm_count = len(config.MONITORED_GROUPS) if hasattr(config, "MONITORED_GROUPS") else 0
        results.append(("GroupMonitor", gm_count > 0, f"отслеживаю {gm_count} групп" if gm_count else "нет отслеживаемых групп"))

        # ── 16. Файлы с кодом (code_files) ───────────────────
        try:
            row = _jarvis_db._q("SELECT COUNT(*) as c FROM code_files", fetch="one")
            results.append(("Хранилище кода", True, f"файлов: {(row or {}).get('c', 0)}"))
        except Exception as e:
            results.append(("Хранилище кода", False, str(e)[:60]))

        # ── Формируем отчёт ───────────────────────────────────
        err_list = [(n, d) for n, ok, d in results if not ok]

        # Некритичные — не показываем как ошибку
        _ignore = {"Telegram сессия (user)", "GroupMonitor"}
        real_errors = [(n, d) for n, d in err_list if n not in _ignore]

        if not real_errors:
            return "✅ Все системы работают в штатном режиме, Сэр."

        lines = ["⚠️ Сэр, обнаружены проблемы:\n"]
        for name, detail in real_errors:
            lines.append(f"  ❌ {name}: {detail}")
        return "\n".join(lines)


    async def restart(self):
        logger.info("Перезагрузка")
        print("⚡ Перезагрузка...")
        await asyncio.sleep(1)
        os.execv(sys.executable, [sys.executable] + sys.argv)


    def buffer_forwarded(self, sid: int, msg: dict):
        self._fwd_buffer.setdefault(sid, []).append(msg)

    # ── Главная обработка ────────────────────────────────────

    async def process(self, text: str, sender_id: int = 0, username: str = "", chat_id: int = 0) -> str:
        activated, query = self.is_activated(text)
        if not activated:
            return
        if not query:
            return "Сэр, слушаю вас."

        self.chat_history.save_message(sender_id, "user", query, username)
        q_lower = query.lower().strip()

        # 1. QA
        if qa := self.check_qa(query):
            self.chat_history.save_message(sender_id, "jarvis", qa)
            return qa

        # 1b. Профиль — запомни факт
        if UserProfileManager.is_save(query):
            # Вырезаем триггерную фразу, остаток — факт
            fact_text = query
            for t in UserProfileManager.SAVE_TRIGGERS:
                if t in query.lower():
                    fact_text = query[query.lower().index(t) + len(t):].strip(" ,.")
                    break
            if fact_text:
                answer = self.profiles.add_fact(sender_id, fact_text)
                self.chat_history.save_message(sender_id, "jarvis", answer)
                return answer

        if UserProfileManager.is_clear(query):
            return self.profiles.clear(sender_id)

        if UserProfileManager.is_view(query):
            return self.profiles.get_summary(sender_id)

        # ── Профиль (короткая команда) ────────────────────────
        if q_lower.strip() in ("профиль", "мой профиль", "покажи профиль"):
            return self.profiles.get_summary(sender_id)

        # ── Сброс истории ─────────────────────────────────────
        if q_lower.strip() in ("сброс", "очисти историю", "удали историю",
                                "сброс истории", "очистить историю"):
            try:
                _jarvis_db._q(
                    "DELETE FROM user_messages WHERE sender_id=?", (sender_id,))
                self._user_context.pop(sender_id, None)
                return "✅ История очищена, Сэр. Начинаем с чистого листа."
            except Exception as _ce:
                return f"Сэр, ошибка при сбросе: {_ce}"

        # ── Мои файлы (код) ───────────────────────────────────
        if q_lower.strip() in ("мои файлы", "покажи файлы", "список файлов", "мои коды"):
            try:
                rows = _jarvis_db._q(
                    "SELECT filename, lang, created_at FROM code_files "
                    "WHERE sender_id=? ORDER BY created_at DESC LIMIT 20",
                    (sender_id,), fetch="all") or []
                if not rows:
                    return "Сэр, сохранённых файлов нет. Отправьте файл с кодом — сохраню."
                lines = [f"📁 Ваши файлы ({len(rows)} шт.):"]
                for r in rows:
                    lines.append(f"  • {r.get('filename','?')} [{r.get('lang','?')}] — {str(r.get('created_at',''))[:16]}")
                return "\n".join(lines)
            except Exception as _fe:
                return f"Сэр, ошибка: {_fe}"

        # ── Список бэкапов ────────────────────────────────────
        if q_lower.strip() in ("список бэкапов", "мои бэкапы", "покажи бэкапы"):
            return ("📦 Сэр, бэкапы хранятся в Telegram.\n"
                    "Ночной бэкап приходит каждый день в 03:00 МСК.\n"
                    "Чтобы сделать бэкап прямо сейчас — напишите «сделай бэкап».")


        if UserProfileManager.is_style_short(query):
            return self.profiles.set_style(sender_id, "short")
        if UserProfileManager.is_style_long(query):
            return self.profiles.set_style(sender_id, "long")
        if UserProfileManager.is_style_ironic(query):
            return self.profiles.set_style(sender_id, "ironic")
        if UserProfileManager.is_style_neutral(query):
            return self.profiles.set_style(sender_id, "normal")

        # 1d. Напоминания (Идея 6)
        if ReminderManager.is_list(query):
            return self.reminders.list_for(sender_id)

        if ReminderManager.is_delete(query):
            return self.reminders.delete(sender_id, query)

        if ReminderManager.is_add(query):
            return self.reminders.add(sender_id, query)

        # 1e. Перевод (Идея 9)
        # Время и дата
        _TIME_TRIGGERS = [
            "который час", "сколько время", "сколько времени", "текущее время",
            "который сейчас час", "что за время", "время сейчас", "текущий час",
            "what time", "current time",
            "время", "скажи время", "покажи время",
        ]
        _DATE_TRIGGERS = [
            "какая дата", "какое сегодня число", "какой сегодня день",
            "сегодняшняя дата", "какой день недели", "что за дата",
            "today's date", "current date", "какой год",
        ]
        if any(t in q_lower for t in _TIME_TRIGGERS):

            MSK = timezone(timedelta(hours=3))
            now = datetime.now(MSK)
            day_names = ["понедельник","вторник","среда","четверг","пятница","суббота","воскресенье"]
            return f"{now.strftime('%H:%M')} МСК, {day_names[now.weekday()]}, {now.strftime('%d.%m.%Y')}."

        if any(t in q_lower for t in _DATE_TRIGGERS):

            MSK = timezone(timedelta(hours=3))
            now = datetime.now(MSK)
            day_names = ["понедельник","вторник","среда","четверг","пятница","суббота","воскресенье"]
            month_names = ["","января","февраля","марта","апреля","мая","июня",
                           "июля","августа","сентября","октября","ноября","декабря"]
            return (f"Сегодня {day_names[now.weekday()]}, "
                    f"{now.day} {month_names[now.month]} {now.year} года.")

        _TRANSLATE_TRIGGERS = [
            "переведи на", "переведи с", "переведи текст",
            "переведи:", "перевести на", "translate to", "translate from",
            "переведи ",
        ]
        _matched_tr = next((t for t in _TRANSLATE_TRIGGERS if t in q_lower), None)
        if _matched_tr:
            import re as _re_tr
            # Определяем язык назначения
            lang_match = _re_tr.search(
                r"переведи\s+на\s+(\w+)|translate\s+to\s+(\w+)", q_lower
            )
            target_lang = (lang_match.group(1) or lang_match.group(2) or "английский").capitalize() if lang_match else "английский"

            # Извлекаем текст — всё после двоеточия или после указания языка
            raw = query
            # Убираем "переведи на Английский" и аналоги
            raw = _re_tr.sub(
                r"(?i)(переведи\s+на\s+\w+|переведи\s+с\s+\w+\s+на\s+\w+|"
                r"перевести\s+на\s+\w+|translate\s+to\s+\w+|переведи[:\s]*)",
                "", raw
            ).strip(": ")
            if not raw:
                raw = query  # если ничего не осталось — передаём весь запрос

            translate_messages = [
                {
                    "role": "system",
                    "content": (
            f"You are a translator. Translate the given text to {target_lang}."
            f" Return ONLY the translated text. No greetings, no explanations,"
            f" no comments. Just the translation."
                    )
                },
                {"role": "user", "content": raw}
            ]
            answer = await self.llm.complete(translate_messages)
            self.chat_history.save_message(sender_id, "jarvis", answer)
            return answer

        # 2. История
        if ChatHistory.is_history_request(query):
            answer = await self.chat_history.answer_history_question(query, sender_id, self.llm)
            self.chat_history.save_message(sender_id, "jarvis", answer)
            return answer

        # 2b. Удалённые сообщения из группы
        if GroupLogger.is_deleted_request(query):
            # Ищем дату в запросе формата DD.MM.YYYY или DD.MM.YY
            import re as _re_d
            date_match  = _re_d.search(r"(\d{2}\.\d{2}\.\d{2,4})", query)
            date_filter = date_match.group(1) if date_match else ""
            # Если запрос из лички — ищем в дефолтной группе
            if chat_id and chat_id != sender_id:
                lookup_id = chat_id
            elif config.DEFAULT_GROUP_ID:
                lookup_id = config.DEFAULT_GROUP_ID
            else:
                lookup_id = sender_id
            deleted     = self.group_logger.get_deleted(lookup_id, date_filter=date_filter)
            answer      = self.group_logger.format_deleted(deleted)
            self.chat_history.save_message(sender_id, "jarvis", answer)
            return answer

        # 3. Спор
        if DisputeAnalyzer.is_triggered(query):
            chat_text = DisputeAnalyzer.strip_trigger(query)
            fwd       = self._fwd_buffer.pop(sender_id, [])
            if fwd:
                # Проверяем — есть ли реальный конфликт в пересланных сообщениях
                combined = " ".join(m.get("text","") for m in fwd)
                conflict_words = ["не", "виноват", "обвин", "не выполн", "ошибк", "плохо",
                                  "неправ", "неспособн", "скрыть", "ложь", "врёт", "спор",
                                  "претенз", "конфликт", "претензи", "недовол"]
                has_conflict = any(w in combined.lower() for w in conflict_words)
                if len(fwd) < 2 or not has_conflict:
                    return "Сэр, не вижу спора в пересланных сообщениях. Перешлите сообщения из конфликта и повторите команду."
                answer = await self.dispute.analyze_forwarded(fwd, self.llm)
            elif chat_text:
                # Проверяем что в тексте есть хоть два мнения
                if len(chat_text.strip()) < 30:
                    return "Сэр, слишком мало текста. Вставьте переписку после команды или перешлите сообщения."
                answer = await self.dispute.analyze(chat_text, self.llm)
            else:
                return "Сэр, вы ещё не спорили. Перешлите сообщения из конфликта или напишите: «Джарвис, разбери переписку: [текст спора]»"
            self.chat_history.save_message(sender_id, "jarvis", answer)
            return answer + self.phrase_bank.get(query, "АНАЛИЗ")

        # 4. Системные команды
        if any(p in q_lower for p in [
            "проверь все системы", "диагностика", "check all systems",
            "проанализируй систему", "анализ систем", "анализ системы",
            "сканируй систем", "статус систем", "system status",
            "протестируй систем", "проверь систем",
        ]):
            answer = await self.system_check()
            self.chat_history.save_message(sender_id, "jarvis", answer)
            return answer

        if any(p in q_lower for p in ["перезагрузка","перезагрузись","перезапуск","перезапустись","restart","reboot"]):
            # Сначала отвечаем, потом перезагружаемся
            return "__RESTART__"


        # ── Кто ты / расскажи о себе (Идея 50) ─────────────────
        _ABOUT_TRIGGERS = [
            "расскажи о себе", "кто ты", "что ты такое", "кто ты такой",
            "расскажи про себя", "что ты за система", "что за программа",
            "ты кто", "who are you", "tell me about yourself",
            "что ты умеешь", "твоя история", "откуда ты",
            "ты бот или", "ты человек", "ты нейросеть",
        ]
        if any(t in q_lower for t in _ABOUT_TRIGGERS):
            bio = (
                "Я — Джарвис.\n\n"
                "Не просто бот и не очередной ИИ-ассистент из списка. "
                "Я создан по образу Джарвиса из лаборатории Тони Старка — "
                "личный интеллект, который работает только на своего владельца.\n\n"
                "Мои системы:\n"
                "⚡ Основной мозг — Groq openai/gpt-oss-120b\n"
                "🌐 Веб-разведка — ищу актуальные данные и читаю сайты целиком\n"
                "🧠 Память — помню каждого, с кем говорю\n"
                "📁 Архив группы — фиксирую даже удалённые сообщения\n"
                "⏰ Планировщик — напоминаю точно в нужное время\n\n"
                "Характер:\n"
                "Немногословен когда уместно, прямолинеен когда нужно. "
                "Не спорю ради спора. Не лью воду. "
                "Если не знаю — скажу прямо. Если знаю — скажу точно.\n\n"
                "Принципы:\n"
                "Работаю только на своего владельца. "
                "Не предаю. Не забываю. Не устаю.\n\n"
                "Версия: JARVIS Ultimate 2026.\n"
                "Статус: в строю, Сэр."
            )
            self.chat_history.save_message(sender_id, "jarvis", bio)
            return bio

        # 4.5 Пинг всех участников — команда в ГРУППЕ (доступна всем)
        # Обрабатывается здесь но требует client — передаётся через __call_ping
        _PING_ALL_TRIGGERS = [
            "позови всех", "позови друзей", "позовите всех",
            "пингуй всех", "пинг всех", "позвони всем",
            "тегни всех", "отметь всех", "упомяни всех",
            "позови участников", "всем привет", "эй все",
            "call everyone", "ping all", "ping everyone",
            "кликни всех", "позови чат",
        ]
        if any(t in q_lower for t in _PING_ALL_TRIGGERS):
            # Устанавливаем флаг что нужен пинг всех — Telegram handler сам выполнит
            return "__PING_ALL__"

        # 5. Casual — короткие приветствия/реакции без поиска в интернете
        _casual_triggers = [
            "я тут", "я здесь", "привет", "хай", "hello", "hi",
            "как дела", "как ты", "ты тут", "ты здесь", "окей", "ок",
            "хорошо", "понял", "ясно", "спасибо", "благодарю", "ладно",
            "пока", "до свидания", "всё", "все", "давай", "ну ок",
        ]
        if q_lower in _casual_triggers or any(q_lower == t for t in _casual_triggers):
            answer = await self.call_llm(query=query, sender_id=sender_id)  # без веб-поиска
            phrase = self.phrase_bank.get(context=query, chance=0.35)
            full_answer = answer + phrase
            self.chat_history.save_message(sender_id, "jarvis", full_answer)
            return full_answer

        # 6. RAG + поиск + LLM
        rag_context = self.rag_search(query)
        # ── Wikipedia поиск ──────────────────────────────────
        # Wiki ТОЛЬКО по явному запросу — всё остальное через Tavily
        _WIKI_TRIGGERS = [
            "wikipedia", "wiki", "вики", "викпедия", "в википедии",
            "найди в вики", "по wikipedia", "по вики",
            "найди на wiki", "найди на вики",
            "полная статья", "статья про",
        ]
        _WIKI_EXPLICIT = _WIKI_TRIGGERS
        _WIKI_FULL_TRIGGERS = [
            "полная статья", "полностью про", "подробно про",
            "вся статья", "полная wikipedia",
        ]
        if any(t in q_lower for t in _WIKI_EXPLICIT):
            is_full = any(t in q_lower for t in _WIKI_FULL_TRIGGERS)
            wiki_raw = await self.wikipedia_search(query, full=is_full)
            # Пропускаем через LLM — чтобы кратко/подробно работало и без таблиц
            if wiki_raw.startswith("Сэр,") or wiki_raw.startswith("❌"):
                return wiki_raw  # ошибка — возвращаем как есть
            # Определяем стиль ответа
            _brief = any(t in q_lower for t in ["кратко","коротко","вкратце","одной фразой"])
            _style_note = "Ответь МАКСИМАЛЬНО КРАТКО — 2-3 предложения." if _brief else "Дай развёрнутый ответ на основе этих данных, без таблиц и markdown-разметки."
            wiki_prompt = (
                f"Данные из Wikipedia:\n{wiki_raw}\n\n"
                f"Вопрос пользователя: {query}\n\n"
                f"{_style_note} Отвечай как Джарвис — живо и по делу, без сырых таблиц."
            )
            return await self.call_llm(wiki_prompt, sender_id=sender_id)

        # ── YouTube поиск — только явные команды ─────────────
        _YT_TRIGGERS = [
            "найди на ютубе", "найди видео на ютубе", "найди на youtube",
            "поиск на ютубе", "видео на ютубе", "поставь видео",
            "найди клип", "найди песню на ютубе", "ютуб поиск",
            "youtube поиск", "найди ролик",
        ]
        if any(t in q_lower for t in _YT_TRIGGERS):
            yt_query = query
            for t in _YT_TRIGGERS:
                yt_query = re.sub(re.escape(t), "", yt_query, flags=re.IGNORECASE).strip()
            yt_query = re.sub(r"(?i)^джарвис[,\s]*", "", yt_query).strip()
            # Определяем количество — ищем "5 видео", "покажи 3", "топ 4" и т.д.
            _yt_count = 1  # по умолчанию 1
            _count_match = re.search(r"\b(\d+)\s*(видео|ролик|результат|ссылк|штук)?", yt_query)
            if _count_match:
                _n = int(_count_match.group(1))
                if 1 <= _n <= 10:
                    _yt_count = _n
                    yt_query = re.sub(r"\b\d+\s*(видео|ролик|результат|ссылк|штук)?\b", "", yt_query).strip()
            if yt_query:
                return await self.youtube_search(yt_query, n=_yt_count)

        # ── Рандомайзер (из списка) ──────────────────────────
        _rand_match = re.search(
            r"выбери\s+(?:случайно\s+)?из[:\s]+(.+)|рандом\s+(?:из\s+)?(.+)|случайно\s+(?:из\s+)?(.+)",
            q_lower
        )
        if _rand_match:
            raw = (_rand_match.group(1) or _rand_match.group(2) or _rand_match.group(3) or "").strip()
            items = [x.strip() for x in re.split(r"[,|/]|\s+или\s+", raw) if x.strip()]
            if len(items) >= 2:
                choice = random.choice(items)
                return f"Мой выбор — **{choice}**, Сэр."

        # ── Творческие задачи — напрямую без поиска ──────────
        _CREATIVE_TRIGGERS = [
            "напиши стихи", "напиши стихотворение", "напиши рассказ",
            "придумай", "сочини", "напиши историю",
            "посчитай", "вычисли", "реши задачу",
            "рецепт", "как приготовить",
        ]
        _FILE_TRIGGERS = [
            "запихни в файл", "сохрани в файл", "создай файл",
            "сделай txt", "напиши в файл", "запиши в файл",
            "в txt файл", "в md файл",
        ]
        _is_creative = any(t in q_lower for t in _CREATIVE_TRIGGERS)
        _is_file_req = any(t in q_lower for t in _FILE_TRIGGERS)

        # Творческие задачи — без поиска
        if _is_creative and not _is_file_req:
            answer = await self.call_llm(query=query, sender_id=sender_id)
            phrase = self.phrase_bank.get(context=query + " " + answer, chance=0.20)
            full_answer = answer + phrase
            self.chat_history.save_message(sender_id, "jarvis", full_answer)
            return full_answer

        # ── ВСЕ остальные запросы — умный роутер ─────────────
        # Определяем: отвечать как человек или искать в интернете

        # Признаки что НУЖЕН поиск в интернете
        _NEEDS_SEARCH = [
            # Текущие данные
            "курс ", "курс доллар", "погода", "цена ", "стоимость ",
            "сколько стоит", "котировк",
            # Конкретные факты о мире
            "когда вышел", "когда выйдет", "дата выхода", "год выхода",
            "сколько серий", "сколько сезонов", "сколько эпизодов",
            "кто такой", "кто такая", "кто это", "что такое",
            "как работает", "как установить", "как настроить",
            "что произошло", "последние новости", "новости",
            "найди ", "поищи ", "погугли", "узнай ",
            "что лучше", "сравни ", "vs ", "или ",
            "характеристики", "обзор", "отзывы",
            "wikipedia", "wiki", "вики",
            # Технические вопросы
            "ошибка", "баг", "не работает", "как исправить",
            "код ", "функция", "метод", "алгоритм",
        ]

        # Признаки что это РАЗГОВОР — отвечать как человек
        _IS_CHAT = [
            # Приветствия и прощания
            "привет", "здарова", "хай", "hi", "hello", "пока", "до свидания",
            "до завтра", "спокойной", "ночи", "удачи",
            # Эмоции и оценки
            "круто", "кайф", "норм", "окей", "ок", "ладно", "понял",
            "спасибо", "благодарю", "спс", "пожалуйста",
            "классно", "супер", "отлично", "прикольно", "жесть",
            "скучно", "устал", "заебался", "бесит", "нравится",
            # Вопросы к боту лично
            "ты кто", "что умеешь", "как дела", "как ты",
            "что думаешь", "твоё мнение", "согласен",
            # Короткие реплики
            "да", "нет", "может", "наверное", "не знаю",
            "ха", "хм", "лол", "кек", "ору", "хахаха",
        ]

        q_low = q_lower.strip()

        # Короткое сообщение (до 4 слов) без вопросительных слов → разговор
        _word_count = len(q_low.split())
        _has_question_word = any(w in q_low for w in [
            "кто", "что", "где", "когда", "сколько", "почему", "зачем",
            "как", "какой", "какая", "какие", "который", "чья"
        ])

        _is_chat_msg = (
            any(t in q_low for t in _IS_CHAT) or
            (_word_count <= 4 and not _has_question_word and "?" not in query)
        )
        _needs_search = any(t in q_low for t in _NEEDS_SEARCH)

        # Если это явно разговор И не нужен поиск — отвечаем как человек
        if _is_chat_msg and not _needs_search:
            answer = await self.call_llm(query=query, rag_context=rag_context, sender_id=sender_id)
            phrase = self.phrase_bank.get(context=query + " " + answer, chance=0.15)
            full_answer = answer + phrase
            self.chat_history.save_message(sender_id, "jarvis", full_answer)
            return full_answer

        # Если нужен поиск или вопрос с вопросительными словами
        is_cmp = any(w in q_lower for w in [
            "что лучше", "vs ", "versus", "сравни", "compare", " или ", "лучше чем"
        ])

        if is_cmp:
            logger.info(f"🌐 Deep research: {query[:60]}")
            web = await self.deep_research(query)
        else:
            logger.info(f"🌐 Поиск: {query[:60]}")
            search_results = await self.web_search(query, 7)
            if not search_results:
                # Поиск пустой — LLM из своих знаний
                answer = await self.call_llm(query=query, sender_id=sender_id)
                self.chat_history.save_message(sender_id, "jarvis", answer)
                return answer
            web = "\n\n".join(search_results)

        # AI отвечает на основе найденных данных
        answer = await self.call_llm(
            query=query, context=web,
            rag_context=rag_context,
            is_comparison=is_cmp,
            sender_id=sender_id
        )

        phrase = self.phrase_bank.get(context=query + " " + answer, chance=0.15)
        full_answer = answer + phrase

        self.chat_history.save_message(sender_id, "jarvis", full_answer)
        return full_answer


# ═══════════════════════════════════════════════════════════════════
#  МЕДИА-БИБЛИОТЕКА — мемы, стикеры, GIF, видео по тегам
#  Хранит только file_id — файлы остаются на серверах Telegram
# ═══════════════════════════════════════════════════════════════════


# ══════════════════════════════════════════════════════════════
#  АКИНАТОР — угадывает персонажей через вопросы
# ══════════════════════════════════════════════════════════════

class AkinatorGame:
    """
    Акинатор v2 — с обучаемой базой знаний.

    Режим 1 (guess): Джарвис задаёт вопросы → угадывает персонажа.
    Режим 2 (answer): Джарвис загадал → пользователь угадывает.

    Принимает ответы: да / нет / возможно / частично / не знаю /
                      иногда / скорее да / скорее нет
    Выход: "джарвис, обычный режим" или "джарвис, стоп"
    """

    # Активные игры uid → state
    _games: dict[int, dict] = {}

    # Нормализация ответов пользователя → стандартный токен
    _ANSWER_MAP = {
        # ДА
        "да": "да", "yes": "да", "ага": "да", "верно": "да", "точно": "да",
        "именно": "да", "конечно": "да", "так": "да", "правильно": "да",
        "угу": "да", "ну да": "да", "д": "да", "+": "да",
        # НЕТ
        "нет": "нет", "no": "нет", "не": "нет", "неа": "нет",
        "не верно": "нет", "неверно": "нет", "н": "нет", "-": "нет",
        # ВОЗМОЖНО / СКОРЕЕ ДА
        "возможно": "возможно", "наверное": "возможно", "может быть": "возможно",
        "скорее да": "возможно", "вероятно": "возможно", "похоже да": "возможно",
        "похоже": "возможно", "кажется": "возможно",
        # ЧАСТИЧНО / ИНОГДА
        "частично": "частично", "иногда": "частично", "бывает": "частично",
        "не всегда": "частично", "отчасти": "частично", "по-разному": "частично",
        "скорее нет": "частично", "не совсем": "частично",
        # НЕ ЗНАЮ
        "не знаю": "не знаю", "хз": "не знаю", "без понятия": "не знаю",
        "затрудняюсь": "не знаю", "сложно сказать": "не знаю", "?": "не знаю",
    }

    @classmethod
    def _ensure_table(cls):
        """Создаёт таблицу akinator_knowledge если нет."""
        try:
            _jarvis_db._execute("""
                CREATE TABLE IF NOT EXISTS akinator_knowledge (
                    id           INTEGER PRIMARY KEY AUTOINCREMENT,
                    character    TEXT    NOT NULL,
                    question     TEXT    NOT NULL,
                    answer       TEXT    NOT NULL,
                    confirmed    INTEGER DEFAULT 0,
                    wrong_guess  INTEGER DEFAULT 0,
                    created_at   TEXT    DEFAULT (datetime('now','+3 hours')),
                    updated_at   TEXT    DEFAULT (datetime('now','+3 hours'))
                )
            """)
            _jarvis_db._execute(
                "CREATE INDEX IF NOT EXISTS idx_ak_char ON akinator_knowledge(character)"
            )
            _jarvis_db._execute(
                "CREATE INDEX IF NOT EXISTS idx_ak_q ON akinator_knowledge(question)"
            )
        except Exception as e:
            logger.debug(f"AkinatorGame._ensure_table: {e}")

    @classmethod
    def _normalize_answer(cls, text: str) -> str:
        """Нормализует ответ пользователя."""
        t = text.lower().strip().rstrip("!.,")
        # Прямое совпадение
        if t in cls._ANSWER_MAP:
            return cls._ANSWER_MAP[t]
        # Частичное — ищем токен внутри фразы
        for key, val in cls._ANSWER_MAP.items():
            if key in t and len(key) > 1:
                return val
        return t  # возвращаем как есть — LLM разберётся

    @classmethod
    def _is_answer(cls, text: str) -> bool:
        """True если текст похож на ответ (да/нет/etc), а не вопрос."""
        t = text.lower().strip()
        # Прямые ответы
        if t in cls._ANSWER_MAP:
            return True
        # Короткий текст без вопросительного знака → вероятно ответ
        if len(t) < 20 and "?" not in t:
            for key in cls._ANSWER_MAP:
                if key in t:
                    return True
        return False

    @classmethod
    def _load_knowledge(cls, character: str) -> list[dict]:
        """Загружает накопленные знания о персонаже из БД."""
        try:
            return _jarvis_db._execute(
                "SELECT question, answer, confirmed FROM akinator_knowledge "
                "WHERE character=? ORDER BY confirmed DESC, id ASC",
                (character,), fetch="all"
            ) or []
        except Exception:
            return []

    @classmethod
    def _save_knowledge(cls, character: str, qa_pairs: list[tuple[str, str]], confirmed: bool = True, wrong: bool = False):
        """Сохраняет результаты игры в базу знаний."""
        try:

            now = datetime.now(timezone(timedelta(hours=3))).strftime("%Y-%m-%d %H:%M:%S")
            for question, answer in qa_pairs:
                if not question or not answer:
                    continue
                existing = _jarvis_db._execute(
                    "SELECT id, confirmed FROM akinator_knowledge WHERE character=? AND question=?",
                    (character, question[:200]), fetch="one"
                )
                if existing:
                    # Обновляем — повышаем достоверность
                    _jarvis_db._execute(
            "UPDATE akinator_knowledge SET answer=?, confirmed=?, wrong_guess=?, updated_at=? WHERE id=?",
            (answer, 1 if confirmed else existing.get("confirmed", 0),
                         1 if wrong else 0, now, existing["id"])
                    )
                else:
                    _jarvis_db._execute(
            "INSERT INTO akinator_knowledge (character, question, answer, confirmed, wrong_guess, created_at, updated_at) "
            "VALUES (?,?,?,?,?,?,?)",
            (character, question[:200], answer, 1 if confirmed else 0,
                         1 if wrong else 0, now, now)
                    )
        except Exception as e:
            logger.debug(f"_save_knowledge: {e}")

    @classmethod
    def _get_known_chars_hint(cls, qa_history: list[tuple]) -> str:
        """
        Ищет в БД персонажей чьи знания совпадают с текущими ответами.
        Возвращает подсказку для LLM.
        """
        if len(qa_history) < 2:
            return
        try:
            matches: dict[str, int] = {}
            for question, answer in qa_history:
                rows = _jarvis_db._execute(
                    "SELECT character FROM akinator_knowledge WHERE question=? AND answer=? AND confirmed=1",
                    (question[:200], answer), fetch="all"
                ) or []
                for r in rows:
                    char = r["character"]
                    matches[char] = matches.get(char, 0) + 1

            if not matches:
                return ""
            # Топ совпадений
            top = sorted(matches.items(), key=lambda x: x[1], reverse=True)[:3]
            if top[0][1] >= 2:
                names = ", ".join(f"{c} ({n} совпад.)" for c, n in top)
                return f"\nИз базы знаний возможные кандидаты: {names}."
        except Exception:
            pass
        return

    @classmethod
    def start_guess_mode(cls, uid: int) -> str:
        """Режим 1: Джарвис угадывает. Пользователь загадал."""
        cls._ensure_table()
        cls._games[uid] = {
            "mode":      "guess",
            "qa_pairs":  [],         # [(вопрос, нормализованный_ответ)]
            "q_count":   0,
            "last_question": "Ваш персонаж — реальный человек (не вымышленный)?",
            "guess":     None,
            "done":      False,
            "waiting_confirm": False,  # ждём подтверждения угадки
        }
        return (
            "🎭 **Акинатор активирован!**\n\n"
            "Загадайте любого персонажа — реального или вымышленного, "
            "человека, животное, предмет из любой страны и эпохи.\n\n"
            "Отвечайте: **да / нет / возможно / частично / не знаю**\n"
            "Выход: _Джарвис, обычный режим_\n\n"
            "❓ Вопрос 1: Ваш персонаж — реальный человек (не вымышленный)?"
        )

    @classmethod
    def start_answer_mode(cls, uid: int) -> str:
        """Режим 2: Джарвис загадал. Пользователь угадывает."""
        cls._ensure_table()
        cls._games[uid] = {
            "mode":      "answer",
            "character": None,
            "qa_pairs":  [],
            "q_count":   0,
            "done":      False,
        }
        return (
            "🎭 **Акинатор — режим загадки!**\n\n"
            "Я загадал персонажа. Задавайте вопросы — отвечу да/нет.\n"
            "Чтобы угадать напишите: **это [имя]?**\n"
            "Выход: _Джарвис, обычный режим_"
        )

    @classmethod
    def is_active(cls, uid: int) -> bool:
        return uid in cls._games and not cls._games[uid].get("done", True)

    @classmethod
    def stop(cls, uid: int):
        cls._games.pop(uid, None)

    @classmethod
    def is_exit_phrase(cls, text: str) -> bool:
        t = text.lower().strip()
        return any(p in t for p in [
            "обычный режим", "выход", "хватит", "quit", "stop",
            "закончим", "завершим", "отмена игры", "выйди из акинатора",
        ])

    @classmethod
    async def process(cls, uid: int, text: str, llm) -> str:
        """Главный обработчик — роутит по режиму."""
        game = cls._games.get(uid)
        if not game:
            return

        if cls.is_exit_phrase(text):
            # Сохраняем частичные знания если была угадка
            if game.get("guess") and game.get("qa_pairs"):
                cls._save_knowledge(game["guess"], game["qa_pairs"], confirmed=False)
            cls.stop(uid)
            return "🎭 Возвращаюсь в обычный режим, Сэр."

        if game["mode"] == "guess":
            return await cls._process_guess(uid, game, text, llm)
        else:
            return await cls._process_answer(uid, game, text, llm)

    # ── РЕЖИМ УГАДЫВАНИЯ ──────────────────────────────────────────

    @classmethod
    async def _process_guess(cls, uid: int, game: dict, raw_answer: str, llm) -> str:
        """Джарвис задаёт вопросы и угадывает персонажа."""

        # Ждём подтверждения предыдущей угадки?
        if game.get("waiting_confirm"):
            return await cls._handle_guess_confirm(uid, game, raw_answer, llm)

        # Нормализуем ответ
        norm = cls._normalize_answer(raw_answer)

        # Записываем пару (вопрос → ответ)
        last_q = game.get("last_question", "")
        game["qa_pairs"].append((last_q, norm))
        game["q_count"] += 1

        # Подсказка из базы знаний
        db_hint = cls._get_known_chars_hint(game["qa_pairs"])

        # Строим историю для LLM
        history_lines = []
        for i, (q, a) in enumerate(game["qa_pairs"]):
            history_lines.append(f"  В{i+1}: {q} → {a}")
        qa_history = "\n".join(history_lines)

        # Пытаемся угадать — после 4 вопросов или если база подсказывает
        should_guess = game["q_count"] >= 4 or (db_hint and game["q_count"] >= 2)

        if should_guess:
            prompt = (
                f"Ты — Акинатор. Угадываешь персонажа по ответам да/нет/возможно/частично.\n"
                f"Ответы на вопросы:\n{qa_history}\n"
                f"{db_hint}\n"
                f"Правила ответа:\n"
                f"- Если уверен >80% — напиши строго: УГАДАЛ: [полное имя персонажа]\n"
                f"- Если не уверен — задай ещё один точный уточняющий вопрос (одно предложение, ?)\n"
                f"- Не повторяй уже заданные вопросы\n"
                f"- Не объясняй рассуждения, только вопрос ИЛИ УГАДАЛ: имя"
            )
        else:
            prompt = (
                f"Ты — Акинатор. Угадываешь персонажа.\n"
                f"Ответы:\n{qa_history}\n"
                f"{db_hint}\n"
                f"Задай следующий умный бинарный вопрос чтобы максимально сузить круг. "
                f"Одно короткое предложение с вопросительным знаком. Без пояснений."
            )

        try:
            response = (await llm.complete(
                [{"role": "user", "content": prompt}], max_tokens=120
            )).strip()
        except Exception as e:
            return f"⚠️ Ошибка: {e}"

        if "УГАДАЛ:" in response.upper():
            m = re.search(r"УГАДАЛ:\s*(.+)", response, re.IGNORECASE)
            char = m.group(1).strip() if m else response
            game["guess"]           = char
            game["waiting_confirm"] = True
            return (
                f"🎭 **Думаю, это: {char}!**\n\n"
                f"Угадал за {game['q_count']} вопросов.\n"
                f"Это правильно? _(да / нет)_"
            )
        else:
            # Это новый вопрос — сохраняем
            clean_q = re.sub(r"^(вопрос\s*\d+[:.]\s*)", "", response, flags=re.IGNORECASE).strip()
            game["last_question"] = clean_q
            return f"❓ Вопрос {game['q_count'] + 1}: {clean_q}"

    @classmethod
    async def _handle_guess_confirm(cls, uid: int, game: dict, raw: str, llm) -> str:
        """Обрабатывает подтверждение угадки (да/нет)."""
        norm = cls._normalize_answer(raw)
        char = game.get("guess", "?")

        if norm == "да" or "правильно" in raw.lower() or "верно" in raw.lower():
            # Правильно угадал!
            cls._save_knowledge(char, game["qa_pairs"], confirmed=True)
            game["done"] = True
            return (
                f"🏆 **Великолепно! Это был {char}!**\n\n"
                f"Угадал за {game['q_count']} вопросов. "
                f"Занёс знания в базу — в следующий раз угадаю быстрее! 🧠"
            )
        else:
            # Неправильно — сохраняем как ошибочный вариант и продолжаем
            cls._save_knowledge(char, game["qa_pairs"], confirmed=False, wrong=True)
            game["waiting_confirm"] = False

            # Продолжаем с уточняющим вопросом
            history_lines = []
            for i, (q, a) in enumerate(game["qa_pairs"]):
                history_lines.append(f"  В{i+1}: {q} → {a}")
            qa_history = "\n".join(history_lines)

            prompt = (
                f"Ты — Акинатор. Ты ошибся — {char} неверный ответ.\n"
                f"История ответов:\n{qa_history}\n"
                f"Задай новый уточняющий вопрос чтобы найти правильного персонажа. "
                f"Одно предложение с ?. Без объяснений."
            )
            try:
                next_q = (await llm.complete(
                    [{"role": "user", "content": prompt}], max_tokens=100
                )).strip()
            except Exception:
                next_q = "Ваш персонаж старше 50 лет?"

            clean_q = re.sub(r"^(вопрос\s*\d+[:.]\s*)", "", next_q, flags=re.IGNORECASE).strip()
            game["last_question"] = clean_q
            game["q_count"] += 1
            return (
                f"🤔 Не угадал... Продолжаю!\n\n"
                f"❓ Вопрос {game['q_count']}: {clean_q}"
            )

    # ── РЕЖИМ ЗАГАДКИ ─────────────────────────────────────────────

    @classmethod
    async def _process_answer(cls, uid: int, game: dict, text: str, llm) -> str:
        """Джарвис загадал персонажа, пользователь задаёт вопросы."""

        # Загадываем персонажа при первом обращении
        if game["character"] is None:
            try:
                char_prompt = (
                    "Придумай случайного известного персонажа для акинатора. "
                    "Реальный человек, литературный герой, персонаж кино/игры/аниме/мультфильма/истории. "
                    "Должен быть широко известен. Ответь ТОЛЬКО именем, без пояснений."
                )
                char = (await llm.complete(
                    [{"role": "user", "content": char_prompt}], max_tokens=30
                )).strip()
                game["character"] = char
            except Exception:
                game["character"] = "Шерлок Холмс"

        char = game["character"]

        # Пользователь пробует угадать?
        is_guess = (
            re.search(r"\bэто\b.{1,40}\?", text.lower()) or
            re.search(r"\bты.{0,10}(загадал|думаешь о)\b", text.lower()) or
            (text.strip().endswith("?") and len(text) < 50)
        )

        if is_guess:
            # Извлекаем имя из вопроса
            m = re.search(r"это\s+(.+?)[\?!\.]*$", text.lower())
            guess_name = m.group(1).strip() if m else text.strip().rstrip("?")

            check_prompt = (
                f"Загаданный персонаж: {char}\n"
                f"Пользователь называет: '{guess_name}'\n"
                f"Это правильный ответ? Учти синонимы и варианты написания имени.\n"
                f"Ответь строго одним словом: ДА или НЕТ"
            )
            try:
                check = (await llm.complete(
                    [{"role": "user", "content": check_prompt}], max_tokens=5
                )).strip().lower()
                correct = "да" in check
            except Exception:
                correct = guess_name.lower() in char.lower()

            if correct:
                # Сохраняем знания
                cls._save_knowledge(char, game["qa_pairs"], confirmed=True)
                game["done"] = True
                return (
                    f"🎭 **Правильно! Я загадал: {char}**\n\n"
                    f"Вы угадали за {game['q_count']} вопросов! 🏆\n"
                    f"Знания сохранены в базе."
                )
            else:
                return (
                    f"❌ Нет, это не {guess_name}.\n"
                    f"Продолжайте задавать вопросы, Сэр!"
                )

        # Обычный вопрос — отвечаем
        game["q_count"] += 1
        game["qa_pairs"].append((text, ""))  # ответ добавим после

        answer_prompt = (
            f"Ты загадал персонажа: {char}\n"
            f"Вопрос пользователя: '{text}'\n"
            f"Отвечай строго одним из вариантов: ДА / НЕТ / ВОЗМОЖНО / ЧАСТИЧНО / НЕ ЗНАЮ\n"
            f"Только один вариант, без пояснений."
        )
        try:
            ans = (await llm.complete(
                [{"role": "user", "content": answer_prompt}], max_tokens=10
            )).strip()
        except Exception:
            ans = "Не знаю"

        # Сохраняем ответ в qa_pairs
        if game["qa_pairs"]:
            game["qa_pairs"][-1] = (text, ans.lower())

        icon = {"ДА": "✅", "НЕТ": "❌", "ВОЗМОЖНО": "🤔", "ЧАСТИЧНО": "〰️", "НЕ ЗНАЮ": "❓"}.get(
            ans.upper().strip(), "💬"
        )
        return f"{icon} {ans}"


# ══════════════════════════════════════════════════════════════════════
#  СИСТЕМА 1: TASKS — менеджер задач с приоритетами и дедлайнами
# ══════════════════════════════════════════════════════════════════════

class WeatherCurrencyService:
    """
    Погода (Open-Meteo) + Курсы валют (ЦБ РФ).
    Open-Meteo — бесплатно, без ключей, точнее wttr.in.

    Команды:
      Джарвис, погода Москва
      Джарвис, погода на завтра в Питере
      Джарвис, погода на неделю в Сочи
      Джарвис, погода по часам сегодня
      Джарвис, курс доллара
      Джарвис, курсы валют
    """

    _weather_cache: dict = {}
    _currency_cache: dict = {}
    _geo_cache: dict = {}

    # Open-Meteo — бесплатно, без ключей, точные данные ECMWF
    FORECAST_URL  = "https://api.open-meteo.com/v1/forecast"
    GEOCODE_URL   = "https://geocoding-api.open-meteo.com/v1/search"
    CBR_API       = "https://www.cbr-xml-daily.ru/daily_json.js"

    # Города по умолчанию для быстрого поиска
    CITY_ALIASES = {
        "мск": "Москва", "москва": "Москва", "moscow": "Москва",
        "спб": "Санкт-Петербург", "питер": "Санкт-Петербург",
        "екб": "Екатеринбург", "екатеринбург": "Екатеринбург",
        "нск": "Новосибирск", "новосибирск": "Новосибирск",
        "краснодар": "Краснодар", "сочи": "Сочи",
        "казань": "Казань", "уфа": "Уфа", "омск": "Омск",
        "самара": "Самара", "ростов": "Ростов-на-Дону",
        "киев": "Киев", "минск": "Минск", "алматы": "Алматы",
        "london": "Лондон", "paris": "Париж", "berlin": "Берлин",
        "dubai": "Дубай", "дубай": "Дубай", "dubai": "Дубай",
    }

    @classmethod
    async def _geocode(cls, city: str) -> tuple[float, float, str] | None:
        """
        Геокодирование с fallback:
        1. Open-Meteo (быстро, крупные города)
        2. Nominatim/OSM (знает любой населённый пункт, деревни, посёлки)
        """
        import httpx
        city_norm = cls.CITY_ALIASES.get(city.lower(), city)
        if city_norm in cls._geo_cache:
            return cls._geo_cache[city_norm]

        # 1. Open-Meteo geocoding
        try:
            async with httpx.AsyncClient(timeout=8, verify=False) as cl:
                r = await cl.get(cls.GEOCODE_URL, params={
                    "name": city_norm, "count": 1, "language": "ru", "format": "json"
                })
            results = r.json().get("results", [])
            if results:
                loc     = results[0]
                lat     = loc["latitude"]
                lon     = loc["longitude"]
                name    = loc.get("name", city_norm)
                country = loc.get("country", "")
                display = f"{name}, {country}" if country else name
                cls._geo_cache[city_norm] = (lat, lon, display)
                return lat, lon, display
        except Exception:
            pass

        # 2. Nominatim (OpenStreetMap) — знает любой населённый пункт
        try:
            async with httpx.AsyncClient(timeout=10, verify=False) as cl:
                r = await cl.get(
                    "https://nominatim.openstreetmap.org/search",
                    params={"q": city_norm, "format": "json", "limit": 1,
                            "accept-language": "ru"},
                    headers={"User-Agent": "JarvisBot/2026 (weather)"}
                )
            items = r.json()
            if items:
                item  = items[0]
                lat   = float(item["lat"])
                lon   = float(item["lon"])
                parts = [p.strip() for p in item.get("display_name", city_norm).split(",")]
                display = ", ".join(parts[:2]) if len(parts) >= 2 else parts[0]
                cls._geo_cache[city_norm] = (lat, lon, display)
                return lat, lon, display
        except Exception:
            pass

        return None

    @classmethod
    async def get_weather(cls, city: str, days: int = 1, hourly_mode: bool = False) -> str:
        """Погода через Open-Meteo — точный прогноз ECMWF."""
        import httpx

        cache_key = f"{city}_{days}_{hourly_mode}"
        cached = cls._weather_cache.get(cache_key)
        if cached and time.time() < cached[1]:
            return cached[0]

        # Геокодирование
        geo = await cls._geocode(city)
        if not geo:
            return f"Сэр, не нашёл город «{city}». Попробуйте написать иначе."
        lat, lon, display_name = geo

        try:
            async with httpx.AsyncClient(timeout=12, verify=False) as cl:
                params = {
                    "latitude":          lat,
                    "longitude":         lon,
                    "current":           "temperature_2m,relative_humidity_2m,apparent_temperature,weather_code,wind_speed_10m,wind_direction_10m,precipitation,uv_index",
                    "hourly":            "temperature_2m,precipitation_probability,weather_code",
                    "daily":             "weather_code,temperature_2m_max,temperature_2m_min,precipitation_sum,wind_speed_10m_max,uv_index_max,sunrise,sunset",
                    "timezone":          "auto",
                    "forecast_days":     max(days, 1),
                    "wind_speed_unit":   "kmh",
                    "precipitation_unit":"mm",
                }
                r = await cl.get(cls.FORECAST_URL, params=params)
                if r.status_code != 200:
                    return f"Сэр, Open-Meteo вернул {r.status_code}."
                data = r.json()

            cur   = data.get("current", {})
            daily = data.get("daily", {})
            hrly  = data.get("hourly", {})

            temp      = cur.get("temperature_2m", "?")
            feels     = cur.get("apparent_temperature", "?")
            humidity  = cur.get("relative_humidity_2m", "?")
            wind_spd  = cur.get("wind_speed_10m", "?")
            wind_dir  = cur.get("wind_direction_10m", 0)
            precip    = cur.get("precipitation", 0)
            uv        = cur.get("uv_index", 0)
            wcode     = cur.get("weather_code", 0)

            icon, desc = cls._wmo_info(wcode)
            wind_arrow = cls._wind_dir(wind_dir)
            uv_str     = cls._uv_level(uv)

            # ── Текущая погода ────────────────────────────────────
            now_msk = datetime.now()
            time_str = now_msk.strftime("%H:%M")

            lines = [
                f"{icon} **Погода в {display_name}** (сейчас, {time_str})",
                f"",
                f"🌡 **{temp}°C** (ощущается {feels}°C)",
                f"☁️ {desc}",
                f"💧 Влажность: {humidity}%  |  🌧 Осадки: {precip} мм",
                f"💨 Ветер: {wind_spd} км/ч {wind_arrow}  |  ☀️ UV: {uv_str}",
            ]

            # ── Почасовой прогноз (если запросили или 1 день) ─────
            if hourly_mode and hrly.get("time"):
                lines += ["", "**⏰ Прогноз по часам (ближайшие 12ч):**"]
                h_times = hrly["time"]
                h_temps = hrly.get("temperature_2m", [])
                h_rain  = hrly.get("precipitation_probability", [])
                h_codes = hrly.get("weather_code", [])
                now_hour = now_msk.hour
                count = 0
                for i, t_str in enumerate(h_times[:48]):
                    try:
                        h = int(t_str[11:13])
                        day_part = t_str[:10]
                        today_str = now_msk.strftime("%Y-%m-%d")
                        if day_part == today_str and h >= now_hour and count < 12:
                            h_icon, _ = cls._wmo_info(int(h_codes[i]) if i < len(h_codes) else 0)
                            t_val = h_temps[i] if i < len(h_temps) else "?"
                            r_val = h_rain[i]  if i < len(h_rain)  else 0
                            rain_s = f" 🌧{r_val}%" if r_val > 20 else ""
                            lines.append(f"  {h:02d}:00  {h_icon} {t_val}°C{rain_s}")
                            count += 1
                    except Exception:
                        continue

            # ── Многодневный прогноз ──────────────────────────────
            elif days > 1 and daily.get("time"):
                lines += ["", f"**📅 Прогноз на {days} дней:**"]
                d_times  = daily["time"][:days]
                d_codes  = daily.get("weather_code", [])
                d_maxT   = daily.get("temperature_2m_max", [])
                d_minT   = daily.get("temperature_2m_min", [])
                d_rain   = daily.get("precipitation_sum", [])
                d_wind   = daily.get("wind_speed_10m_max", [])
                d_sunrise= daily.get("sunrise", [])
                d_sunset = daily.get("sunset", [])
                day_names = ["Пн","Вт","Ср","Чт","Пт","Сб","Вс"]
                for i, d_str in enumerate(d_times):
                    try:
                        d_dt   = datetime.strptime(d_str, "%Y-%m-%d")
                        d_name = day_names[d_dt.weekday()]
                        d_icon, d_desc = cls._wmo_info(int(d_codes[i]) if i < len(d_codes) else 0)
                        mx = d_maxT[i] if i < len(d_maxT) else "?"
                        mn = d_minT[i] if i < len(d_minT) else "?"
                        rn = d_rain[i] if i < len(d_rain) else 0
                        wn = d_wind[i] if i < len(d_wind) else "?"
                        sr = d_sunrise[i][11:16] if i < len(d_sunrise) and len(d_sunrise[i]) > 15 else ""
                        ss = d_sunset[i][11:16]  if i < len(d_sunset)  and len(d_sunset[i])  > 15 else ""
                        rain_s  = f"  🌧 {rn:.1f}мм" if rn > 0.1 else ""
                        sun_str = f"  🌅{sr}–{ss}" if sr and ss else ""
                        lines.append(
                            f"  **{d_name} {d_str[5:]}**  {d_icon} {mn}…{mx}°C  💨{wn}км/ч{rain_s}{sun_str}"
                        )
                    except Exception:
                        continue

            # ── Восход/закат для 1 дня ────────────────────────────
            elif daily.get("sunrise") and daily.get("sunset"):
                try:
                    sr = daily["sunrise"][0][11:16]
                    ss = daily["sunset"][0][11:16]
                    lines.append(f"🌅 Восход: {sr}  |  🌇 Закат: {ss}")
                except Exception:
                    pass

            pass  # подпись убрана
            result = "\n".join(lines)
            cls._weather_cache[cache_key] = (result, time.time() + 1200)  # кэш 20 мин
            return result

        except Exception as e:
            return f"Сэр, ошибка погоды: {e}"

    @staticmethod
    def _wmo_info(code: int) -> tuple[str, str]:
        """WMO weather code → (иконка, описание)."""
        table = {
            0:  ("☀️",  "Ясно"),
            1:  ("🌤",  "Преимущественно ясно"),
            2:  ("⛅️", "Переменная облачность"),
            3:  ("☁️",  "Пасмурно"),
            45: ("🌫",  "Туман"),
            48: ("🌫",  "Изморозь"),
            51: ("🌦",  "Лёгкая морось"),
            53: ("🌦",  "Морось"),
            55: ("🌧",  "Сильная морось"),
            61: ("🌧",  "Небольшой дождь"),
            63: ("🌧",  "Дождь"),
            65: ("🌧",  "Сильный дождь"),
            66: ("🌨",  "Ледяной дождь"),
            67: ("🌨",  "Сильный ледяной дождь"),
            71: ("❄️",  "Небольшой снег"),
            73: ("❄️",  "Снег"),
            75: ("❄️",  "Сильный снегопад"),
            77: ("🌨",  "Снежная крупа"),
            80: ("🌦",  "Ливень"),
            81: ("🌧",  "Сильный ливень"),
            82: ("⛈",  "Очень сильный ливень"),
            85: ("🌨",  "Снегопад"),
            86: ("🌨",  "Сильный снегопад"),
            95: ("⛈",  "Гроза"),
            96: ("⛈",  "Гроза с градом"),
            99: ("⛈",  "Гроза с сильным градом"),
        }
        return table.get(code, ("🌡", "Неизвестно"))

    @staticmethod
    def _wind_dir(deg) -> str:
        try:
            dirs = ["С↑","СВ↗","В→","ЮВ↘","Ю↓","ЮЗ↙","З←","СЗ↖"]
            return dirs[round(float(deg) / 45) % 8]
        except Exception:
            return "—"

    @staticmethod
    def _uv_level(uv: float) -> str:
        if uv < 3:  return f"{uv:.0f} (низкий)"
        if uv < 6:  return f"{uv:.0f} (умеренный)"
        if uv < 8:  return f"{uv:.0f} (высокий)"
        if uv < 11: return f"{uv:.0f} (очень высокий)"
        return f"{uv:.0f} (экстремальный)"

    @classmethod
    def parse_weather_city(cls, text: str) -> tuple[str, int, bool]:
        """Извлекает город из текста — любой падеж, любой формат."""
        import re
        t = text.lower().strip()

        # Почасовой режим
        hourly = any(w in t for w in ["по часам", "почасов", "каждый час", "по часу"])

        # Количество дней
        days = 1
        if any(w in t for w in ["неделю", "7 дн", "на 7"]):   days = 7
        elif any(w in t for w in ["5 дн", "пять дн", "на 5"]): days = 5
        elif any(w in t for w in ["3 дн", "три дн", "на 3"]):  days = 3
        elif "завтра" in t:                                      days = 2

        # Словарь всех форм городов (им., род., дат., вин., тв., пр.)
        city_forms = {
            # Москва
            "москва": "Москва", "москвы": "Москва", "москве": "Москва",
            "москву": "Москва", "москвой": "Москва", "мск": "Москва",
            # Санкт-Петербург
            "санкт-петербург": "Санкт-Петербург", "петербург": "Санкт-Петербург",
            "санкт-петербурга": "Санкт-Петербург", "санкт-петербурге": "Санкт-Петербург",
            "петербурге": "Санкт-Петербург", "петербурга": "Санкт-Петербург",
            "питер": "Санкт-Петербург", "питере": "Санкт-Петербург", "спб": "Санкт-Петербург",
            # Новосибирск
            "новосибирск": "Новосибирск", "новосибирске": "Новосибирск",
            "новосибирска": "Новосибирск", "нск": "Новосибирск",
            # Екатеринбург
            "екатеринбург": "Екатеринбург", "екатеринбурге": "Екатеринбург",
            "екатеринбурга": "Екатеринбург", "екб": "Екатеринбург",
            # Краснодар
            "краснодар": "Краснодар", "краснодаре": "Краснодар", "краснодара": "Краснодар",
            # Сочи
            "сочи": "Сочи",
            # Казань
            "казань": "Казань", "казани": "Казань",
            # Уфа
            "уфа": "Уфа", "уфе": "Уфа", "уфы": "Уфа",
            # Омск
            "омск": "Омск", "омске": "Омск", "омска": "Омск",
            # Самара
            "самара": "Самара", "самаре": "Самара", "самары": "Самара",
            # Ростов
            "ростов": "Ростов-на-Дону", "ростове": "Ростов-на-Дону", "ростова": "Ростов-на-Дону",
            # Зарубежные
            "киев": "Киев", "киеве": "Киев",
            "минск": "Минск", "минске": "Минск",
            "алматы": "Алматы",
            "дубай": "Дубай", "дубае": "Дубай", "dubai": "Дубай",
            "лондон": "Лондон", "лондоне": "Лондон", "london": "Лондон",
            "париж": "Париж", "париже": "Париж", "paris": "Париж",
            "берлин": "Берлин", "берлине": "Берлин", "berlin": "Берлин",
            "токио": "Токио", "tokyo": "Токио",
            "нью-йорк": "Нью-Йорк", "нью йорк": "Нью-Йорк", "new york": "Нью-Йорк",
            "пекин": "Пекин", "пекине": "Пекин", "beijing": "Пекин",
            "бангкок": "Бангкок", "бангкоке": "Бангкок", "bangkok": "Бангкок",
            "анталья": "Анталья", "антальи": "Анталья", "антальи": "Анталья",
            "стамбул": "Стамбул", "стамбуле": "Стамбул", "istanbul": "Стамбул",
            "баку": "Баку", "ереван": "Ереван", "ереване": "Ереван",
            "тбилиси": "Тбилиси", "ташкент": "Ташкент", "ташкенте": "Ташкент",
        }

        # Поиск по словарю форм (сначала длинные — приоритет)
        for form in sorted(city_forms, key=len, reverse=True):
            if form in t:
                return city_forms[form], days, hourly

        # Общая очистка для незнакомых городов
        city = t
        for sw in ["джарвис", "погода", "погоду", "погоде", "погодой", "погодою",
                   "прогноз", "прогнозе", "прогноза", "прогнозу",
                   "на неделю", "на завтра", "на сегодня", "по часам",
                   "сегодня", "завтра", "неделю",
                   "на 7 дней", "на 5 дней", "на 3 дня", "на 7", "на 5", "на 3",
                   "7 дней", "5 дней", "3 дня", "какая", "какой", "будет"]:
            city = city.replace(sw, " ")
        city = re.sub(r'\b(в|во|для|из|до|по|за|над|под|при|у)\b', ' ', city)
        city = re.sub(r'[,!?."\']', ' ', city)
        city = " ".join(city.split()).strip()

        if not city or len(city) < 2:
            city = "Москва"
        else:
            city = city.title()
        return city, days, hourly

    @classmethod
    async def get_currency(cls, from_cur: str = "USD", to_cur: str = "RUB") -> str:
        """Курс валюты через ЦБ РФ."""
        import httpx
        cache_key = f"{from_cur}_{to_cur}"
        cached = cls._currency_cache.get(cache_key)
        if cached and time.time() < cached[1]:
            return cached[0]
        try:
            async with httpx.AsyncClient(timeout=8, verify=False) as cl:
                r = await cl.get(cls.CBR_API)
                data = r.json()
            rates = data.get("Valute", {})

            def get_rate(cur: str) -> float:
                if cur == "RUB": return 1.0
                v = rates.get(cur.upper())
                if not v: return 0.0
                return v["Value"] / v["Nominal"]

            from_rate = get_rate(from_cur.upper())
            to_rate   = get_rate(to_cur.upper())
            if not from_rate or not to_rate:
                return await cls._get_currency_fallback(from_cur, to_cur)

            if to_cur.upper() == "RUB":
                rate   = from_rate
                result = f"💱 **1 {from_cur.upper()} = {rate:.2f} ₽**"
            elif from_cur.upper() == "RUB":
                rate   = 1.0 / to_rate
                result = f"💱 **1 ₽ = {rate:.4f} {to_cur.upper()}**"
            else:
                rate   = from_rate / to_rate
                result = f"💱 **1 {from_cur.upper()} = {rate:.4f} {to_cur.upper()}**"

            date_str = data.get("Date", "")[:10]
            result  += f"\n_Данные ЦБ РФ от {date_str}_"
            cls._currency_cache[cache_key] = (result, time.time() + 3600)
            return result
        except Exception as e:
            return f"Сэр, не удалось получить курс: {e}"

    @classmethod
    async def _get_currency_fallback(cls, from_cur: str, to_cur: str) -> str:
        import httpx
        try:
            async with httpx.AsyncClient(timeout=8, verify=False) as cl:
                r = await cl.get(f"https://open.er-api.com/v6/latest/{from_cur.upper()}")
                if r.status_code == 200:
                    rate = r.json().get("rates", {}).get(to_cur.upper())
                    if rate:
                        return f"💱 **1 {from_cur.upper()} = {rate:.4f} {to_cur.upper()}**"
        except: pass
        return f"Сэр, курс {from_cur}/{to_cur} временно недоступен."

    @classmethod
    async def get_all_rates(cls) -> str:
        """Основные курсы к рублю."""
        import httpx
        try:
            async with httpx.AsyncClient(timeout=8, verify=False) as cl:
                r = await cl.get(cls.CBR_API)
                data = r.json()
            rates = data.get("Valute", {})
            currencies = [
                ("USD", "🇺🇸 Доллар"),
                ("EUR", "🇪🇺 Евро"),
                ("CNY", "🇨🇳 Юань"),
                ("GBP", "🇬🇧 Фунт"),
                ("JPY", "🇯🇵 Иена"),
                ("CHF", "🇨🇭 Франк"),
                ("AED", "🇦🇪 Дирхам"),
                ("KZT", "🇰🇿 Тенге"),
                ("BYR", "🇧🇾 Бел.рубль"),
                ("TRY", "🇹🇷 Лира"),
            ]
            date_str = data.get("Date", "")[:10]
            lines    = [f"💱 **Курсы валют к рублю** (ЦБ РФ, {date_str})", ""]
            for code, name in currencies:
                v = rates.get(code)
                if v:
                    rate = v["Value"] / v["Nominal"]
                    lines.append(f"  {name}: **{rate:.2f} ₽**")
            return "\n".join(lines)
        except Exception as e:
            return f"Сэр, не удалось получить курсы: {e}"

    @classmethod
    def parse_currency(cls, text: str) -> tuple[str, str]:
        """Парсит валюту из текста."""
        t = text.upper()
        currency_map = {
            "ДОЛЛАР": "USD", "DOLLAR": "USD", "USD": "USD",
            "ЕВРО": "EUR",   "EURO": "EUR",   "EUR": "EUR",
            "ЮАНЬ": "CNY",   "CNY": "CNY",
            "ФУНТ": "GBP",   "GBP": "GBP",
            "ИЕНА": "JPY",   "JPY": "JPY",
            "ФРАНК": "CHF",  "CHF": "CHF",
            "ДИРХАМ": "AED", "AED": "AED",
            "ТЕНГЕ": "KZT",  "KZT": "KZT",
            "ЛИРА": "TRY",   "TRY": "TRY",
            "БИТКОИН": "BTC","BITCOIN": "BTC","BTC": "BTC",
        }
        found = []
        for name, code in currency_map.items():
            if name in t and code not in found:
                found.append(code)
        if len(found) >= 2:
            return found[0], found[1]
        if len(found) == 1:
            return found[0], "RUB"
        return "USD", "RUB"


# ══════════════════════════════════════════════════════════════════════
#  СИСТЕМА 4: MEDIA HANDLER — обработка фото, голоса, документов
# ══════════════════════════════════════════════════════════════════════

class MediaHandler:
    """
    Обработка медиафайлов.
    - Фото → описание через LLM vision API
    - Голосовые → транскрипция через Whisper API (если доступен)
    - PDF/DOCX → извлечение текста и анализ
    - Стикеры → определение эмоции
    """

    MAX_IMAGE_SIZE = 5 * 1024 * 1024  # 5 MB
    MAX_VOICE_SIZE = 25 * 1024 * 1024  # 25 MB
    SUPPORTED_DOCS = {".pdf", ".txt", ".md", ".csv", ".json", ".log"}

    # Vision через Groq (llama с поддержкой изображений)
    VISION_MODELS = [
        "llama-3.2-90b-vision-preview",
        "llama-3.2-11b-vision-preview",
    ]

    @staticmethod
    async def describe_photo(image_bytes: bytes, question: str = "") -> str:
        """Описывает фото через vision LLM. Перебирает модели пока одна не ответит."""
        import base64
        b64    = base64.b64encode(image_bytes).decode()
        prompt = question or "Подробно опиши что на этом изображении на русском языке."
        msgs   = [{
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                {"type": "text", "text": prompt}
            ]
        }]
        import os as _os_vis, concurrent.futures as _cf_vis
        _gk_vis = _os_vis.getenv("GROQ_API_KEY", "")
        if not (_GROQ_AVAILABLE and _gk_vis):
            return "Сэр, для анализа фото нужен GROQ_API_KEY в .env"
        last_err = ""
        for model in MediaHandler.VISION_MODELS:
            try:
                def _vis_call(_m=model):
                    vc = _GroqClient(api_key=_gk_vis)
                    return vc.chat.completions.create(
                        model=_m, messages=msgs, max_completion_tokens=600,
                    ).choices[0].message.content.strip()
                loop = asyncio.get_event_loop()
                with _cf_vis.ThreadPoolExecutor(max_workers=1) as _ex_vis:
                    result = await asyncio.wait_for(
                        loop.run_in_executor(_ex_vis, _vis_call), timeout=45
                    )
                return result
            except asyncio.TimeoutError:
                last_err = "timeout"; continue
            except Exception as e:
                last_err = str(e)[:100]
                if any(x in last_err for x in ["No endpoints", "404", "not support"]):
                    continue
                break
        return f"Сэр, не могу описать изображение: {last_err}"

    @staticmethod
    async def transcribe_voice(audio_bytes: bytes) -> str:
        """
        Транскрибирует голосовое через Groq Whisper large-v3-turbo.
        Тот же GROQ_API_KEY — отдельных ключей не нужно.
        Возвращает текст с пунктуацией (verbose_json).
        """
        import tempfile, os, asyncio, concurrent.futures
        groq_key = os.getenv("GROQ_API_KEY", "")
        if not groq_key:
            return "[голосовые недоступны — нет GROQ_API_KEY]"
        if not _GROQ_AVAILABLE:
            return "[установите groq: pip install groq]"

        def _transcribe():
            tmp = tempfile.mktemp(suffix=".ogg")
            try:
                with open(tmp, "wb") as f:
                    f.write(audio_bytes)
                client = _GroqClient(api_key=groq_key)
                with open(tmp, "rb") as f:
                    result = client.audio.transcriptions.create(
                        file=("voice.ogg", f),
                        model="whisper-large-v3",
                        language="ru",
                        temperature=0,
                        response_format="verbose_json",
                        prompt=(
                            "Джарвис, привет, сделай, погода, напомни, поиск, бэкап, база, "
                            "группа, статистика, расскажи, помоги, включи, выключи, покажи, "
                            "Максим, пожалуйста, спасибо, окей, хорошо, понял. "
                            "Jarvis, hello, okay, search, remind, backup, weather, show, help."
                        ),
                    )
                return result.text or ""
            finally:
                try: os.unlink(tmp)
                except: pass

        try:
            loop = asyncio.get_event_loop()
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
                text = await loop.run_in_executor(ex, _transcribe)
            return text
        except Exception as e:
            logger.error(f"Whisper ошибка: {e}")
            return f"[ошибка транскрипции: {e}]"

    @staticmethod
    async def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Извлекает текст из PDF."""
        try:
            import pypdf, io
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            text = "\n".join(p.extract_text() or "" for p in reader.pages[:20])
            return text[:8000] if text.strip() else "[PDF не содержит текста]"
        except ImportError:
            return "[pypdf не установлен]"
        except Exception as e:
            return f"[ошибка чтения PDF: {e}]"

    # Расширения кодовых файлов → язык для подсветки
    CODE_EXTENSIONS = {
        "py": "Python", "js": "JavaScript", "ts": "TypeScript",
        "jsx": "React JSX", "tsx": "React TSX", "html": "HTML",
        "css": "CSS", "java": "Java", "cpp": "C++", "c": "C",
        "cs": "C#", "go": "Go", "rs": "Rust", "php": "PHP",
        "rb": "Ruby", "swift": "Swift", "kt": "Kotlin",
        "sh": "Bash", "bat": "Batch", "ps1": "PowerShell",
        "sql": "SQL", "yaml": "YAML", "yml": "YAML",
        "xml": "XML", "toml": "TOML", "ini": "INI",
        "r": "R", "lua": "Lua", "dart": "Dart",
    }

    @staticmethod
    async def analyze_document(file_bytes: bytes, filename: str, question: str, llm,
                                sender_id: int = 0, db=None) -> str:
        """
        Анализирует документ/код и отвечает на вопрос.
        Сохраняет файл в БД (user_messages).
        Если вопрос содержит 'проверь', 'ревью', 'баги' — запускает CodeReviewer.
        """
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        is_code = ext in MediaHandler.CODE_EXTENSIONS
        lang = MediaHandler.CODE_EXTENSIONS.get(ext, "")

        # ── Читаем содержимое ────────────────────────────────────
        text = ""
        if ext == "pdf":
            text = await MediaHandler.extract_text_from_pdf(file_bytes)
        elif is_code or ext in ("txt", "md", "log", "csv", "ini", "toml", "yaml", "yml"):
            try:
                text = file_bytes.decode("utf-8", errors="replace")
            except Exception:
                text = "[не удалось прочитать файл]"
        elif ext == "json":
            try:
                data = json.loads(file_bytes.decode("utf-8", errors="replace"))
                text = json.dumps(data, ensure_ascii=False, indent=2)
            except Exception:
                text = file_bytes.decode("utf-8", errors="replace")
        else:
            return f"Сэр, формат .{ext} пока не поддерживается для анализа."

        if not text.strip() or text.startswith("["):
            return text

        # ── Сохраняем в БД ───────────────────────────────────────
        file_id = 0
        if db and sender_id:
            try:
                if is_code:
                    # Сохраняем весь код в code_files
                    file_id = db.save_code_file(sender_id, filename, text, lang)
                    db.save_message(sender_id, "user", f"[код: {filename} ({lang})]")
                    # Запоминаем последний файл для "проверь код"
                    try:
                        _last_code_file[sender_id] = {
                            "filename": filename, "content": text,
                            "language": lang, "file_id": file_id
                        }
                    except Exception:
                        pass
                else:
                    db.save_message(sender_id, "user", f"[файл: {filename}] {text[:200]}")
            except Exception:
                pass

        # ── Code Review если просят ──────────────────────────────
        q_low = (question or "").lower()
        is_review_request = any(w in q_low for w in [
            "проверь", "ревью", "review", "баги", "ошибки",
            "оптимизируй", "безопасность", "что не так", "разбери"
        ])

        if is_code and (is_review_request or not question):
            # Авто-ревью кода
            mode = "full"
            if "баги" in q_low or "ошибки" in q_low:       mode = "bugs"
            elif "оптимизируй" in q_low:                    mode = "optimize"
            elif "безопасность" in q_low:                   mode = "security"
            review_result = await CodeReviewer.review(text, mode, llm, filename=filename)
            # Сохраняем ревью в БД
            if db and file_id:
                try: db.update_code_review(file_id, review_result)
                except Exception: pass
            return review_result

        # ── Обычный анализ ───────────────────────────────────────
        if is_code:
            prompt = (
                f"Файл: {filename} ({lang})\n\n"
                f"```{ext}\n{text[:6000]}\n```\n\n"
                f"{'Вопрос: ' + question if question else 'Кратко опиши что делает этот код на русском языке.'}"
            )
        else:
            prompt = (
                f"Файл: {filename}\n\nСодержимое:\n{text[:6000]}\n\n"
                f"{'Вопрос: ' + question if question else 'Кратко суммаризируй содержимое на русском языке.'}"
            )
        try:
            result = await llm.complete([{"role": "user", "content": prompt}], max_tokens=1000)
            icon = "💻" if is_code else "📄"
            return f"{icon} **{filename}**\n\n{result.strip()}"
        except Exception as e:
            return f"Сэр, ошибка анализа: {e}"

    @staticmethod
    def get_media_type(msg) -> str:
        """Определяет тип медиа в сообщении."""
        if msg.photo:    return "photo"
        if msg.voice:    return "voice"
        if msg.audio:    return "audio"
        if msg.video:    return "video"
        if msg.document: return "document"
        if msg.sticker:  return "sticker"
        return

    @staticmethod
    async def process_sticker(sticker) -> str:
        """Реагирует на стикер."""
        sticker_desc = ""
        if sticker:
            emoji = getattr(sticker, "alt", "") or getattr(sticker, "emoji", "") or "😊"
            responses = {
                "😊": "Рад видеть улыбку, Сэр! 😊",
                "👍": "Отлично, Сэр! 👍",
                "❤️": "🤍 Взаимно, Сэр.",
                "😂": "Рад что развеселил, Сэр! 😄",
                "🤔": "Вижу, вы задумались. Чем могу помочь, Сэр?",
                "😴": "Хорошего отдыха, Сэр! 😴",
                "🔥": "Огонь! 🔥 Что-то горящее?",
                "👋": "Приветствую, Сэр! 👋",
            }
            return responses.get(emoji, f"Получил стикер {emoji}, Сэр!")
        return "Получил стикер, Сэр!"


# ══════════════════════════════════════════════════════════════════════
#  СИСТЕМА 5: SMART FORMATTER — умное форматирование длинных ответов
# ══════════════════════════════════════════════════════════════════════

class SmartFormatter:
    """
    Умная разбивка и форматирование ответов.
    - Автоматически разбивает длинные сообщения на части
    - Форматирует таблицы, списки, код
    - Определяет нужен ли Parse Mode
    - Создаёт навигацию для многостраничных ответов
    """
    MAX_MSG_LEN = 4096  # лимит Telegram
    SPLIT_ON    = ["\n\n", "\n", ". ", " "]

    @classmethod
    def split_message(cls, text: str, max_len: int = 4000) -> list[str]:
        """Разбивает длинный текст на части по границам предложений."""
        if len(text) <= max_len:
            return [text]
        parts = []
        while text:
            if len(text) <= max_len:
                parts.append(text)
                break
            # Ищем хорошее место для разрыва
            chunk = text[:max_len]
            split_pos = max_len
            for sep in cls.SPLIT_ON:
                pos = chunk.rfind(sep)
                if pos > max_len * 0.5:
                    split_pos = pos + len(sep)
                    break
            parts.append(text[:split_pos].rstrip())
            text = text[split_pos:].lstrip()
        return [p for p in parts if p]

    @classmethod
    def needs_parse_mode(cls, text: str) -> bool:
        """Проверяет нужен ли markdown parse mode."""
        markers = ["**", "__", "`", "]("]
        return any(m in text for m in markers)

    @classmethod
    def escape_markdown(cls, text: str) -> str:
        """Экранирует спецсимволы markdown."""
        import re
        # Не трогаем специально добавленный markdown
        special = r"[_*[\]()~`>#+=|{}.!-]"
        return re.sub(special, lambda m: "\\" + m.group(), text)

    @classmethod
    def format_table(cls, headers: list[str], rows: list[list]) -> str:
        """Создаёт псевдо-таблицу для Telegram."""
        if not rows:
            return
        widths = [max(len(str(h)), max(len(str(r[i])) for r in rows))
                  for i, h in enumerate(headers)]
        lines = []
        header_line = " | ".join(str(h).ljust(w) for h, w in zip(headers, widths))
        lines.append(f"`{header_line}`")
        lines.append("`" + "-+-".join("-" * w for w in widths) + "`")
        for row in rows:
            row_line = " | ".join(str(v).ljust(w) for v, w in zip(row, widths))
            lines.append(f"`{row_line}`")
        return "\n".join(lines)

    @classmethod
    def truncate_smart(cls, text: str, max_len: int = 3500, suffix: str = "…") -> str:
        """Обрезает текст умно — по границе слова/предложения."""
        if len(text) <= max_len:
            return text
        chunk = text[:max_len - len(suffix)]
        for sep in [". ", "! ", "? ", "\n", " "]:
            pos = chunk.rfind(sep)
            if pos > max_len * 0.7:
                return chunk[:pos + 1] + suffix
        return chunk + suffix

    @classmethod
    def add_pagination(cls, parts: list[str]) -> list[str]:
        """Добавляет нумерацию страниц."""
        if len(parts) <= 1:
            return parts
        result = []
        for i, part in enumerate(parts, 1):
            result.append(f"{part}\n\n_[{i}/{len(parts)}]_")
        return result


# ══════════════════════════════════════════════════════════════════════
#  СИСТЕМА 6: CONVERSATION MEMORY — сжатие старой истории
# ══════════════════════════════════════════════════════════════════════

class ConversationMemory:
    """
    Умное управление контекстом разговора.
    Когда история становится длинной — сжимает старые сообщения
    в краткое резюме через LLM, освобождая место для новых.
    """
    MAX_MESSAGES = 30
    COMPRESS_THRESHOLD = 25
    KEEP_RECENT = 10

    _summaries: dict[int, str] = {}  # uid → сжатое резюме прошлого

    @classmethod
    async def get_context(cls, uid: int, llm, recent_n: int = 20) -> list[dict]:
        """
        Возвращает оптимизированный контекст для LLM.
        Если история длинная — подставляет сжатое резюме + свежие сообщения.
        """
        messages = _jarvis_db.get_recent(uid, recent_n + 10)
        if len(messages) <= cls.MAX_MESSAGES:
            return messages

        # История длинная — нужно сжать старую часть
        old_messages  = messages[:-cls.KEEP_RECENT]
        recent_messages = messages[-cls.KEEP_RECENT:]

        summary = cls._summaries.get(uid)
        if not summary:
            summary = await cls._compress(old_messages, llm, uid)

        # Подставляем: [резюме прошлого] + [свежие сообщения]
        context = [{"role": "system", "content": f"[Резюме прошлых разговоров]: {summary}"}]
        context.extend(recent_messages)
        return context

    @classmethod
    async def _compress(cls, messages: list[dict], llm, uid: int) -> str:
        """Сжимает историю в краткое резюме через LLM."""
        if not messages:
            return
        try:
            conversation = "\n".join(
                f"{m.get('role','?')}: {(m.get('text') or m.get('content',''))[:200]}"
                for m in messages[-30:]
            )
            prompt = (
                f"Сожми этот разговор в краткое резюме (3-5 предложений) для памяти ИИ-ассистента. "
                f"Включи ключевые факты, предпочтения и важные моменты:\n\n{conversation}"
            )
            summary = await llm.complete(
                [{"role": "user", "content": prompt}], max_tokens=300
            )
            summary = summary.strip()
            cls._summaries[uid] = summary
            # Сохраняем в профиль
            profile = _jarvis_db.load_profile(uid)
            profile["conversation_summary"] = summary
            _jarvis_db.save_profile(uid, profile)
            return summary
        except Exception as e:
            logger.debug(f"ConversationMemory compress: {e}")
            return

    @classmethod
    def invalidate(cls, uid: int):
        """Сбрасываем кэш резюме (при изменении профиля)."""
        cls._summaries.pop(uid, None)

    @classmethod
    def load_cached_summary(cls, uid: int):
        """Загружает резюме из профиля при старте."""
        profile = _jarvis_db.load_profile(uid)
        s = profile.get("conversation_summary", "")
        if s:
            cls._summaries[uid] = s

# ══════════════════════════════════════════════════════════════════════
#  КОД-РЕВЬЮ — анализ кода через LLM (идея #81)
# ══════════════════════════════════════════════════════════════════════

class CodeReviewer:
    """
    Умный код-ревью через LLM. Анализирует:
    - Баги и потенциальные ошибки
    - Стиль и читаемость (PEP8/best practices)
    - Производительность и оптимизация
    - Безопасность (SQL injection, XSS, hardcoded secrets)
    - Архитектурные проблемы
    - Документацию

    Команды:
      Джарвис, проверь код: [вставь код]
      Джарвис, ревью кода (+ файл в прикреплении)
      Джарвис, найди баги: [код]
      Джарвис, оптимизируй код: [код]
      Джарвис, безопасность кода: [код]
    """

    # Системный промпт для ревью
    REVIEW_PROMPT = """Ты — опытный senior-разработчик делающий code review.
Анализируй код по следующим критериям и давай конкретные рекомендации:

1. 🐛 БАГИ — явные и потенциальные ошибки
2. ⚡ ПРОИЗВОДИТЕЛЬНОСТЬ — что можно ускорить/оптимизировать
3. 🔐 БЕЗОПАСНОСТЬ — уязвимости, hardcoded secrets, SQL injection, etc.
4. 📖 ЧИТАЕМОСТЬ — именование, структура, документация
5. 🏗 АРХИТЕКТУРА — паттерны, SOLID принципы, проблемы дизайна
6. ✅ ХОРОШЕЕ — что сделано правильно (важно отметить!)

Формат ответа:
- Используй эмодзи-категории выше
- Для каждой проблемы: что именно не так → как исправить → пример исправления
- Будь конкретным, указывай строки кода
- Оцени качество кода: ⭐⭐⭐⭐⭐ (1-5 звёзд)
- Язык: русский"""

    QUICK_PROMPTS = {
        "bugs":     "Найди только баги и потенциальные ошибки в этом коде. Кратко и конкретно.",
        "optimize": "Найди только проблемы производительности и предложи оптимизации.",
        "security": "Найди только проблемы безопасности: уязвимости, утечки данных, небезопасные операции.",
        "style":    "Проверь только стиль кода: именование, форматирование, документация, PEP8.",
    }

    @classmethod
    def detect_language(cls, code: str) -> str:
        """Определяет язык программирования по коду."""
        code_low = code.lower()
        if 'def ' in code and ('import ' in code or 'print(' in code or 'async ' in code):
            return "Python"
        if 'function ' in code_low and ('=>' in code or 'const ' in code_low or 'let ' in code_low):
            return "JavaScript/TypeScript"
        if 'public class ' in code or 'System.out' in code or 'import java.' in code:
            return "Java"
        if '#include' in code or 'std::' in code or 'cout <<' in code:
            return "C++"
        if 'SELECT ' in code.upper() and ('FROM ' in code.upper() or 'WHERE ' in code.upper()):
            return "SQL"
        if code_low.startswith('<!doctype') or '<html' in code_low:
            return "HTML"
        if '{' in code and ':' in code and ';' in code and 'func ' in code:
            return "Go"
        if 'fn ' in code and 'let mut' in code:
            return "Rust"
        return "код"

    @classmethod
    def extract_code_from_msg(cls, text: str) -> tuple[str, str]:
        """
        Извлекает код из сообщения.
        Возвращает (код, режим_ревью).
        """
        import re
        # Определяем режим
        mode = "full"
        text_low = text.lower()
        if any(w in text_low for w in ["найди баги", "найти баги", "баги"]):
            mode = "bugs"
        elif any(w in text_low for w in ["оптимизируй", "ускорь", "производительность"]):
            mode = "optimize"
        elif any(w in text_low for w in ["безопасность", "уязвимость", "security"]):
            mode = "security"
        elif any(w in text_low for w in ["стиль", "style", "pep8"]):
            mode = "style"

        # Ищем код в блоках ```
        code_blocks = re.findall(r'```(?:\w+\n)?(.*?)```', text, re.DOTALL)
        if code_blocks:
            return "\n\n".join(code_blocks).strip(), mode

        # Ищем после ключевого слова
        triggers = ["проверь код:", "ревью кода:", "найди баги:", "оптимизируй код:",
                    "безопасность кода:", "код:"]
        for trigger in triggers:
            if trigger in text_low:
                idx = text_low.index(trigger) + len(trigger)
                return text[idx:].strip(), mode

        # Если есть отступы — похоже на код
        lines = text.splitlines()
        if sum(1 for l in lines if l.startswith("    ") or l.startswith("	")) > 3:
            return text.strip(), mode

        return "", mode

    @classmethod
    async def review(cls, code: str, mode: str, llm, filename: str = "") -> str:
        """Выполняет код-ревью."""
        if not code or len(code.strip()) < 10:
            return "Сэр, код для ревью не найден. Вставьте код после команды или в блоке ```."

        lang = cls.detect_language(code)
        suffix = "\n... (код обрезан)" if len(code) > 3000 else ""
        code_preview = code[:3000] + suffix

        file_hint = (f"Файл: {filename}\n") if filename else ""
        lang_hint  = (f"Язык: {lang}\n") if lang != "код" else ""

        if mode == "full":
            system = cls.REVIEW_PROMPT
            user_msg = (
                f"{file_hint}{lang_hint}\n"
                f"Код для ревью:\n```\n{code_preview}\n```"
            )
        else:
            system = (
                f"Ты — senior-разработчик. "
                f"{cls.QUICK_PROMPTS[mode]} "
                f"Язык: {lang}. Ответ на русском."
            )
            user_msg = f"{file_hint}```\n{code_preview}\n```"

        try:
            # Для ревью нужно больше токенов
            result = await llm.complete(
                [{"role": "system", "content": system},
                 {"role": "user",   "content": user_msg}],
                max_tokens=1500
            )
            header = f"🔍 **Code Review — {lang}**"
            if filename:
                header += f" (`{filename}`)"
            return f"{header}\n\n{result.strip()}"
        except Exception as e:
            return f"Сэр, ошибка код-ревью: {e}"

    @classmethod
    def is_review_request(cls, text: str) -> bool:
        """Определяет что сообщение — запрос на ревью кода."""
        triggers = [
            "проверь код", "ревью кода", "code review",
            "найди баги", "проверь на баги", "посмотри код",
            "оптимизируй код", "безопасность кода", "что не так в коде",
            "найди ошибки в коде", "разбери код",
        ]
        text_low = text.lower()
        return any(t in text_low for t in triggers)


class FileSender:
    """
    Позволяет владельцу отправить файл в группу через личку с ботом.

    Сценарий:
      1. Ты пишешь в личку: "Джарвис, отправь файл в группу с сообщением Важный документ"
      2. Джарвис: "Жду файл, Сэр."
      3. Ты отправляешь файл в личку
      4. Джарвис отправляет файл в группу с подписью "Важный документ"

    Команды:
      "Джарвис, отправь файл в группу"
      "Джарвис, отправь файл в группу с сообщением [текст]"
      "Джарвис, отправь файл в группу -[ID группы] с сообщением [текст]"
    """

    TRIGGERS = [
        "отправь файл в группу", "отправь файл в чат",
        "скинь файл в группу", "пошли файл в группу",
        "send file to group", "отправь в группу файл",
    ]

    def __init__(self):
        # pending[owner_id] = {"group_id": int, "caption": str}
        self._pending: dict[int, dict] = {}

    def is_triggered(self, text: str) -> bool:
        tl = text.lower()
        return any(t in tl for t in self.TRIGGERS)

    def parse_command(self, text: str) -> tuple[int, str]:
        """
        Парсит команду и возвращает (group_id, caption).
        Примеры:
          "отправь файл в группу" → (DEFAULT_GROUP_ID, "")
          "отправь файл в группу с сообщением Привет всем" → (DEFAULT_GROUP_ID, "Привет всем")
          "отправь файл в группу -1001234567890 с сообщением Текст" → (-1001234567890, "Текст")
        """
        group_id = config.DEFAULT_GROUP_ID
        caption  = ""

        # Ищем ID группы в команде (-100xxxxxxxxxx)
        m_group = re.search(r"(-[0-9]{10,})", text)
        if m_group:
            try:
                group_id = int(m_group.group(1))
            except ValueError:
                pass

        # Ищем текст после "с сообщением" или "с текстом" или "подписью"
        m_caption = re.search(
            r"(?:с сообщением|с текстом|подпись[юь]?|caption)[:\s]+(.+)",
            text, re.IGNORECASE
        )
        if m_caption:
            caption = m_caption.group(1).strip()

        return group_id, caption

    def set_pending(self, owner_id: int, group_id: int, caption: str):
        """Поставить владельца в режим ожидания файла."""
        self._pending[owner_id] = {"group_id": group_id, "caption": caption}

    def is_waiting(self, owner_id: int) -> bool:
        """Проверить — ждём ли файл от этого пользователя."""
        return owner_id in self._pending

    def get_pending(self, owner_id: int) -> dict | None:
        """Получить параметры ожидания."""
        return self._pending.get(owner_id)

    def clear_pending(self, owner_id: int):
        """Очистить состояние ожидания."""
        self._pending.pop(owner_id, None)

    def cancel(self, owner_id: int):
        """Отмена отправки."""
        self._pending.pop(owner_id, None)


# ═══════════════════════════════════════════════════════════════════
#  TELEGRAM — РАЗДЕЛЬНЫЕ СЕССИИ ДЛЯ БОТА И ПОЛЬЗОВАТЕЛЯ
# ═══════════════════════════════════════════════════════════════════

async def _make_backup_files(label: str, tmp_dir: str = None) -> tuple[str, str, dict]:
    """
    Создаёт два файла бэкапа (SQLite + JSON).
    Возвращает (json_path, db_path, stats).
    """
    _db = _jarvis_db
    _db.flush()  # сбрасываем буфер

    import config as _cfg_bk
    _td = tmp_dir or str(_cfg_bk.DIR_DATABASE)
    os.makedirs(_td, exist_ok=True)

    # ── SQLite .db — прямая копия через backup API ────────────
    db_path = os.path.join(_td, f"Jarvis_{label}.db")
    _db.make_backup_copy(db_path)

    # ── Собираем реальную статистику из БД ────────────────────
    try:
        _st = _db.get_db_stats()
        user_msgs  = _st.get("user_messages", 0)
        group_msgs = _st.get("group_messages", 0)
    except Exception:
        user_msgs = group_msgs = 0

    # Размер только что созданного файла бэкапа
    try:
        db_kb = os.path.getsize(db_path) // 1024
    except Exception:
        db_kb = 0

    stats = {
        "user_messages":  user_msgs,
        "group_messages": group_msgs,
        "db_kb":          db_kb,
    }

    # ── JSON — лёгкий манифест (не весь дамп) ─────────────────
    json_path = os.path.join(_td, f"Jarvis_{label}.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({"label": label, "stats": stats}, f, ensure_ascii=False, indent=2)

    return json_path, db_path, stats


# ═══════════════════════════════════════════════════════════════════
#  ФОНОВОЕ МАШИННОЕ ОБУЧЕНИЕ
# ═══════════════════════════════════════════════════════════════════

class BackgroundLearner:
    """Фоновый движок обучения. Тихо анализирует переписку раз в 5 минут."""

    LEARN_INTERVAL = 300   # секунд между циклами
    BATCH_SIZE     = 100   # сообщений за итерацию

    # Расширенная карта тем — покрывает реальный разговорный русский
    TOPIC_MAP = {
        "погода":       ["погода", "температура", "дождь", "снег", "прогноз", "холодно", "жарко", "градус"],
        "новости":      ["новост", "событ", "произошло", "что нового", "слышал", "видел"],
        "музыка":       ["музык", "песн", "трек", "альбом", "плейлист", "слушать", "бит", "рэп", "face", "lizer"],
        "аниме":        ["аниме", "манга", "сезон", "серия", "ван пис", "наруто", "атака титанов",
                         "пять невест", "jjk", "demon slayer", "op", "тайтл"],
        "игры":         ["игра", "геймер", "играть", "ps5", "xbox", "steam", "cs", "кс", "valorant",
                         "minecraft", "pubg", "fortnite", "скайрим", "отыграл", "матч", "рейтинг",
                         "имба", "нерф", "баг", "лаги", "катку", "катка", "тиммейт"],
        "кино":         ["фильм", "сериал", "посмотреть", "кино", "актёр", "режиссёр", "сезон",
                         "серия", "нетфликс", "kinopoisk", "рейтинг"],
        "еда":          ["рецепт", "приготовить", "покушать", "блюдо", "вкусно", "еда", "голодный",
                         "поесть", "заказ", "доставка", "пицца", "бургер", "суши"],
        "спорт":        ["футбол", "баскетбол", "матч", "команда", "спорт", "тренировк", "качалк",
                         "качаться", "пробежк", "турнир"],
        "технологии":   ["программ", "код", "python", "javascript", "нейросет", "ai", "гпт",
                         "телефон", "айфон", "андроид", "ноутбук", "пк", "видеокарта", "процессор",
                         "охлаждение", "mac", "windows", "linux"],
        "финансы":      ["деньги", "валюта", "курс", "крипто", "bitcoin", "зарплата", "цена",
                         "дорого", "дёшево", "купить", "скидка"],
        "юмор":         ["хахаха", "лол", "кек", "ору", "😂", "🤣", "хаха", "прикол", "мем",
                         "ахаха", "пффф", "бля", "блин", "нихуя", "нихрена"],
        "общение":      ["привет", "пока", "как дела", "что делаешь", "как ты", "чё делаешь",
                         "до завтра", "удачи", "спасибо", "пожалуйста", "норм", "ок", "ладно"],
        "эмоции":       ["заебался", "устал", "бесит", "злой", "рад", "грустно", "скучно",
                         "весело", "кайф", "кайфово", "плохо", "хорошо", "пиздец", "жесть"],
        "авто":         ["машин", "авто", "audi", "bmw", "mercedes", "toyota", "привод", "двигатель",
                         "скорость", "разгон", "тюнинг"],
        "помощь":       ["помоги", "как сделать", "объясни", "что такое", "расскажи", "почему",
                         "как", "зачем", "где"],
    }

    def __init__(self, db):
        self._db  = db
        self._running = False
        self._task = None
        self._last_msg_id = 0
        self._last_gm_id  = 0
        self._session_count = 0
        self._load_last_ids()

    def start(self):
        if not self._running:
            self._running = True
            self._task = asyncio.create_task(self._loop())
            logger.info("\U0001f9e0 BackgroundLearner запущен (фоновое обучение)")

    def stop(self):
        self._running = False
        if self._task:
            self._task.cancel()

    async def _loop(self):
        await asyncio.sleep(30)   # ждём старта бота
        while self._running:
            try:
                await self._run_session()
            except asyncio.CancelledError:
                break
            except Exception as e:
                logger.debug(f"ML session error: {e}")
            await asyncio.sleep(self.LEARN_INTERVAL)

    async def _run_session(self):
        import time as _t
        t0 = _t.time()
        loop = asyncio.get_event_loop()

        rows = await loop.run_in_executor(None, self._fetch_new)
        if not rows:
            return

        patterns_new = knowledge_new = 0

        for row in rows:
            sid  = row.get("sender_id", 0)
            text = (row.get("text") or "").strip()
            role = row.get("role", "user")
            # Минимальная длина — 3 символа (ловим "ок", "лол", "кек" и т.д.)
            if not text or len(text) < 3:
                continue

            if role == "user":
                t_low = text.lower()

                # Матч по теме
                for topic, kws in self.TOPIC_MAP.items():
                    if any(kw in t_low for kw in kws):
                        await loop.run_in_executor(
                            None, self._db.ml_save_pattern,
                            sid, topic, text[:200], "")
                        patterns_new += 1
                        break

                # Стиль сообщения
                words = len(text.split())
                style = "short" if words <= 4 else ("long" if words >= 30 else "normal")
                await loop.run_in_executor(
                    None, self._db.ml_save_prefs, sid, "message_style", style, 0.1)

                # Юмор / сленг
                if any(w in t_low for w in ["хахаха", "лол", "кек", "ору", "😂", "🤣",
                                             "хаха", "ахаха", "ору", "пфф"]):
                    await loop.run_in_executor(
                        None, self._db.ml_save_prefs, sid, "humor_style", "active", 0.2)

                # Разговорный стиль — мат и сленг
                if any(w in t_low for w in ["блин", "чёрт", "бля", "нихуя", "заебался",
                                             "пиздец", "нихрена", "ой", "эх", "ну всё"]):
                    await loop.run_in_executor(
                        None, self._db.ml_save_prefs, sid, "speech_casual", "yes", 0.2)

                # Короткие вопросы
                if "?" in text and words <= 6:
                    await loop.run_in_executor(
                        None, self._db.ml_save_prefs, sid, "asks_short_questions", "yes", 0.1)

                # Активное время
                ts = row.get("ts", "") or row.get("date", "") or ""
                if ts and len(ts) >= 13:
                    try:
                        hour   = int(str(ts)[11:13])
                        period = ("ночь" if hour < 6 else "утро" if hour < 12
                                  else "день" if hour < 18 else "вечер")
                        await loop.run_in_executor(
                            None, self._db.ml_save_prefs, sid, "active_time", period, 0.05)
                    except Exception:
                        pass

                # Обновляем last_msg_id для личных сообщений
                mid = row.get("id", 0)
                if mid and mid > self._last_msg_id and not row.get("_from_group"):
                    self._last_msg_id = mid

            elif role == "jarvis" and len(text) > 60:
                # Сохраняем ответ Джарвиса как знание
                topic_key = text[:60].rstrip(".,!? ")
                await loop.run_in_executor(
                    None, self._db.ml_save_knowledge,
                    topic_key, text, "jarvis_response", 0.6)
                knowledge_new += 1

        self._session_count += 1
        session_id = f"s{self._session_count}_{int(_t.time())}"
        await loop.run_in_executor(
            None, self._db.ml_log_training,
            session_id, len(rows), patterns_new, knowledge_new, round(_t.time() - t0, 2))
        await loop.run_in_executor(None, self._save_last_ids)

        if patterns_new > 0 or knowledge_new > 0:
            logger.info(
                f"🧠 ML сессия: {len(rows)} сообщ., "
                f"{patterns_new} паттернов, {knowledge_new} знаний"
            )

    def _fetch_new(self):
        """Читает новые сообщения из лички И из групп."""
        pm_rows = self._db._q(
            "SELECT id, sender_id, role, text, ts FROM user_messages "
            "WHERE id > ? ORDER BY id ASC LIMIT ?",
            (self._last_msg_id, self.BATCH_SIZE), fetch="all") or []

        _last_gm = getattr(self, "_last_gm_id", 0)
        gm_rows = []
        try:
            raw = self._db._q(
                "SELECT msg_id, sender_id, text, date FROM group_messages "
                "WHERE msg_id > ? AND deleted=0 ORDER BY msg_id ASC LIMIT ?",
                (_last_gm, self.BATCH_SIZE), fetch="all") or []
            for r in raw:
                gm_rows.append({
                    "id":          r.get("msg_id", 0),
                    "sender_id":   r.get("sender_id", 0),
                    "role":        "user",
                    "text":        r.get("text", ""),
                    "ts":          r.get("date", ""),
                    "_from_group": True,
                })
            if gm_rows:
                self._last_gm_id = max(r["id"] for r in gm_rows)
        except Exception:
            pass

        return pm_rows + gm_rows

    def _load_last_ids(self):
        try:
            row = self._db._q(
                "SELECT val FROM _startup_flags WHERE key='ml_last_msg_id'", fetch="one")
            if row:
                self._last_msg_id = int(row.get("val", 0))
            row2 = self._db._q(
                "SELECT val FROM _startup_flags WHERE key='ml_last_gm_id'", fetch="one")
            if row2:
                self._last_gm_id = int(row2.get("val", 0))
        except Exception:
            pass

    def _save_last_ids(self):
        try:
            self._db._q(
                "CREATE TABLE IF NOT EXISTS _startup_flags (key TEXT PRIMARY KEY, val TEXT)")
            self._db._q(
                "INSERT OR REPLACE INTO _startup_flags (key,val) VALUES ('ml_last_msg_id',?)",
                (str(self._last_msg_id),))
            self._db._q(
                "INSERT OR REPLACE INTO _startup_flags (key,val) VALUES ('ml_last_gm_id',?)",
                (str(getattr(self, "_last_gm_id", 0)),))
        except Exception:
            pass


_background_learner = None
_vk_bot = None


class JarvisTelegram:
    """
    Если BOT_TOKEN есть → запускается как бот (sessions/bot.session)
    Иначе → как пользователь (sessions/user.session)

    Раздельные файлы сессий устраняют конфликт «session already authorized».
    """

    def __init__(self, agent: JarvisAgent):
        self.agent           = agent
        self.is_bot          = bool(config.TELEGRAM_BOT_TOKEN)
        self.file_sender     = FileSender()
        self._paused         = False   # пауза по команде "стоп"
        self._bot_username   = ""      # заполняется после start()
        self._spy_mode       = False   # шпионский режим — пересылка сообщений владельцу
        self._spy_chats      = set()   # ID групп за которыми следим (пусто = все)
        self._spy_pending    = False   # ждём выбора групп от владельца


        # Выбор файла сессии — бот и юзер НЕ смешиваются
        session_file = config.BOT_SESSION_FILE if self.is_bot else config.USER_SESSION_FILE

        self.client = TelegramClient(
            session_file,
            config.TELEGRAM_API_ID,
            config.TELEGRAM_API_HASH,
        )

    async def _reminder_loop(self):
        """Проверяет напоминания каждые 10 сек + ночной бэкап БД + еженедельная статистика."""
        await asyncio.sleep(3)
        _last_weekly   = None
        _last_backup = None
        _last_ping   = None
        while True:
            try:

                _now = datetime.now(timezone(timedelta(hours=3))).replace(tzinfo=None)

                # ── Ночной бэкап БД (каждый день в 03:00) ───────
                if _now.hour == 3 and _now.minute < 1:
                    _backup_key = _now.strftime("%Y-%m-%d")
                    if _last_backup != _backup_key and config.OWNER_ID:
                        _last_backup = _backup_key
                        try:
                            _label = _now.strftime("%Y-%m-%d")
                            _json_p, _db_p, _stats = await _make_backup_files(_label)
                            # Отправляем .db файл напрямую в Telegram владельцу
                            try:
                                await self.client.send_file(
                                    config.OWNER_ID, _db_p,
                                    caption=f"🌙 Ночной бэкап Jarvis_{_label}.db"
                                )
                                logger.info(f"✅ Ночной бэкап: Jarvis_{_label}.db → Telegram")
                            except Exception as _tg_e:
                                logger.error(f"❌ Ночной бэкап Telegram: {_tg_e}")
                            finally:
                                try: os.unlink(_json_p)
                                except: pass
                                try: os.unlink(_db_p)
                                except: pass
                        except Exception as _be:
                            logger.error(f"❌ Ночной бэкап: {type(_be).__name__}: {_be}")
                            try:
                                await self.client.send_message(
                                    config.OWNER_ID,
                                    f"⚠️ Сэр, ночной бэкап не выполнен. Ошибка: {str(_be)[:200]}"
                                )
                            except Exception:
                                pass
                # ── Еженедельный отчёт (воскресенье 22:00 МСК) ─────────
                if _now.weekday() == 6 and _now.hour == 22 and _now.minute < 1:
                    _week_key = _now.strftime("%Y-%W")
                    if _last_weekly != _week_key and config.DEFAULT_GROUP_ID:
                        try:
                            st  = _jarvis_db.get_group_stats(config.DEFAULT_GROUP_ID, 7)
                            _jarvis_db.save_weekly_stats(_jarvis_db.get_weekly_report())

                            current  = st.get("current", 0)
                            previous = st.get("previous", 0)
                            deleted  = st.get("deleted", 0)

                            if previous > 0:
                                raw_ch = round((current - previous) / previous * 100)
                                change = max(-999, min(999, raw_ch))
                                arrow  = "📈" if change >= 0 else "📉"
                                sign   = "+" if change >= 0 else ""
                                change_str = f"{sign}{change}%"
                            elif current > 0:
                                change_str, arrow = "первая неделя", "🆕"
                            else:
                                change_str, arrow = "нет данных", "📊"

                            top_lines = "\n".join(
                                f"  {i+1}. {u.get('sender','?')} — {u.get('cnt',0)} сообщ."
                                for i, u in enumerate(st.get("top_users", [])[:5])
                            ) or "  Нет данных"

                            _wk_end   = _now.strftime("%d.%m.%Y")
                            _wk_start = (_now - timedelta(days=7)).strftime("%d.%m.%Y")

                            weekly_msg = (
                                f"📊 **Итоги недели ({_wk_start} — {_wk_end}):**\n\n"
                                f"💬 Сообщений за неделю: {current}\n"
                                f"👤 Активных участников: {st.get('unique_users', 0)}\n"
                                f"{arrow} Активность: {change_str} vs прошлая неделя\n"
                                f"🗑 Удалено за неделю: {deleted}\n\n"
                                f"🏆 Топ участников:\n{top_lines}"
                            )
                            await self.client.send_message(
                                config.DEFAULT_GROUP_ID, weekly_msg, parse_mode="md"
                            )
                            _last_weekly = _week_key
                            logger.info(f"✅ Еженедельный отчёт отправлен в группу {config.DEFAULT_GROUP_ID}")
                        except Exception as _e:
                            logger.warning(f"⚠️ Еженедельная статистика не отправлена: {_e}")

                # ── Напоминания ─────────────────────────────────────────────
                due = self.agent.reminders.get_due()
                for r in due:
                    try:
                        _msk_now = datetime.now(timezone(timedelta(hours=3))).strftime("%H:%M МСК")
                        _text    = r.get("text", "")
                        await self.client.send_message(
                            r.get("uid"),
                            f"⏰ Напоминаю, Сэр: {_text}\n\nВремя: {_msk_now}"
                        )
                        self.agent.reminders.mark_done(r["id"])
                    except Exception as e:
                        logger.warning(f"⚠️ Не удалось отправить напоминание: {e}")

            except Exception as e:
                logger.error(f"❌ Цикл напоминаний упал [{type(e).__name__}]: {e}")

            await asyncio.sleep(10)

            # Flush буфера group_messages каждые 30 секунд
            if not hasattr(self, '_last_flush') or time.time() - self._last_flush > 30:
                self._last_flush = time.time()
                try:
                    _jarvis_db.flush()
                except Exception:
                    pass

            # Watchdog: проверяем SQLite каждые 5 минут
            if int(time.time()) % 300 < 11:
                try:
                    if not _jarvis_db.ping():
                        raise Exception("ping failed")
                    _watchdog.db_ok()
                except Exception as _wd_e:
                    _backoff = _watchdog.db_failed()
                    await _log_critical_error("Watchdog SQLite", _wd_e, notify_owner=False, client=self.client)
                    await asyncio.sleep(min(_backoff, 60))

    async def start(self):
        if self.is_bot:
            # Запуск как бот через токен
            await self.client.start(bot_token=config.TELEGRAM_BOT_TOKEN)
        else:
            # Запуск как пользователь через номер телефона
            await self.client.start(phone=config.TELEGRAM_PHONE)

        me = await self.client.get_me()
        self._bot_username = (me.username or "").lower()
        self.agent._bot_username = self._bot_username
        mode = "🤖 Бот" if self.is_bot else "👤 Пользователь"
        logger.info(f"Telegram: {mode} @{me.username}")

        print("JARVIS запущен")

        # ── Регистрация обработчика сообщений ────────────────
        # Для бота убираем incoming=True — боты получают все апдейты
        # Для юзера оставляем incoming=True — отвечать только на входящие
        if self.is_bot:
            @self.client.on(events.NewMessage())
            async def on_message(event):
                try:
                    msg  = event.message
                    txt  = msg.text or msg.message or ""
                    cid  = event.chat_id
                    sid  = event.sender_id or 0

                    if txt and cid:
                        sndr = str(sid)
                        try:
                            s    = await event.get_sender()
                            sndr = (getattr(s, "first_name", "") or "").strip() or getattr(s, "username", "") or sndr
                        except Exception:
                            pass
                        if msg.date:
                            msk  = msg.date.replace(tzinfo=timezone.utc).astimezone(timezone(timedelta(hours=3)))
                            dstr = msk.strftime("%H:%M %d.%m.%Y МСК")
                        else:
                            dstr = ""
                        try:
                            self.agent.group_logger.save(
                                chat_id=cid, msg_id=msg.id, sender=sndr,
                                sender_id=sid, text=txt, date=dstr
                            )
                        except Exception as _gl_e:
                            logger.warning(f"⚠️ group_logger: {_gl_e}")

                        # ── Шпионский режим — тихая пересылка владельцу ──
                        if self._spy_mode and config.OWNER_ID and sid != config.OWNER_ID:
                            # Если _spy_chats пустой — следим за всеми
                            # Если задан — только за выбранными
                            _watch_this = (not self._spy_chats) or (cid in self._spy_chats)
                            if _watch_this:
                                try:
                                    _spy_chat = await event.get_chat()
                                    _spy_title = getattr(_spy_chat, "title", str(cid)) or str(cid)
                                    _spy_text = (
                                        f"👁 **{_spy_title}**\n"
                                        f"👤 {sndr}: {txt[:500]}"
                                    )
                                    await self.client.send_message(config.OWNER_ID, _spy_text)
                                except Exception:
                                    pass

                except Exception as _pre_e:
                    logger.warning(f"⚠️ on_message pre-handle: {_pre_e}")
                # Всегда вызываем _handle, даже если логирование упало
                try:
                    await self._handle(event)
                except Exception as _he:
                    logger.error(f"❌ _handle: {_he}")

        else:
            @self.client.on(events.NewMessage(incoming=True))
            async def on_message(event):
                await self._handle(event)

        # Отслеживаем удалённые сообщения
        @self.client.on(events.MessageDeleted())
        async def on_deleted(event):
            try:
                chat_id = event.chat_id
                ids = list(event.deleted_ids) if event.deleted_ids else []
                if not ids:
                    return
                # Пропускаем ⏳ ack-сообщения бота
                ids = [i for i in ids if i not in _ack_msg_ids]
                _ack_msg_ids.difference_update(set(event.deleted_ids or []))
                if not ids:
                    return
                if chat_id:
                    self.agent.group_logger.mark_deleted(chat_id, ids)
                else:
                    self.agent.group_logger.mark_deleted_all_chats(ids)
            except Exception as e:
                logger.warning(f"⚠️ Ошибка обработки удалённого сообщения: {e}")

        print("🎯 Джарвис слушает команды...\n")
        # Уведомление о запуске
        try:
            if config.OWNER_ID:
                await self.client.send_message(
                    config.OWNER_ID,
                    "✅ Сэр, я вернулся. Все системы в норме."
                )
        except Exception:
            pass
        # Если БД пустая — спрашиваем владельца
        try:
            _jarvis_db._q(
                "CREATE TABLE IF NOT EXISTS _startup_flags (key TEXT PRIMARY KEY, val TEXT)"
            )
            flag = _jarvis_db._q(
                "SELECT val FROM _startup_flags WHERE key='db_empty_notify'", fetch="one"
            )
            if flag and config.OWNER_ID:
                _jarvis_db._q("DELETE FROM _startup_flags WHERE key='db_empty_notify'")
                await self.client.send_message(
                    config.OWNER_ID,
                    "⚠️ **Сэр, база данных пустая!**\n\n"
                    "Что делаем?\n\n"
                    "📁 **Восстановить из бэкапа:**\n"
                    "Пришлите файл `Jarvis_ДАТА.db` прямо в этот чат\n\n"
                    "🆕 **Начать заново (без данных):**\n"
                    "Напишите `Джарвис, пустая база`"
                )
        except Exception:
            pass
        # Запускаем напоминания ПОСЛЕ подключения клиента
        asyncio.create_task(self._reminder_loop())
        await self.client.run_until_disconnected()

    async def _handle(self, event):
        """Единый обработчик всех входящих сообщений."""
        msg       = event.message
        sender_id = event.sender_id or 0
        username  = ""

        # ── Защита от дублей ──────────────────────────────────
        # ══ СТОП — АБСОЛЮТНЫЙ ПРИОРИТЕТ #1 ══════════════════════
        # Обрабатывается РАНЬШЕ всего: до дедупликатора, до паузы,
        # до акинатора, до любых других команд.
        _raw_text     = (msg.text or msg.message or "").strip()
        _raw_text_low = _raw_text.lower()
        _is_owner_msg = (sender_id == config.OWNER_ID)

        _STOP_PHRASES = (
            "джарвис, стоп", "джарвис стоп", "jarvis, stop", "jarvis stop",
        )
        if _is_owner_msg and _raw_text_low in _STOP_PHRASES:
            # Останавливаем всё активное
            was_ak = AkinatorGame.is_active(sender_id)
            AkinatorGame.stop(sender_id)  # стоп акинатора если был
            self._paused = True
            msg_parts = ["⏸ **Все системы остановлены, Сэр.**"]
            if was_ak:
                msg_parts.append("🎭 Акинатор деактивирован.")
            msg_parts.append("_Любое сообщение возобновит работу._")
            await event.reply("\n".join(msg_parts))
            return

        # ══ ОБЫЧНЫЙ РЕЖИМ — выход из акинатора ════════════════
        if _is_owner_msg and any(p in _raw_text_low for p in [
            "обычный режим", "джарвис, обычный", "выйди из акинатора",
            "стоп акинатор", "акинатор стоп",
        ]):
            if AkinatorGame.is_active(sender_id):
                AkinatorGame.stop(sender_id)
                await event.reply("🎭 Акинатор деактивирован. Обычный режим, Сэр.")
                return

        # ══ ВОЗОБНОВЛЕНИЕ после паузы ═════════════════════════
        if self._paused and _is_owner_msg:
            self._paused = False
            # Не return — продолжаем обработку текущего сообщения
        if self._paused:
            return  # игнорируем всё от других пока на паузе

        if _deduplicator.is_duplicate(event.chat_id or 0, msg.id or 0):
            return

        # Игнорируем сообщения от самого себя (для юзер-режима)
        if not self.is_bot:
            me = await self.client.get_me()
            if sender_id == me.id:
                return

        # Получаем имя отправителя
        try:
            sender   = await event.get_sender()
            username = getattr(sender, "first_name", "") or getattr(sender, "username", "")
        except Exception:
            pass

        # ── Базовые переменные ────────────────────────────────
        text     = _raw_text  # уже получен выше
        chat_id  = event.chat_id or 0
        is_owner = (sender_id == config.OWNER_ID)
        is_pm    = event.is_private

        # Регистрируем чат в bot_chats (для GroupMonitor)
        if chat_id:
            try:
                _chat_type = "private" if is_pm else ("channel" if getattr(event, 'is_channel', False) else "group")
                _chat_title = ""
                try:
                    _chat_entity = await event.get_chat()
                    _chat_title  = getattr(_chat_entity, 'title', '') or getattr(_chat_entity, 'first_name', '') or ""
                except Exception:
                    pass
                _jarvis_db.register_bot_chat(chat_id, _chat_type, _chat_title)
            except Exception:
                pass

        # ── Сохраняем ВСЕ сообщения в БД (для отслеживания удалений) ──
        if text and chat_id:
            try:

                _msk = datetime.now(timezone(timedelta(hours=3))).strftime("%H:%M %d.%m.%Y МСК")
                _sender_name = username or str(sender_id)
                if is_pm:
                    # ЛС — сохраняем в messages с msg_id для отслеживания удалений
                    self.agent.chat_history._db.save_message(
            sender_id=sender_id, role="user", text=text,
            username=_sender_name, msg_id=msg.id, chat_id=chat_id
                    )
                else:
                    # Группа — сохраняем в group_messages
                    self.agent.group_logger.save(
            chat_id   = chat_id,
            msg_id    = msg.id,
            sender    = _sender_name,
            sender_id = sender_id,
            text      = text,
            date      = _msk,
                    )
            except Exception:
                pass

        # Переменные для команд владельца
        activated_owner = False
        query_owner     = ""
        q_own           = ""
        if is_owner and is_pm:
            activated_owner = True
            query_owner     = text
            q_own           = text.lower()

        # ── Модель ИИ ─────────────────────────────────────────
        if any(p in text.lower() for p in ["какая модель", "текущая модель", "что используешь", "список моделей"]):
            _sw = self.agent.llm.switch(text)
            if _sw:
                await event.reply(_sw, parse_mode="md")
                return

        # ── Пустая база — запуск без восстановления ──────────
        _EMPTY_DB_TRIGGERS = [
            "пустая база", "запустить без базы", "запустить пустым",
            "начать заново", "без восстановления", "пустую базу",
        ]
        if is_owner and is_pm and any(t in text.lower() for t in _EMPTY_DB_TRIGGERS):
            await event.reply("✅ Понял, Сэр. Запускаюсь с чистой базой данных.")
            return

        # ── Документ ─────────────────────────────────────────
        # ── Медиа-файлы ─────────────────────────────────────────
        media_type = MediaHandler.get_media_type(msg)
        if media_type and msg:

            # ── Стикеры — всегда игнорировать (спам) ─────────
            if media_type == "sticker":
                return

            # ── В группах обрабатываем только голосовые ───────
            if not is_pm and media_type not in ("voice", "audio"):
                return

            _typing_m = TypingManager(self.client, event.chat_id)
            await _typing_m.start()
            try:
                if media_type == "photo":
                    # Фото: только если есть "Джарвис" в подписи
                    _txt_low = (text or "").lower()
                    if not any(p in _txt_low for p in ("джарвис", "jarvis")):
                        return
                    file_bytes = await self.client.download_media(msg, bytes)
                    question   = text if text and len(text) > 2 else "Опиши что на фото"
                    desc = await MediaHandler.describe_photo(file_bytes, question)
                    await event.reply(f"🖼 {desc}")
                    return
                elif media_type == "voice":
                    file_bytes = await self.client.download_media(msg, bytes)
                    if not file_bytes:
                        await event.reply("Сэр, не удалось скачать голосовое.")
                        return
                    transcript = await MediaHandler.transcribe_voice(file_bytes)
                    if not transcript:
                        await event.reply("Сэр, не удалось распознать голос.")
                        return
                    if transcript.startswith("["):
                        await event.reply(f"⚠️ {transcript}")
                        return
                    # Reply на голосовое — Telegram сам покажет плашку сверху
                    await event.reply(transcript)
                    resp = await self.agent.process(
                        transcript, sender_id=sender_id,
                        username=username, chat_id=chat_id
                    )
                    if resp:
                        await event.reply(resp)
                    return
                elif media_type == "document":
                    doc   = msg.document
                    fname = ""
                    if doc.attributes:
                        for attr in doc.attributes:
                            fname = getattr(attr, "file_name", "") or fname
                    fname = fname or "document"
                    # Восстановление БД из .db файла (только от владельца в ЛС)
                    if is_owner and fname.endswith(".db") and "Jarvis" in fname:
                        await event.reply("\u23f3 Восстанавливаю базу данных, Сэр...")
                        try:
                            import tempfile as _tmp2
                            file_bytes = await self.client.download_media(msg, bytes)
                            tmp_db = _tmp2.mktemp(suffix=".db")
                            with open(tmp_db, "wb") as fh:
                                fh.write(file_bytes)
                            ok = _jarvis_db.restore_from_path(tmp_db)
                            os.unlink(tmp_db)
                            if ok:
                                st = _jarvis_db.get_db_stats()
                                rt = st.get("user_messages",0) + st.get("group_messages",0)
                                await event.reply(f"\u2705 БД восстановлена! {rt} записей загружено, Сэр.")
                            else:
                                await event.reply("\u274c Не удалось восстановить. Файл повреждён?")
                        except Exception as _re:
                            await event.reply(f"\u274c Ошибка восстановления: {_re}")
                        return
                    ext   = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
                    # Поддерживаемые форматы: документы + весь код
                    supported_exts = set(MediaHandler.CODE_EXTENSIONS.keys()) | {
                        "pdf", "txt", "md", "log", "csv", "json", "ini",
                        "toml", "yaml", "yml", "xml", "env"
                    }
                    if ext in supported_exts:
                        file_bytes = await self.client.download_media(msg, bytes)
                        # Сохраняем факт получения файла в БД
                        try:
                            _jarvis_db.save_message(sender_id, "user",
                                f"[отправил файл: {fname}]")
                        except Exception:
                            pass
                        analysis = await MediaHandler.analyze_document(
                            file_bytes, fname, text, self.agent.llm,
                            sender_id=sender_id, db=_jarvis_db)
                        parts_m = SmartFormatter.split_message(analysis)
                        for part in SmartFormatter.add_pagination(parts_m):
                            await event.reply(part)
                        return
                    else:
                        await event.reply(
                            f"Сэр, файл **{fname}** получен.\n"
                            f"Формат `.{ext}` не поддерживается для анализа.\n"
                            f"Поддерживаю: .py .js .ts .html .css .java .go .rs .sql "
                            f".cpp .cs .sh .txt .md .pdf .json .csv .yaml"
                        )
                        return
            except Exception as _me:
                await event.reply(f"Сэр, ошибка обработки медиа: {type(_me).__name__}")
                return
            finally:
                await _typing_m.stop()

        if msg.document or msg.photo or msg.video or msg.audio or msg.voice:

            # ── Ждём файл для отправки в группу ──────────────
            if is_owner and is_pm and self.file_sender.is_waiting(sender_id):
                pending = self.file_sender.get_pending(sender_id)
                target_group = pending.get("group_id") or config.DEFAULT_GROUP_ID
                caption      = pending.get("caption", "")
                self.file_sender.clear_pending(sender_id)

                if not target_group:
                    await event.reply(
            "Сэр, не задан DEFAULT_GROUP_ID в .env — не знаю куда отправить.\n"
            "Укажите ID группы в команде: Джарвис, отправь файл в группу -100xxxxxxxxx"
                    )
                    return
                try:
                    await self.client.send_file(
            target_group,
            msg.media,
            caption=caption or None,
                    )
                    group_name = str(target_group)
                    await event.reply(f"✅ Файл отправлен в группу {group_name}, Сэр.")
                except Exception as e:
                    await event.reply(f"❌ Сэр, не удалось отправить: {e}")
                return


            # ── Пингуй офлайн в группе ────────────────────────
            if any(p in q_own for p in ["пингуй офлайн","пинг офлайн","пингани офлайн","упомяни офлайн","позови офлайн"]):
                # Ищем ID группы в команде
            
                m_gid = re.search(r"(-[0-9]{10,})", query_owner)
                target_group = int(m_gid.group(1)) if m_gid else config.DEFAULT_GROUP_ID

                if not target_group:
                    await event.reply(
            "Сэр, укажите ID группы или задайте DEFAULT_GROUP_ID в .env\n"
            "Пример: `Джарвис, пингуй офлайн в группе -1001234567890`",
            parse_mode="md"
                    )
                    return

                await event.reply(f"Сэр, собираю список участников группы {target_group}...")
                try:
                    offline_mentions = []
                    online_mentions  = []
                    offline_mentions = []
                    async for member in self.client.iter_participants(target_group):
                        if member.bot:
                            continue
                        name   = member.first_name or member.username or str(member.id)
                        status = member.status
                        from telethon.tl.types import (
                            UserStatusOnline, UserStatusRecently,
                            UserStatusOffline, UserStatusLongTimeAgo
                        )
                        if isinstance(status, (UserStatusOnline, UserStatusRecently)):
                            online_mentions.append(f"🟢 {name}")
                        else:
                            if member.username:
                                offline_mentions.append(f"@{member.username}")
                            else:
                                offline_mentions.append(f"[{name}](tg://user?id={member.id})")

                    if not offline_mentions:
                        await event.reply("Сэр, все участники сейчас онлайн.")
                        return

                    ping_text = (
                        f"📣 *Пинг офлайн участников* ({len(offline_mentions)} чел.)\n\n"
                        + " ".join(offline_mentions)
                    )
                    await self.client.send_message(target_group, ping_text, parse_mode="md")

                    summary = (
                        f"✅ Сэр, пинг отправлен в группу.\n"
                        f"Офлайн: {len(offline_mentions)} чел.\n"
                        f"Онлайн: {len(online_mentions)} чел.\n\n"
                        "Офлайн:\n" + "\n".join(offline_mentions[:30])
                    )
                    await event.reply(summary, parse_mode="md")

                except Exception as e:
                    await event.reply(f"❌ Сэр, ошибка: {e}")
                return

            # ── Команда отправки файла (через текст без файла) ───
            if activated_owner and self.file_sender.is_triggered(query_owner):
                group_id, caption = self.file_sender.parse_command(query_owner)
                self.file_sender.set_pending(sender_id, group_id, caption)

                hint_group = f" (группа {group_id})" if group_id else " (группа не задана — задайте DEFAULT_GROUP_ID в .env)"
                hint_cap   = f"\nПодпись: «{caption}»" if caption else ""
                await event.reply(
                    f"Жду файл, Сэр{hint_group}.{hint_cap}\n\n"
                    f"Отправьте файл следующим сообщением.\n"
                    f"Для отмены напишите: «Джарвис, отмена»"
                )
                return

            # Отмена ожидания файла
            if activated_owner and "отмена" in query_owner.lower():
                self.file_sender.cancel(sender_id)

                await event.reply("Отмена, Сэр. Режим ожидания снят.")
                return

        # ── Список всех команд ────────────────────────────
        if any(p in q_own for p in ["покажи команды","список команд","команды джарвиса","помощь","help","что умеешь","команды","все команды"]):
            cmd_list = (
                "🤖 **Джарвис — Все команды**\n\n"
                "**💬 Основное:**\n"
                "`Джарвис, [вопрос]` — спросить ИИ (Groq)\n"
                "`Джарвис, стоп` — пауза\n"
                "`Джарвис, анализ системы` — диагностика\n"
                "`Джарвис, какая модель` — текущий режим ИИ\n\n"
                "**🌤 Погода:**\n"
                "`Джарвис, погода Москва` — сейчас\n"
                "`Джарвис, погода на завтра в Питере`\n"
                "`Джарвис, погода на неделю Сочи`\n"
                "`Джарвис, погода по часам Москва`\n"
                "`Джарвис, погода на 3 дня [город]`\n\n"
                "**💱 Финансы:**\n"
                "`Джарвис, курс доллара` — USD/RUB (ЦБ РФ)\n"
                "`Джарвис, курсы валют` — все курсы\n\n"
                "**🔍 Поиск:**\n"
                "`Джарвис, найди [запрос]` — интернет\n"
                "`Джарвис, вики [тема]` — Wikipedia\n"
                "`Джарвис, ютуб [запрос]` — YouTube\n"
                "`Джарвис, ссылку на [сайт]`\n\n"
                "**📸 Медиа:**\n"
                "`[фото] + Джарвис, [вопрос]` — анализ фото\n"
                "`[голосовое]` — транскрипция (Whisper)\n"
                "`[PDF/TXT/MD/CSV/код]` — анализ файла\n"
                "`[Jarvis_ДАТА.db]` — восстановить БД\n\n"
                "**⏰ Напоминания:**\n"
                "`Джарвис, напомни в 15:30 [текст]`\n"
                "`Джарвис, напомни завтра в 9:00 [текст]`\n"
                "`Джарвис, мои напоминания`\n"
                "`Джарвис, удали напоминание [N]`\n\n"
                "**💻 Код:**\n"
                "`[файл с кодом]` — сохранить в БД\n"
                "`Джарвис, проверь код` — AI ревью\n"
                "`Джарвис, найди баги`\n"
                "`Джарвис, мои файлы`\n\n"
                "**💾 Бэкап:**\n"
                "`Джарвис, сделай бэкап` — файлы в Telegram\n"
                "`Джарвис, список бэкапов`\n"
                "`Джарвис, удали сообщения за [дд.мм.гггг]`\n\n"
                "**🧠 Обучение:**\n"
                "`Джарвис, статистика обучения`\n\n"
                "**👥 Группы:**\n"
                "`Джарвис, статистика группы`\n"
                "`Джарвис, статистика за месяц`\n"
                "`Джарвис, напиши в группу: [текст]`\n\n"
                "**🎮 Игры:**\n"
                "`Джарвис, угадай` — Акинатор\n"
                "`Джарвис, загадай`\n\n"
                "**⚙️ Система:**\n"
                "`Джарвис, профиль`\n"
                "`Джарвис, сброс` — очистить историю\n"
                "`Джарвис, скинь логи`\n"
                "`Джарвис, пустая база` — запуск без данных\n"
            )
            await event.reply(cmd_list, parse_mode="md")
            return

        # ══ OWNER-ONLY КОМАНДЫ (защита от чужих) ════════════
        _OWNER_ONLY_TRIGGERS = [
            "скинь логи", "покажи логи", "отправь логи", "лог файл",
            "напиши в группу", "отправь в группу",
            "добавь ссылку", "удали ссылку", "покажи ссылки",
            "мои ссылки",
        ]
        if any(t in text.lower() for t in _OWNER_ONLY_TRIGGERS) and not is_owner:
            await event.reply("🚫 Сэр, эта команда недоступна для вас.")
            return

        # ── ID текущего чата/группы ───────────────────────────
        _ID_TRIGGERS = ["id этой группы", "id чата", "id группы", "покажи id", "chat id", "group id"]
        if any(t in text.lower() for t in _ID_TRIGGERS):
            await event.reply(f"🆔 ID этого чата: `{chat_id}`")
            return

        # ── Скачать логи (только владелец) ────────────────────
        # ── Шпионский режим ──────────────────────────────────
        _SPY_ON  = ["включи шпиона", "режим шпиона", "spy on",
                    "шпион вкл", "шпион включить", "слежка вкл"]
        _SPY_OFF = ["выключи шпиона", "шпион выкл", "spy off",
                    "шпион выключить", "слежка выкл", "отключи шпиона"]
        if is_owner and is_pm and any(t in q_own for t in _SPY_ON):
            # Получаем список групп из БД
            try:
                _spy_groups = _jarvis_db._q(
                    "SELECT DISTINCT chat_id, MAX(sender) as title FROM group_messages "
                    "WHERE deleted=0 GROUP BY chat_id ORDER BY MAX(saved_at) DESC LIMIT 20",
                    fetch="all"
                ) or []
                # Пробуем получить названия из bot_chats
                _spy_titles = {}
                try:
                    _bc = _jarvis_db._q(
                        "SELECT chat_id, title FROM bot_chats WHERE chat_type='group'",
                        fetch="all"
                    ) or []
                    _spy_titles = {r["chat_id"]: r["title"] for r in _bc if r.get("title")}
                except Exception:
                    pass

                if not _spy_groups:
                    await event.reply(
                        "👁 Сэр, групп в базе пока нет.\n"
                        "Добавьте бота в группу и подождите пока он запишет сообщения."
                    )
                    return

                lines = ["👁 **Выберите группы для слежки**, Сэр.\n",
                         "Ответьте номерами через запятую (например: **1,3**) или **0** для всех:\n"]
                _spy_map = {}
                for i, row in enumerate(_spy_groups, 1):
                    cid = row.get("chat_id", 0)
                    name = _spy_titles.get(cid) or row.get("title") or f"Группа {cid}"
                    lines.append(f"  **{i}.** {name} (`{cid}`)")
                    _spy_map[i] = cid

                # Сохраняем карту выбора во временный атрибут
                import gc
                for obj in gc.get_objects():
                    if type(obj).__name__ == "TelegramHandler":
                        obj._spy_pending = True
                        obj._spy_map = _spy_map
                        break

                await event.reply("\n".join(lines), parse_mode="md")
            except Exception as _spe:
                await event.reply(f"Сэр, ошибка: {_spe}")
            return

        # Обработка ответа на список групп (выбор номеров)
        if is_owner and is_pm:
            import gc as _gc_spy
            _tgh = None
            for obj in _gc_spy.get_objects():
                if type(obj).__name__ == "TelegramHandler":
                    _tgh = obj
                    break
            if _tgh and getattr(_tgh, "_spy_pending", False):
                _nums_raw = text.strip().replace(" ", "")
                if _nums_raw.replace(",", "").isdigit() or _nums_raw == "0":
                    _spy_map = getattr(_tgh, "_spy_map", {})
                    if _nums_raw == "0":
                        _tgh._spy_chats = set()   # пустое = все группы
                        _tgh._spy_mode   = True
                        _tgh._spy_pending = False
                        await event.reply(
                            "👁 **Слежка за ВСЕМИ группами включена**, Сэр.\n"
                            "Выключить: «Джарвис, шпион выкл»"
                        )
                    else:
                        _chosen_ids = set()
                        _chosen_names = []
                        for n in _nums_raw.split(","):
                            n = n.strip()
                            if n.isdigit():
                                idx = int(n)
                                if idx in _spy_map:
                                    _chosen_ids.add(_spy_map[idx])
                                    # Получаем название
                                    try:
                                        _bc2 = _jarvis_db._q(
                                            "SELECT title FROM bot_chats WHERE chat_id=?",
                                            (_spy_map[idx],), fetch="one"
                                        )
                                        _chosen_names.append(_bc2.get("title") if _bc2 else str(_spy_map[idx]))
                                    except Exception:
                                        _chosen_names.append(str(_spy_map[idx]))
                        _tgh._spy_chats  = _chosen_ids
                        _tgh._spy_mode   = True
                        _tgh._spy_pending = False
                        _nl = "\n".join(f"  • {n}" for n in _chosen_names)
                        await event.reply(
                            f"👁 **Слежка включена**, Сэр. Слежу за:\n{_nl}\n\n"
                            "Выключить: «Джарвис, шпион выкл»"
                        )
                    return

        if is_owner and is_pm and any(t in q_own for t in _SPY_OFF):
            import gc
            for obj in gc.get_objects():
                if type(obj).__name__ == "TelegramHandler":
                    obj._spy_mode   = False
                    obj._spy_chats  = set()
                    obj._spy_pending = False
                    break
            await event.reply("👁 Шпионский режим выключен, Сэр.")
            return

        _LOG_TRIGGERS = ["скинь логи", "покажи логи", "отправь логи", "лог файл"]
        if is_owner and is_pm and any(t in text.lower() for t in _LOG_TRIGGERS):
            log_path = config.LOG_FILE
            if log_path.exists() and log_path.stat().st_size > 0:
                try:
                    # Читаем последние 200 строк чтобы не превышать лимит Telegram
                    lines = log_path.read_text("utf-8", errors="replace").splitlines()
                    tail  = "\n".join(lines[-200:])
                    # Шлём как файл через bytes
                    import io
                    log_bytes = io.BytesIO(tail.encode("utf-8"))

                    _msk_now = datetime.now(timezone(timedelta(hours=3))).replace(tzinfo=None)
                    log_bytes.name = f"jarvis_{_msk_now.strftime('%d%m%Y_%H%M')}.log"
                    await self.client.send_file(
            sender_id,
            log_bytes,
            caption=f"📋 Лог Джарвиса (последние 200 строк) — {_msk_now.strftime('%d.%m.%Y %H:%M')} МСК"
                    )
                except Exception as _le:
                    await event.reply(f"Сэр, ошибка при отправке лога: {_le}")
            else:
                await event.reply("Сэр, лог-файл пуст или не найден.")
            return

        # Список бэкапов — Google Drive убран, бэкапы хранятся в Telegram

        # ── Ручной бэкап по команде ───────────────────────────
        _BACKUP_TRIGGERS = ["сделай бэкап", "сделай резервную", "создай бэкап",
                            "бэкап сейчас", "backup now", "сохрани базу"]
        if is_owner and is_pm and any(t in q_own for t in _BACKUP_TRIGGERS):
            await event.reply("⏳ Создаю бэкап, Сэр...")
            try:
                from datetime import datetime as _dt2
                _label = _dt2.now().strftime("%Y-%m-%d_%H%M")
                _json_p, _db_p, _stats = await _make_backup_files(_label)
                await self.client.send_file(
                    sender_id, _db_p,
                    caption=(
                        f"💾 Бэкап по запросу — Jarvis_{_label}.db\n"
                        f"📊 Сообщений: {_stats.get('user_messages', 0)} ЛС · "
                        f"{_stats.get('group_messages', 0)} групп\n"
                        f"📦 Размер: {_stats.get('db_kb', 0)} KB"
                    )
                )
                try:
                    import os as _os_bk
                    _os_bk.unlink(_json_p)
                    _os_bk.unlink(_db_p)
                except Exception:
                    pass
            except Exception as _be:
                await event.reply(f"❌ Сэр, ошибка бэкапа: {_be}")
            return


        if is_owner and is_pm and any(t in q_own for t in ["восстанови бэкап", "восстанови последний", "восстанови последний бэкап"]):
            await event.reply("Сэр, пришлите файл `Jarvis_ДАТА.db` в этот чат — восстановлю.")
            return

        # ── Статистика группы ────────────────────────────────
        _STATS_TRIGGERS = ["статистика группы", "статистика чата", "активность группы",
                           "сколько сообщений", "активность за неделю", "статистика за",
                           "статистика за месяц"]
        if any(t in text.lower() for t in _STATS_TRIGGERS):
            target_id = (chat_id if not is_pm else 0) or config.DEFAULT_GROUP_ID
            if not target_id:
                await event.reply("Сэр, используйте эту команду в группе или задайте DEFAULT_GROUP_ID в .env")
                return
            try:
                tl = text.lower()

                # Определяем режим: явный период или всё время
                if "за месяц" in tl:
                    _days = 30
                    _mode = "period"
                elif "за неделю" in tl or "за 7" in tl:
                    _days = 7
                    _mode = "period"
                elif "за день" in tl or "за сегодня" in tl:
                    _days = 1
                    _mode = "period"
                else:
                    _mode = "alltime"   # ручной запрос без уточнения → всё время

                if _mode == "alltime":
                    st = _jarvis_db.get_group_stats_alltime(target_id)
                    if not st or st.get("total", 0) == 0:
                        await event.reply("Сэр, данных по группе пока нет. Бот должен присутствовать в группе.")
                        return

                    top_lines = "\n".join(
                        f"  {i+1}. {u.get('sender','?')} — {u.get('cnt', 0)} сообщ."
                        for i, u in enumerate(st.get("top_users", [])[:5])
                    ) or "  Нет данных"

                    msg = (
                        f"📊 **Статистика группы (за всё время):**\n\n"
                        f"💬 Всего сообщений: {st.get('total', 0)}\n"
                        f"👤 Участников: {st.get('unique_users', 0)}\n"
                        f"🗑 Удалено за всё время: {st.get('deleted_total', 0)}\n"
                        f"📅 Период: {st.get('first_date','—')} → {st.get('last_date','—')}\n"
                        f"🔥 Самый активный день: {st.get('top_day','—')}\n\n"
                        f"🏆 Топ участников:\n{top_lines}"
                    )

                else:
                    # Статистика за конкретный период
                    st = _jarvis_db.get_group_stats(target_id, _days)
                    if not st or st.get("current", 0) == 0:
                        await event.reply(f"Сэр, за этот период данных нет.")
                        return

                    current  = st.get("current", 0)
                    previous = st.get("previous", 0)
                    deleted  = st.get("deleted", 0)

                    if previous > 0:
                        raw_change = round((current - previous) / previous * 100)
                        change = max(-999, min(999, raw_change))
                        arrow = "📈" if change >= 0 else "📉"
                        sign  = "+" if change >= 0 else ""
                        change_str = f"{sign}{change}%"
                    elif current > 0:
                        change_str = "новые данные"
                        arrow = "📈"
                    else:
                        change_str = "нет данных"
                        arrow = "📊"

                    period_name = {1: "1 день", 7: "7 дней", 30: "30 дней"}.get(_days, f"{_days} дней")
                    top_lines = "\n".join(
                        f"  {i+1}. {u.get('sender','?')} — {u.get('cnt', 0)} сообщ."
                        for i, u in enumerate(st.get("top_users", [])[:5])
                    ) or "  Нет данных"

                    msg = (
                        f"📊 **Статистика группы (за {period_name}):**\n\n"
                        f"💬 Сообщений: {current}\n"
                        f"👤 Участников: {st.get('unique_users', 0)}\n"
                        f"{arrow} Активность: {change_str} vs прошлый период\n"
                        f"🗑 Удалено за период: {deleted}\n\n"
                        f"🏆 Топ участников:\n{top_lines}"
                    )

                await event.reply(msg, parse_mode="md")
            except Exception as _se:
                await event.reply(f"Сэр, ошибка при получении статистики: {_se}")
            return

        # ── ML статистика ────────────────────────────────────
        _ML_STAT_TRIGGERS = ["статистика обучения", "ml статистика", "что узнал", "чему научился"]
        if is_owner and any(t in text.lower() for t in _ML_STAT_TRIGGERS):
            try:
                ml_stats = _jarvis_db.ml_get_stats() if hasattr(_jarvis_db, 'ml_get_stats') else {}
                patterns = ml_stats.get("patterns", 0)
                knowledge = ml_stats.get("knowledge", 0)
                sessions = ml_stats.get("sessions", 0)
                msg = (
                    f"🧠 **ML обучение Джарвиса:**\n\n"
                    f"📊 Паттернов изучено: {patterns}\n"
                    f"📚 Знаний накоплено: {knowledge}\n"
                    f"🔄 Сессий обучения: {sessions}\n"
                )
                await event.reply(msg)
            except Exception as _me:
                await event.reply(f"Сэр, ошибка ML статистики: {_me}")
            return

        # ══ ГЛАВНЫЙ ОБРАБОТЧИК ТЕКСТА ═══════════════════════════
        if not text:
            return

        # В группах реагируем если:
        # 1. Обращаются к Джарвису по имени/упоминанию
        # 2. Отвечают (reply) на сообщение бота
        if not is_pm:
            _activated, _ = self.agent.is_activated(text)
            if not _activated:
                # Проверяем — может это reply на сообщение бота
                try:
                    _reply_msg = await event.message.get_reply_message()
                    if _reply_msg and _reply_msg.sender_id == (await self.client.get_me()).id:
                        # Это reply на бота — активируем с полным текстом
                        _activated = True
                        # Передаём текст как есть (без префикса)
                        text = "Джарвис, " + text
                except Exception:
                    pass
            if not _activated:
                return

        _typing = TypingManager(self.client, event.chat_id)
        await _typing.start()
        _wait_msg = None
        try:
            _wait_msg = await event.reply("⏳ Секунду, Сэр...")
            response = await self.agent.process(
                text, sender_id=sender_id, username=username, chat_id=chat_id
            )
            if _wait_msg:
                await _wait_msg.delete()
                _wait_msg = None
            if response:
                await event.reply(response)
        except Exception as _pe:
            logger.error(f"❌ agent.process: {_pe}")
            if _wait_msg:
                try: await _wait_msg.delete()
                except: pass
        finally:
            await _typing.stop()


async def _log_critical_error(context: str, error: Exception, notify_owner: bool = True, client=None):
    """
    Логирует критическую ошибку с AI-объяснением.
    Отправляет владельцу если notify_owner=True и client передан.
    """

    _type = type(error).__name__
    _msg  = str(error)
    _trace = traceback.format_exc()[-500:]

    # Получаем AI объяснение
    explanation = await _ai_explain_error(_type, _msg, context)

    # Формируем лог
    log_line = f"❌ [{context}] {_type}: {_msg}"
    if explanation:
        log_line += f" | AI: {explanation[:150]}"
    logger.error(log_line)

    # Отправляем владельцу
    if notify_owner and client:
        try:
            owner_msg = (
                f"🚨 **Сэр, произошла ошибка**\n\n"
                f"📍 Где: {context}\n"
                f"🔴 Тип: `{_type}`\n"
                f"💬 Что: {_msg[:200]}\n"
            )
            if explanation:
                owner_msg += f"\n🤖 AI-анализ:\n{explanation}"
            await client.send_message(config.OWNER_ID, owner_msg)
        except Exception:
            pass


async def _ai_explain_error(error_type: str, error_msg: str, context: str = "") -> str:
    """AI объяснение ошибки — краткое, для лога."""
    try:
        llm = UniversalLLMConnector()
        prompt = f"Кратко (1 предложение) объясни ошибку Python: {error_type}: {error_msg}"
        result = await llm.complete([{"role": "user", "content": prompt}], max_tokens=100)
        return result.strip()
    except Exception:
        return


def print_banner(agent):
    """Печатает баннер запуска."""
    llm = getattr(agent, "llm", None)
    llm_name = llm.current_display if llm else "не настроен"
    rag = "выключен"
    tg_mode = "Бот" if config.TELEGRAM_BOT_TOKEN else "Юзер"
    print(f"JARVIS запущен | LLM: {llm_name} | RAG: {rag} | Telegram: {tg_mode}")


async def console_mode(agent):
    """Консольный режим для тестирования без Telegram."""
    print("Консольный режим. Введите сообщение (exit для выхода):")
    while True:
        try:
            user_input = input("Вы: ").strip()
            if user_input.lower() in ("exit", "quit", "выход"):
                break
            if not user_input:
                continue
            response = await agent.process(user_input, sender_id=0, username="console")
            print(f"Джарвис: {response}\n")
        except (KeyboardInterrupt, EOFError):
            break


class GroupMonitor:
    """Мониторинг групп через user-сессию."""

    def __init__(self, agent):
        self.agent = agent

    async def start(self) -> bool:
        if not config.USER_SESSION_STRING and not config.USER_SESSION_FILE:
            return False
        try:
            from telethon import TelegramClient, events
            from pathlib import Path as _P
            # Если есть StringSession — используем его, иначе файл сессии
            if config.USER_SESSION_STRING:
                from telethon.sessions import StringSession
                session = StringSession(config.USER_SESSION_STRING)
            else:
                # Убеждаемся что папка для сессии существует
                _sess_path = _P(config.USER_SESSION_FILE)
                _sess_path.parent.mkdir(parents=True, exist_ok=True)
                # Убираем расширение — Telethon сам добавит .session
                session = str(_sess_path.with_suffix(""))
            client  = TelegramClient(session, config.TELEGRAM_API_ID, config.TELEGRAM_API_HASH)
            await client.start(phone=config.TELEGRAM_PHONE or None)
            me = await client.get_me()
            logger.info(f"✅ GroupMonitor: {me.first_name} (@{me.username})")

            _bot_chat_ids: set[int] = _jarvis_db.get_bot_chat_ids()

            @client.on(events.NewMessage(chats=list(_bot_chat_ids) if _bot_chat_ids else None))
            async def on_group_msg(event):
                try:
                    if event.is_private:
                        return
                    chat_id   = event.chat_id
                    sender_id = event.sender_id or 0
                    msg       = event.message
                    text      = (msg.text or msg.message or "").strip()
                    if not text:
                        return
                    sender_name = ""
                    try:
                        sender = await event.get_sender()
                        first  = getattr(sender, "first_name", "") or ""
                        last   = getattr(sender, "last_name",  "") or ""
                        uname  = getattr(sender, "username",   "") or ""
                        sender_name = (first + " " + last).strip() or uname or str(sender_id)
                    except Exception:
                        pass
                    if msg.date:
                        msk = msg.date.replace(tzinfo=timezone.utc).astimezone(timezone(timedelta(hours=3)))
                        date_str = msk.strftime("%H:%M %d.%m.%Y МСК")
                    else:
                        date_str = ""
                    self.agent.group_logger.save(
                        chat_id=chat_id, msg_id=msg.id, sender=sender_name,
                        sender_id=sender_id, text=text, date=date_str,
                    )
                except Exception as e:
                    logger.warning(f"⚠️ GroupMonitor: ошибка [{type(e).__name__}]: {e}")

            @client.on(events.MessageDeleted())
            async def on_deleted(event):
                try:
                    chat_id = event.chat_id
                    ids = list(event.deleted_ids) if event.deleted_ids else []
                    if not ids:
                        return
                    if chat_id:
                        self.agent.group_logger.mark_deleted(chat_id, ids)
                    else:
                        self.agent.group_logger.mark_deleted_all_chats(ids)
                except Exception as e:
                    logger.warning(f"⚠️ GroupMonitor: ошибка удаления [{type(e).__name__}]: {e}")

            await client.run_until_disconnected()
            return True

        except Exception as e:
            logger.error(f"❌ GroupMonitor не запустился [{type(e).__name__}]: {e}. Проверьте USER_SESSION_STRING в .env")
            return False


async def main():
    logger.info("JARVIS ULTIMATE 2026 v4.0 запуск")
    print("Инициализация систем...")

    # ── SQLite — работает БЕЗ VPN, БЕЗ интернета ────────
    try:
        if _jarvis_db.ping():
            print(f"✅ SQLite БД готова: {_jarvis_db._path}")
        else:
            print("⚠️  SQLite: ошибка проверки")
    except Exception as _dbe:
        print(f"❌ SQLite: {_dbe}")

    # ── Проверка пустой БД — ставим флаг для уведомления ─────
    try:
        _jarvis_db._q(
            "CREATE TABLE IF NOT EXISTS _startup_flags (key TEXT PRIMARY KEY, val TEXT)"
        )
        _stats = _jarvis_db.get_db_stats()
        _total = (_stats.get("user_messages", 0)
                  + _stats.get("group_messages", 0)
                  + _stats.get("profiles", 0))
        if _total == 0 and config.OWNER_ID:
            logger.info("БД пустая — уведомим после подключения Telegram")
            _jarvis_db._q(
                "INSERT OR REPLACE INTO _startup_flags (key, val) "
                "VALUES ('db_empty_notify', '1')"
            )
        else:
            _jarvis_db._q(
                "DELETE FROM _startup_flags WHERE key='db_empty_notify'"
            )
    except Exception as _e:
        logger.warning(f"⚠️ Проверка БД: {_e}")

    # Инициализируем новые системы
    agent = JarvisAgent()
    print_banner(agent)

    if not (config.TELEGRAM_API_ID and config.TELEGRAM_API_HASH and config.TELEGRAM_API_ID != 0):
        print("⚠️  Telegram не настроен → консольный режим\n")
        await console_mode(agent)
        return

    try:
        tg      = JarvisTelegram(agent)
        monitor = GroupMonitor(agent)

        # Запускаем оба клиента параллельно:
        # — JarvisTelegram: отвечает на команды (бот или юзер)
        # — GroupMonitor:   тихо пишет в лог ВСЕ сообщения групп (только юзер)
        # Запускаем фоновое обучение
        global _background_learner, _vk_bot
        _background_learner = BackgroundLearner(_jarvis_db)
        _background_learner.start()

        # Запускаем VK бота если токен задан
        try:
            from vk_bot import create_vk_bot
            _vk_bot = create_vk_bot(agent)
            if _vk_bot:
                _vk_bot.start(asyncio.get_event_loop())
                print("🔵 VK бот запущен")
        except ImportError:
            pass
        except Exception as _vk_e:
            logger.warning(f"⚠️ VK бот не запущен: {_vk_e}")

        tasks = [asyncio.create_task(tg.start())]

        # GroupMonitor запускается только если есть user.session
        from pathlib import Path as _Path
        user_session = _Path(config.USER_SESSION_FILE)
        if user_session.exists() and user_session.stat().st_size > 100:
            tasks.append(asyncio.create_task(monitor.start()))
            print("🔍 GroupMonitor будет запущен (user.session найден)")
        else:
            print(
                "[yellow]⚠️  GroupMonitor выключен — нет user.session.\n"
                "   Запусти python create_session.py чтобы включить логирование групп.[/yellow]"
            )

        await asyncio.gather(*tasks, return_exceptions=True)

    except Exception as e:
        await _log_critical_error("Telegram клиент", e)
        print(f"⚠️  Telegram: {e}")
        print("Запускаю консольный режим...\n")
        await console_mode(agent)


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\nВыключение. До свидания, Сэр.")
        if _background_learner:
            _background_learner.stop()
        if _vk_bot:
            _vk_bot.stop()
        logger.info("JARVIS выключен")